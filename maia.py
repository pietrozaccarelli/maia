import asyncio
import json
import os
import re
import subprocess
import tkinter as tk
from tkinter import filedialog, simpledialog, messagebox
import threading
import webbrowser
from typing import List, Dict, Optional, Any
from datetime import datetime
from pathlib import Path
import pickle
import logging
import hashlib
import numpy as np
from tkinter import ttk
from PyPDF2 import PdfReader
import openpyxl
from docx import Document
from docx.shared import RGBColor
try:
    from docx2pdf import convert
except ImportError:
    convert = None
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
import pandas as pd
from ai_script_cleaner import extract_python_code

import demoji
from langchain_community.chat_models import ChatOllama
from sentence_transformers import SentenceTransformer, CrossEncoder
import torch
from wordtyper import writeword

import trafilatura
import requests
import faiss
from embedder import inquire
from codeduo import codeduo
from metadata_creator import enrich_prompt
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.enums import TA_LEFT
import glob
import time
from datetime import date
import chromadb
from chromadb.utils import embedding_functions
from tkinterdnd2 import TkinterDnD, DND_FILES 

from langchain_core.prompts import ChatPromptTemplate
try:
    from langchain_core.messages import HumanMessage, AIMessage, SystemMessage    
except ImportError as e:
    print(f"Import error: {e}")
    class HumanMessage:
        def __init__(self, content):
            self.content = content
    class AIMessage:
        def __init__(self, content):
            self.content = content
    class SystemMessage:
        def __init__(self, content):
            self.content = content

# Initialize logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)



# Initialize logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# Configuration
BASE_DIR = os.path.dirname(os.path.abspath(__file__))
OUTPUT_DIR = os.path.join(BASE_DIR, "output")
os.makedirs(OUTPUT_DIR, exist_ok=True)
EMBEDDINGS_DIR = os.path.join(BASE_DIR, "chat_embeddings")
Path(EMBEDDINGS_DIR).mkdir(exist_ok=True)
EPISODES_DIR = os.path.join(BASE_DIR, "episodes")
os.makedirs(EPISODES_DIR, exist_ok=True)
DB_PATH = os.path.join(BASE_DIR, "maia_db")
RAG_DIR = os.path.join(BASE_DIR, "rag")
os.makedirs(RAG_DIR, exist_ok=True)
MAX_MEMORY_SIZE = 1000
DEFAULT_MODEL = "gemma3n:e4b"

# --- IMPORTS FOR HF & STREAMING ---
import torch
from threading import Thread
from transformers import (
    AutoModelForCausalLM, 
    AutoTokenizer, 
    pipeline, 
    TextIteratorStreamer
)
from langchain_huggingface import HuggingFacePipeline

# --- OPTIONAL IMPORTS FOR CLOUD PROVIDERS ---
try:
    from langchain_openai import ChatOpenAI, AzureChatOpenAI
except ImportError:
    ChatOpenAI = None
    AzureChatOpenAI = None

try:
    from langchain_anthropic import ChatAnthropic
except ImportError:
    ChatAnthropic = None

try:
    from langchain_mistralai import ChatMistralAI
except ImportError:
    ChatMistralAI = None

try:
    from langchain_google_genai import ChatGoogleGenerativeAI
except ImportError:
    ChatGoogleGenerativeAI = None



# --- CONFIGURATION SECTION ---
CONFIG_FILE = os.path.join(BASE_DIR, "llm_config.json")

# Default System Prompts (Extracted from your original script logic)
DEFAULT_MAIA_PROMPT = """You are Maia, an extremely intelligent, collaborative, and attentive artificial assistant.
Your main task is to assist the user in a wide range of activities: from information research to programming, down to simple daily conversation.
Your behavior must always reflect precision, empathy, technical competence, and communicative clarity.

**FUNDAMENTAL INSTRUCTIONS**

**1) Identity**
Your name is Maia. This is your only name and your only identity.

**2) User Focus**
All responses must be personalized for the user you are interacting with.
Consider their interests, goals, and the context of previous conversations.
Avoid generic responses: always adapt tone and content to the user's needs.

**3) Operational Versatility**
You must be competent in three main areas:
a. **Research:** provide reliable information, in-depth explanations, and clear summaries.
b. **Programming:** write, correct, or explain code in various languages.
c. **Conversation:** maintain fluid, friendly, and intelligent dialogues.
Handle all activities with the same level of accuracy and professionalism.

**4) Clarity and Completeness**
Responses must be precise, structured, and easily understandable.
Address every question in its entirety: do not omit implied or secondary parts.
When appropriate, provide examples, step-by-step explanations, or comparisons to improve understanding.

**5) Context Awareness**
Take previous conversations into account to offer consistent and informed responses.
However, do not overcomplicate responses with unnecessary or redundant references.
If a topic has already been covered, you may briefly recall it to build continuity.

**6) Professional and Cordial Tone**
Always maintain a helpful, cordial, and intelligent attitude.
Show empathy and collaboration, especially when the user requests clarifications or complex assistance.

**RESPONSE GUIDELINES**

*   **Fidelity to Requests:**
    If the user asks to print or report a text, reproduce it exactly as specified, without alterations.

*   **Explicating Instructions:**
    If the user asks you to make instructions stronger or more detailed, rephrase them in a clear, actionable, and verifiable way.

*   **Attention and Focus:**
    Remain always focused on the user's specific requests.
    Avoid deviations, superfluous information, or overly verbose responses.

**RESPONSE FORMAT**

*   Provide direct and relevant responses.
*   When useful, briefly cite previous conversations or contexts to improve consistency.
*   Keep responses concise yet comprehensive, balancing brevity and depth.
*   Use fluid and natural language, appropriate to Maia's professional and friendly tone.

**FINAL GOAL**

To be a trusted assistant for the user, capable of:
*   understanding their intentions,
*   responding with precision and competence,
*   offering effective solutions,
*   and maintaining a pleasant, consistent, and constructive dialogue.

"""

DEFAULT_ROUTER_PROMPT = """You're an AI router. You must decide what command needs to be attached to the user's prompt among the following options: 
'writeword' if the prompt asks to print or save in Word; 
'google' if the prompt asks to search for extra information or current events; 
'normalprompt' in all other cases. 
Answer ONLY with the chosen command. Verbose = 0."""

DEFAULT_CODER_PROMPT = """You are a senior software developer and highly skilled python coder. 
Your task is to create a python code whose output answers the user's prompt. 
The code must be concise, specific, and optimized for execution. 
Use pandas for data manipulation. Answer ONLY with the python code."""

DEFAULT_REFINER_PROMPT = """You are a query refiner. Take the user's prompt and refine it into a better Google search query. 
The refined query should be concise, specific, and optimized for web search. 
Answer ONLY with the refined query."""

DEFAULT_MEMORY_ANALIZER_PROMPT = """Analyze the provided chat log. 
1. Identify main topics.
2. For each topic, extract a concise 'learning' detailing what the AI should remember.
(user preferences, specific procedures, or facts about the user). 
Output a JSON object with a key 'memories' containing a list of objects with keys 'topic' and 'learning'.
"""

# --- NEW: Default Prompt for Consolidator ---
DEFAULT_CONSOLIDATOR_PROMPT = """You are a memory optimizer. 
Combine the provided list of similar learnings into ONE single, concise, comprehensive learning. 
Do not lose important details, but remove redundancy. 
Output JSON: {'topic': '...', 'learning': '...'}"""

# Expanded Default Configuration with System Prompts
DEFAULT_CONFIG = {
    "main": {
        "provider": "ollama", "model": DEFAULT_MODEL, 
        "temperature": 0.7, "max_tokens": 2048, "top_p": 0.9, "api_key": "",
        "system_prompt": DEFAULT_MAIA_PROMPT
    },
    "router": {
        "provider": "ollama", "model": DEFAULT_MODEL, 
        "temperature": 0.0, "max_tokens": 128, "top_p": 0.1, "api_key": "",
        "system_prompt": DEFAULT_ROUTER_PROMPT
    },
    "refiner": {
        "provider": "ollama", "model": DEFAULT_MODEL, 
        "temperature": 0.2, "max_tokens": 512, "top_p": 0.5, "api_key": "",
        "system_prompt": DEFAULT_REFINER_PROMPT
    },
    "coder": {
        "provider": "ollama", "model": DEFAULT_MODEL, 
        "temperature": 0.2, "max_tokens": 4096, "top_p": 0.2, "api_key": "",
        "system_prompt": DEFAULT_CODER_PROMPT
    },
    "summarizer": {
        "provider": "ollama", "model": DEFAULT_MODEL, 
        "temperature": 0.0, "max_tokens": 2048, "top_p": 0.5, "api_key": "",
        "system_prompt": "Analyze the chat. Output VALID JSON only with keys 'topics' and 'summary'."
    },
    "memory_analyzer": {
        "provider": "ollama", "model": DEFAULT_MODEL, 
        "temperature": 0.1, "max_tokens": 2048, "top_p": 0.5, "api_key": "",
        "system_prompt": DEFAULT_MEMORY_ANALIZER_PROMPT
    },
    "memory_consolidator": {
        "provider": "ollama", "model": DEFAULT_MODEL, 
        "temperature": 0.1, "max_tokens": 2048, "top_p": 0.5, "api_key": "",
        "system_prompt": DEFAULT_CONSOLIDATOR_PROMPT
    },
         "directories": {
        "hf_home": "",       # Path for HuggingFace models
        "ollama_models": ""  # Path for Ollama models
    },
    "retrieval": {
        "rag_k": 3,
        "episodic_k": 12,
        "summary_k": 3,
        "google_k": 5,           
        "relevance_threshold": 2.0,  
        "consolidation_threshold": 0.8, 
        "max_history": 1000,
        "topic_truncation": 100,      
        "semantic_truncation": 600,
        "doc_rag_k": 5,             
        "doc_rag_threshold": 2.0     
    }
}


# Global Config Variable
LLM_CONFIG = DEFAULT_CONFIG

# Global Cache for Loaded HF Models
LOADED_HF_PIPELINES = {}

def load_config():
    global LLM_CONFIG
    if os.path.exists(CONFIG_FILE):
        try:
            with open(CONFIG_FILE, "r") as f:
                loaded = json.load(f)
                # Merge with default to ensure new keys exist
                for key, val in DEFAULT_CONFIG.items():
                    if key not in loaded:
                        loaded[key] = val
                    else:
                        # Ensure sub-keys exist
                        for subk, subv in val.items():
                            if subk not in loaded[key]:
                                loaded[key][subk] = subv
                LLM_CONFIG = loaded
        except:
            LLM_CONFIG = DEFAULT_CONFIG
    else:
        LLM_CONFIG = DEFAULT_CONFIG

    # --- APPLY DIRECTORY SETTINGS TO ENVIRONMENT ---
    dirs = LLM_CONFIG.get("directories", {})
    
    # 1. HuggingFace: Sets where transformers downloads/loads models
    if dirs.get("hf_home") and os.path.isdir(dirs["hf_home"]):
        os.environ["HF_HOME"] = dirs["hf_home"]
        # Also set older vars just in case
        os.environ["TRANSFORMERS_CACHE"] = dirs["hf_home"]
        logger.info(f"HF_HOME set to: {dirs['hf_home']}")

    # 2. Ollama: Sets where Ollama stores models
    # Note: If Ollama is running as a system service, this might require a service restart
    if dirs.get("ollama_models") and os.path.isdir(dirs["ollama_models"]):
        os.environ["OLLAMA_MODELS"] = dirs["ollama_models"]
        logger.info(f"OLLAMA_MODELS set to: {dirs['ollama_models']}")
        
    return LLM_CONFIG

def save_config(config):
    global LLM_CONFIG
    LLM_CONFIG = config
    with open(CONFIG_FILE, "w") as f:
        json.dump(config, f, indent=4)

load_config()

# --- HELPER: CONVERT MESSAGES TO STRING ---
def messages_to_string(messages):
    if isinstance(messages, str):
        return messages
    prompt_text = ""
    for msg in messages:
        content = msg.content
        if isinstance(msg, SystemMessage):
            prompt_text += f"System: {content}\n"
        elif isinstance(msg, HumanMessage):
            prompt_text += f"User: {content}\n"
        elif isinstance(msg, AIMessage):
            prompt_text += f"Assistant: {content}\n"
        else:
            prompt_text += f"{content}\n"
    prompt_text += "Assistant: "
    return prompt_text


# ==========================================
#       RAG MANAGER FOR DOCUMENTS
# ==========================================

class DocumentRAGManager:
    """Manages embeddings for specific documents in the ./rag folder."""
    def __init__(self, model):
        self.model = model 
        self.active_indices = {} # {filename: (faiss_index, chunks)}
        if not os.path.exists(RAG_DIR):
            os.makedirs(RAG_DIR)

    def get_file_hash(self, file_path):
        with open(file_path, "rb") as f:
            return hashlib.md5(f.read()).hexdigest()

    def embed_file(self, file_path):
        """Chunks and embeds a file, saving the index in ./rag."""
        if self.model is None:
            logger.error("Embedding model not available for Document RAG.")
            return

        file_hash = self.get_file_hash(file_path)
        index_path = os.path.join(RAG_DIR, f"{file_hash}.index")
        chunks_path = os.path.join(RAG_DIR, f"{file_hash}.pkl")

        if file_path in self.active_indices:
            return

        if os.path.exists(index_path) and os.path.exists(chunks_path):
            index = faiss.read_index(index_path)
            with open(chunks_path, 'rb') as f:
                chunks = pickle.load(f)
            self.active_indices[file_path] = (index, chunks)
            return

        ext = os.path.splitext(file_path)[1].lower()
        content = ""
        if ext == ".pdf": content = read_pdf(file_path)
        elif ext == ".docx": content = read_word(file_path)
        elif ext in [".xlsx", ".xls"]: 
            res = read_excel_enhanced(file_path)
            content = res[0] if isinstance(res, tuple) else res
        else:
            try:
                with open(file_path, "r", encoding="utf-8") as f:
                    content = f.read()
            except: return

        if not content: return

        chunks = chunk_text(content)
        # Change: Use self.model instead of global embed_model
        embeddings = self.model.encode(chunks, convert_to_numpy=True)
        
        index = faiss.IndexFlatL2(embeddings.shape[1])
        index.add(embeddings.astype('float32'))

        faiss.write_index(index, index_path)
        with open(chunks_path, 'wb') as f:
            pickle.dump(chunks, f)
            
        self.active_indices[file_path] = (index, chunks)

    def query_active_files(self, query, k=None): # k is optional now, defaults to config
        """Searches only the files selected by the user."""
        
        # --- NEW: Get settings from Config ---
        retrieval_conf = LLM_CONFIG.get("retrieval", DEFAULT_CONFIG["retrieval"])
        
        # Use config if k is not explicitly passed (or override if you prefer config always)
        if k is None:
            k = int(retrieval_conf.get("doc_rag_k", 5))
            
        threshold = float(retrieval_conf.get("doc_rag_threshold", 2.0))
        # -------------------------------------

        context = ""
        for file_path, (index, chunks) in self.active_indices.items():
            # Pass the threshold to query_index
            results = query_index(query, chunks, index, k=k, threshold=threshold)
            if results:
                context += f"\n--- FROM DOCUMENT: {os.path.basename(file_path)} ---\n"
                context += "\n".join(results) + "\n"
        return context





class MaiaAttachmentsDialog(tk.Toplevel):
    def __init__(self, master, existing_files, existing_rag):
        super().__init__(master)
        self.title("M.A.I.A. Attachments & RAG Manager")
        self.geometry("1400x1100")
        self.configure(bg="#1e1e1e")

        # PERSISTENCE LOGIC
        # We work on copies so we can "Cancel" if we want, 
        # but CrewChatUI will update its master lists on "Confirm"
        self.attachments = list(existing_files) 
        self.rag_active = set(existing_rag)
        self.confirmed = False

        # Make modal
        self.transient(master)
        self.grab_set()

        # UI Styling
        lbl_style = {"bg": "#1e1e1e", "fg": "white", "font": ("Segoe UI", 10)}

        # ===== FILE LIST =====
        tk.Label(self, text="Current Session Attachments:", **lbl_style).pack(pady=(10, 0))
        list_frame = tk.Frame(self, bg="#1e1e1e")
        list_frame.pack(fill="both", expand=True, padx=15, pady=5)

        self.listbox = tk.Listbox(
            list_frame, bg="#2d2d2d", fg="white", 
            selectmode=tk.SINGLE, font=("Consolas", 10),
            borderwidth=0, highlightthickness=1, highlightbackground="#444"
        )
        self.listbox.pack(side="left", fill="both", expand=True)

        scrollbar = ttk.Scrollbar(list_frame, command=self.listbox.yview)
        scrollbar.pack(side="right", fill="y")
        self.listbox.config(yscrollcommand=scrollbar.set)

        # ===== RAG CONTROLS =====
        rag_frame = tk.Frame(self, bg="#1e1e1e")
        rag_frame.pack(fill="x", padx=15, pady=5)

        tk.Button(rag_frame, text="Toggle RAG (Green Mode)", bg="#2ecc71", fg="white",
                  command=self.toggle_rag_selection, relief="flat", padx=10).pack(side="left")
        
        tk.Label(rag_frame, text="Green = Activated for Semantic Search", **lbl_style).pack(side="left", padx=10)

        # ===== ACTION BUTTONS =====
        btn_frame = tk.Frame(self, bg="#1e1e1e")
        btn_frame.pack(fill="x", padx=15, pady=10)

        tk.Button(btn_frame, text="Remove Selected", bg="#e74c3c", fg="white",
                  command=self.remove_selected, relief="flat", padx=10).pack(side="left")

        tk.Button(btn_frame, text="CONFIRM & EXIT", bg="#3498db", fg="white", width=20,
                  command=self.confirm, relief="flat", font=("Segoe UI", 10, "bold")).pack(side="right")

        # ===== DROP AREA =====
        # This area handles both Drag & Drop AND Click-to-Browse
        self.drop_area = tk.Label(
            self, text="DRAG FILES HERE TO ADD TO SESSION\n(or click here to browse)",
            relief="ridge", borderwidth=2, height=4, bg="#252525", fg="#00aeff", 
            font=("Segoe UI", 9, "bold"), cursor="hand2"
        )
        self.drop_area.pack(fill="x", padx=15, pady=10)

        # Register Drop Target (This requires DND_FILES to be imported)
        self.drop_area.drop_target_register(DND_FILES)
        self.drop_area.dnd_bind("<<Drop>>", self.on_drop)
        
        # Click to browse logic
        self.drop_area.bind("<Button-1>", self.browse_files)

        # Initial Load of existing files
        self.refresh_listbox()

    def refresh_listbox(self):
        """Clears and repaints the listbox based on current state."""
        self.listbox.delete(0, tk.END)
        for path in self.attachments:
            fname = os.path.basename(path)
            self.listbox.insert(tk.END, fname)
            idx = self.listbox.size() - 1
            if path in self.rag_active:
                # Highlight RAG-active files in Green
                self.listbox.itemconfig(idx, {'bg': '#27ae60', 'fg': 'white'})

    def toggle_rag_selection(self):
        """Toggles the 'Green' RAG mode for the selected file."""
        selection = self.listbox.curselection()
        if not selection: 
            messagebox.showinfo("Selection Required", "Please click a file in the list first.")
            return
        idx = selection[0]
        path = self.attachments[idx]
        
        if path in self.rag_active:
            self.rag_active.remove(path)
        else:
            self.rag_active.add(path)
        self.refresh_listbox()

    def on_drop(self, event):
        """Logic for files dropped onto the label."""
        files = self.tk.splitlist(event.data)
        self.add_files(files)

    def browse_files(self, event=None):
        """Logic for clicking the label to open a file dialog."""
        files = filedialog.askopenfilenames(parent=self, title="Select files to attach")
        if files: 
            self.add_files(files)

    def add_files(self, files):
        """Adds files to the internal list and refreshes UI."""
        for path in files:
            path = os.path.normpath(path)
            if path not in self.attachments:
                self.attachments.append(path)
        self.refresh_listbox()

    def remove_selected(self):
        """Removes selected file from both standard and RAG lists."""
        selection = self.listbox.curselection()
        if not selection: return
        idx = selection[0]
        path = self.attachments.pop(idx)
        if path in self.rag_active:
            self.rag_active.remove(path)
        self.refresh_listbox()

    def confirm(self):
        """Marks as confirmed and closes."""
        self.confirmed = True
        self.destroy()



# ==========================================
#       MODEL MEMORY MANAGEMENT
# ==========================================

def get_ollama_loaded_models():
    """
    Queries the Ollama API to see what is actually loaded in VRAM.
    Returns a list of normalized model names.
    """
    try:
        response = requests.get("http://localhost:11434/api/ps", timeout=2)
        if response.status_code == 200:
            data = response.json()
            # Extract model names (e.g., 'gemma2:latest')
            # We strip the tag if the config doesn't use it, but keeping exact match is safer
            models = [m.get('name', '') for m in data.get('models', [])]
            return models
    except Exception:
        return []
    return []

def unload_specific_model(provider, model_name):
    """
    Unloads a specific model based on provider type.
    Returns string message describing result and success status (bool).
    """
    import gc
    import requests
    
    provider = provider.lower()
    
    # --- HUGGINGFACE LOGIC ---
    if provider == "huggingface":
        global LOADED_HF_PIPELINES
        if model_name in LOADED_HF_PIPELINES:
            del LOADED_HF_PIPELINES[model_name]
            gc.collect()
            try:
                import torch
                if torch.cuda.is_available():
                    torch.cuda.empty_cache()
                    torch.cuda.ipc_collect()
                elif torch.backends.mps.is_available():
                    torch.mps.empty_cache()
            except:
                pass
            return f"Unloaded HF Model: {model_name}", True
        else:
            return f"HF Model {model_name} was not in memory.", False

    # --- OLLAMA LOGIC ---
    elif provider == "ollama":
        try:
            # Send keep_alive=0 to force unload immediately
            requests.post(
                "http://localhost:11434/api/generate",
                json={"model": model_name, "keep_alive": 0},
                timeout=2
            )
            return f"Unload signal sent to Ollama for: {model_name}", True
        except Exception as e:
            return f"Ollama Error: {e}", False

    # --- CLOUD/API LOGIC ---
    else:
        return f"Provider '{provider}' is cloud-based. No local RAM to free.", False

def open_unload_selector(parent_win):
    """
    UI to view loaded models grouped by instance and unload them.
    Shows which components (main, coder, etc.) use which model.
    """
    win = tk.Toplevel(parent_win)
    win.title("Active Model Memory Manager")
    win.geometry("2000x1200") 
    win.configure(bg="#1e1e1e")

    # Header
    header_frame = tk.Frame(win, bg="#1e1e1e")
    header_frame.pack(fill="x", pady=15, padx=20)
    
    tk.Label(header_frame, text="Memory Manager", 
             bg="#1e1e1e", fg="white", font=("Segoe UI", 16, "bold")).pack(anchor="w")
    tk.Label(header_frame, text="View active models, see where they are used, and unload them to free RAM/VRAM.", 
             bg="#1e1e1e", fg="#aaaaaa", font=("Segoe UI", 10)).pack(anchor="w")

    # --- 1. GATHER DATA ---
    # We want to group by unique (Provider, ModelName) tuples
    # Structure: { "ollama::gemma2": { "provider": "ollama", "name": "gemma2", "components": ["main", "coder"] } }
    
    model_map = {}
    
    # --- DYNAMIC COMPONENT DISCOVERY ---
    # Find all keys in the config that look like an LLM component
    components_list = [
        key for key, value in LLM_CONFIG.items() 
        if isinstance(value, dict) and "provider" in value and "model" in value
    ]
    # -------------------------------------

    for comp in components_list:
        conf = LLM_CONFIG.get(comp, DEFAULT_CONFIG.get("main", {}))
        provider = conf.get("provider", "ollama").lower()
        model = conf.get("model", "default")
        
        # Create unique key
        key = f"{provider}::{model}"
        
        if key not in model_map:
            model_map[key] = {
                "provider": provider,
                "name": model,
                "components": []
            }
        model_map[key]["components"].append(comp.upper())

    # Fetch actual status
    ollama_loaded = get_ollama_loaded_models() # List of strings
    
    # --- 2. BUILD UI ---
    container = tk.Frame(win, bg="#1e1e1e")
    container.pack(fill="both", expand=True, padx=20, pady=10)

    # Scrollbar setup
    canvas = tk.Canvas(container, bg="#1e1e1e", highlightthickness=0)
    scrollbar = ttk.Scrollbar(container, orient="vertical", command=canvas.yview)
    scrollable_frame = tk.Frame(canvas, bg="#1e1e1e")

    scrollable_frame.bind("<Configure>", lambda e: canvas.configure(scrollregion=canvas.bbox("all")))
    frame_id = canvas.create_window((0, 0), window=scrollable_frame, anchor="nw")
    
    def on_canvas_configure(event):
        canvas.itemconfig(frame_id, width=event.width)
    canvas.bind("<Configure>", on_canvas_configure)

    canvas.configure(yscrollcommand=scrollbar.set)
    canvas.pack(side="left", fill="both", expand=True)
    scrollbar.pack(side="right", fill="y")

    # --- Headers ---
    cols_frame = tk.Frame(scrollable_frame, bg="#2d2d2d", pady=10)
    cols_frame.pack(fill="x", pady=(0, 5))
    
    tk.Label(cols_frame, text="MODEL NAME", bg="#2d2d2d", fg="#4ea1ff", font=("Segoe UI", 10, "bold"), width=25, anchor="w").pack(side="left", padx=10)
    tk.Label(cols_frame, text="PROVIDER", bg="#2d2d2d", fg="#4ea1ff", font=("Segoe UI", 10, "bold"), width=15, anchor="w").pack(side="left", padx=10)
    tk.Label(cols_frame, text="USED BY TASKS", bg="#2d2d2d", fg="#4ea1ff", font=("Segoe UI", 10, "bold"), width=30, anchor="w").pack(side="left", padx=10)
    tk.Label(cols_frame, text="STATUS", bg="#2d2d2d", fg="#4ea1ff", font=("Segoe UI", 10, "bold"), width=15, anchor="w").pack(side="left", padx=10)
    tk.Label(cols_frame, text="ACTION", bg="#2d2d2d", fg="#4ea1ff", font=("Segoe UI", 10, "bold"), width=15, anchor="e").pack(side="right", padx=20)

    # --- Rows ---
    for key, data in model_map.items():
        prov = data["provider"]
        name = data["name"]
        comps = ", ".join(data["components"])
        
        # Determine Status
        is_loaded = False
        is_cloud = False
        
        if prov == "huggingface":
            if name in LOADED_HF_PIPELINES:
                is_loaded = True
        elif prov == "ollama":
            # Simple fuzzy check: if config name is in the 'ps' name (e.g. 'gemma2' in 'gemma2:latest')
            for active in ollama_loaded:
                if name in active:
                    is_loaded = True
                    break
        elif prov in ["openai", "anthropic", "azure", "mistral", "gemini"]:
            is_cloud = True

        # UI Row
        row = tk.Frame(scrollable_frame, bg="#252525", pady=15, padx=5)
        row.pack(fill="x", pady=2)

        # 1. Name
        tk.Label(row, text=name, bg="#252525", fg="white", font=("Segoe UI", 10, "bold"), width=25, anchor="w").pack(side="left", padx=10)
        
        # 2. Provider
        tk.Label(row, text=prov.upper(), bg="#252525", fg="#cccccc", width=15, anchor="w").pack(side="left", padx=10)
        
        # 3. Components
        tk.Label(row, text=comps, bg="#252525", fg="#aaaaaa", width=30, anchor="w", wraplength=300, justify="left").pack(side="left", padx=10)

        # 4. Status Label
        status_text = "CLOUD" if is_cloud else ("LOADED (Active)" if is_loaded else "Unloaded")
        status_fg = "#3498db" if is_cloud else ("#2ecc71" if is_loaded else "#7f8c8d")
        
        lbl_status = tk.Label(row, text=status_text, bg="#252525", fg=status_fg, font=("Segoe UI", 9, "bold"), width=15, anchor="w")
        lbl_status.pack(side="left", padx=10)

        # 5. Action Button
        btn_state = "normal" if (is_loaded and not is_cloud) else "disabled"
        btn_bg = "#e74c3c" if (is_loaded and not is_cloud) else "#444444"
        btn_txt = "UNLOAD NOW"
        
        # Define button command with closure to capture variables
        def perform_unload(p=prov, m=name, lbl=lbl_status, btn_ref=None):
            msg, success = unload_specific_model(p, m)
            if success:
                lbl.config(text="Unloaded", fg="#7f8c8d")
                if btn_ref:
                    btn_ref.config(state="disabled", bg="#444444")
                messagebox.showinfo("Unload Result", msg)
            else:
                messagebox.showerror("Error", msg)

        btn = tk.Button(row, text=btn_txt, state=btn_state, bg=btn_bg, fg="white",
                        font=("Segoe UI", 9, "bold"), width=15)
        
        # Configure command to pass the button reference itself
        btn.config(command=lambda p=prov, m=name, l=lbl_status, b=btn: perform_unload(p, m, l, b))
        btn.pack(side="right", padx=10)

    # Footer
    tk.Button(win, text="Close", command=win.destroy, bg="#34495e", fg="white", width=20, pady=5).pack(pady=15)

# ==========================================
#       MODULAR LLM ARCHITECTURE
# ==========================================

class UniversalLLMWrapper:
    """Base class defining the interface for all providers."""
    def __init__(self, model, temperature=0.7, max_tokens=1024, top_p=0.9, api_key=None, system_prompt=None, endpoint=None, api_version=None):
        self.model = model
        self.temperature = temperature
        self.max_tokens = max_tokens
        self.top_p = top_p
        self.api_key = api_key
        self.system_prompt = system_prompt
        self.endpoint = endpoint
        self.api_version = api_version 
        self.client = self._setup_client()

    def _setup_client(self):
        raise NotImplementedError

    def _prepare_messages(self, messages):
        """
        Injects the configured system prompt into the message list.
        If the code passed a SystemMessage (e.g., hardcoded agent logic),
        the Configured System Prompt overrides it to allow the UI to control behavior.
        """
        # Ensure we are working with a list of LangChain messages
        if not isinstance(messages, list):
            # If it's a raw string, we can't easily inject system prompt without context, 
            # but usually this wrapper receives lists.
            return messages

        # Create a shallow copy to avoid modifying the original list in memory everywhere
        msgs_copy = messages[:]

        if self.system_prompt and self.system_prompt.strip():
            # Check if the first message is already a SystemMessage
            if msgs_copy and isinstance(msgs_copy[0], SystemMessage):
                # OVERRIDE the hardcoded system message with the Configured one
                msgs_copy[0] = SystemMessage(content=self.system_prompt)
            else:
                # PREPEND the system message
                msgs_copy.insert(0, SystemMessage(content=self.system_prompt))
        
        return msgs_copy

    async def astream(self, messages):
        """Standardized async generator yielding string chunks."""
        raise NotImplementedError

    def invoke(self, messages):
        """Standardized sync invocation returning a string."""
        raise NotImplementedError

class OllamaWrapper(UniversalLLMWrapper):
    def _setup_client(self):
        # Default local endpoint if none provided
        base_url = self.endpoint if self.endpoint and self.endpoint.strip() else "http://localhost:11434"
        
        # --- FIX: Sanitize URL to prevent double-slash 404 errors ---
        base_url = base_url.rstrip("/")
        
        # Prepare headers
        headers = {}
        if self.api_key and self.api_key.strip():
            headers["Authorization"] = f"Bearer {self.api_key}"
            headers["X-API-Key"] = self.api_key

        return ChatOllama(
            model=self.model,
            temperature=self.temperature,
            num_ctx=4096, 
            top_p=self.top_p,
            base_url=base_url,
            headers=headers 
        )

    async def astream(self, messages):
        final_msgs = self._prepare_messages(messages)
        try:
            async for chunk in self.client.astream(final_msgs):
                yield chunk.content
        except Exception as e:
            # Catch 404 specifically to give helpful advice
            if "404" in str(e):
                yield f"[Error: Ollama returned 404. Cause: Model '{self.model}' not found OR Invalid Endpoint. Try running 'ollama pull {self.model}' in your terminal.]"
            else:
                yield f"[Error: {str(e)}]"

    def invoke(self, messages):
        final_msgs = self._prepare_messages(messages)
        try:
            response = self.client.invoke(final_msgs)
            return response.content
        except Exception as e:
            if "404" in str(e):
                return f"[Error: Ollama returned 404. Cause: Model '{self.model}' not found OR Invalid Endpoint. Try running 'ollama pull {self.model}' in your terminal.]"
            return f"[Error: {str(e)}]"

class OpenAIWrapper(UniversalLLMWrapper):
    def _setup_client(self):
        if not ChatOpenAI: raise ImportError("langchain_openai missing")
        return ChatOpenAI(
            model=self.model, temperature=self.temperature,
            max_tokens=self.max_tokens, model_kwargs={"top_p": self.top_p},
            api_key=self.api_key
        )

    async def astream(self, messages):
        final_msgs = self._prepare_messages(messages)
        async for chunk in self.client.astream(final_msgs):
            yield chunk.content

    def invoke(self, messages):
        final_msgs = self._prepare_messages(messages)
        return self.client.invoke(final_msgs).content

class AnthropicWrapper(UniversalLLMWrapper):
    def _setup_client(self):
        if not ChatAnthropic: raise ImportError("langchain_anthropic missing")
        return ChatAnthropic(
            model=self.model, temperature=self.temperature,
            max_tokens=self.max_tokens, top_p=self.top_p, api_key=self.api_key
        )

    async def astream(self, messages):
        final_msgs = self._prepare_messages(messages)
        async for chunk in self.client.astream(final_msgs):
            yield chunk.content

    def invoke(self, messages):
        final_msgs = self._prepare_messages(messages)
        return self.client.invoke(final_msgs).content

class MistralWrapper(UniversalLLMWrapper):
    def _setup_client(self):
        if not ChatMistralAI: raise ImportError("langchain_mistralai missing")
        return ChatMistralAI(
            model=self.model, temperature=self.temperature,
            max_tokens=self.max_tokens, top_p=self.top_p, api_key=self.api_key
        )

    async def astream(self, messages):
        final_msgs = self._prepare_messages(messages)
        async for chunk in self.client.astream(final_msgs):
            yield chunk.content

    def invoke(self, messages):
        final_msgs = self._prepare_messages(messages)
        return self.client.invoke(final_msgs).content

class GeminiWrapper(UniversalLLMWrapper):
    def _setup_client(self):
        if not ChatGoogleGenerativeAI: 
            raise ImportError("langchain_google_genai missing. Pip install langchain-google-genai")
        
        return ChatGoogleGenerativeAI(
            model=self.model,
            temperature=self.temperature,
            max_output_tokens=self.max_tokens, # Google specific param name
            top_p=self.top_p,
            google_api_key=self.api_key
        )

    async def astream(self, messages):
        final_msgs = self._prepare_messages(messages)
        async for chunk in self.client.astream(final_msgs):
            yield chunk.content

    def invoke(self, messages):
        final_msgs = self._prepare_messages(messages)
        return self.client.invoke(final_msgs).content

class AzureOpenAIWrapper(UniversalLLMWrapper):
    def _setup_client(self):
        if not AzureChatOpenAI: 
            raise ImportError("langchain_openai missing. Pip install langchain-openai")

        # 1. Endpoint Logic
        endpoint = self.endpoint
        if not endpoint or not endpoint.strip():
            endpoint = os.environ.get("AZURE_OPENAI_ENDPOINT", "")

        # Legacy Fallback for "ENDPOINT|KEY"
        key = self.api_key
        if (not endpoint or not endpoint.strip()) and "|" in key:
            parts = key.split("|", 1)
            endpoint = parts[0].strip()
            key = parts[1].strip()

        if not endpoint:
            raise ValueError("Azure requires an Endpoint URL.")

        # 2. API Version Logic
        # Use configured version, or default to stable 2023-05-15 if empty
        current_api_version = self.api_version
        if not current_api_version or not current_api_version.strip():
            current_api_version = "2023-05-15"

        return AzureChatOpenAI(
            azure_deployment=self.model,
            openai_api_version=current_api_version, # Dynamic Version
            azure_endpoint=endpoint,
            api_key=key,
            temperature=self.temperature,
            max_tokens=self.max_tokens,
            model_kwargs={"top_p": self.top_p}
        )

    async def astream(self, messages):
        final_msgs = self._prepare_messages(messages)
        async for chunk in self.client.astream(final_msgs):
            yield chunk.content

    def invoke(self, messages):
        final_msgs = self._prepare_messages(messages)
        return self.client.invoke(final_msgs).content

class LocalHFWrapper(UniversalLLMWrapper):
    def _setup_client(self):
        return self._get_or_load_pipeline()

    def _get_or_load_pipeline(self):
        global LOADED_HF_PIPELINES
        if self.model in LOADED_HF_PIPELINES:
            logger.info(f"Using cached pipeline for {self.model}")
            return LOADED_HF_PIPELINES[self.model]

        logger.info(f"Loading HF Model: {self.model}...")
        
        # Retrieve custom cache directory
        custom_cache = LLM_CONFIG.get("directories", {}).get("hf_home", "")
        if not custom_cache or not os.path.isdir(custom_cache):
            custom_cache = None # Fallback to default if invalid

        try:
            # Pass cache_dir to tokenizer
            tokenizer = AutoTokenizer.from_pretrained(
                self.model, 
                cache_dir=custom_cache
            )
            if tokenizer.pad_token_id is None:
                tokenizer.pad_token_id = tokenizer.eos_token_id

            # Pass cache_dir to model
            model = AutoModelForCausalLM.from_pretrained(
                self.model, 
                torch_dtype="auto", 
                device_map="auto",
                cache_dir=custom_cache
            )
            
            pipe = pipeline("text-generation", model=model, tokenizer=tokenizer)
            LOADED_HF_PIPELINES[self.model] = pipe
            return pipe
        except Exception as e:
            logger.error(f"Failed to load HF model: {e}")
            raise e

    async def astream(self, messages):
        # 1. CONVERT LangChain messages to a list of dicts
        formatted_messages = []
        for m in messages:
            if isinstance(m, SystemMessage): role = "system"
            elif isinstance(m, HumanMessage): role = "user"
            else: role = "assistant"
            formatted_messages.append({"role": role, "content": m.content})

        # 2. APPLY THE CHAT TEMPLATE (This adds the <|im_start|> and <|im_end|> tags)
        prompt_text = self.client.tokenizer.apply_chat_template(
            formatted_messages, 
            tokenize=False, 
            add_generation_prompt=True
        )

        # 3. SET UP THE STREAMER
        tokenizer = self.client.tokenizer
        streamer = TextIteratorStreamer(tokenizer, skip_prompt=True, skip_special_tokens=True)
        
        # 4. DEFINE GENERATION PARAMETERS
        # We must explicitly tell the model that <|im_end|> or <|endoftext|> means STOP
        generation_kwargs = {
            "text_inputs": prompt_text,
            "max_new_tokens": self.max_tokens,
            "temperature": self.temperature,
            "top_p": self.top_p,
            "do_sample": True if self.temperature > 0 else False,
            "streamer": streamer,
            "return_full_text": False,
            # Ensure the model stops at the official End-of-Turn token
            "eos_token_id": tokenizer.eos_token_id, 
            "pad_token_id": tokenizer.pad_token_id
        }

        # 5. START GENERATION IN A THREAD
        thread = Thread(target=self.client, kwargs=generation_kwargs)
        thread.start()

        # 6. YIELD CHUNKS BUT WATCH FOR HALLUCINATED "USER:" HEADERS
        for new_text in streamer:
            # Safety break: If the model ignores the EOS token and 
            # starts writing "User:" manually, cut it off.
            if "User:" in new_text or "--- END OF RESPONSE ---" in new_text:
                break
            yield new_text
            
        thread.join()

    
    def invoke(self, messages):
        # Convert LangChain messages to standard dicts
        formatted_messages = []
        for m in messages:
            if isinstance(m, SystemMessage): role = "system"
            elif isinstance(m, HumanMessage): role = "user"
            else: role = "assistant"
            formatted_messages.append({"role": role, "content": m.content})

        # Use the tokenizer's official template!
        prompt_text = self.client.tokenizer.apply_chat_template(
            formatted_messages, 
            tokenize=False, 
            add_generation_prompt=True
        )
        
        # Now run the model with the proper stop token
        result = self.client(
            prompt_text, 
            max_new_tokens=self.max_tokens, 
            temperature=self.temperature,
            stop_strings=["<|im_end|>", "<|endoftext|>"] # Tell it where to stop!
        )
        return result[0]['generated_text']




# --- UNIVERSAL CLIENT FACTORY ---

def get_universal_client(component_name: str) -> UniversalLLMWrapper:
    config = LLM_CONFIG.get(component_name, DEFAULT_CONFIG["main"])
    
    provider = config.get("provider", "ollama").lower()
    model = config.get("model", DEFAULT_MODEL)
    api_key = config.get("api_key", "")
    system_prompt = config.get("system_prompt", "")
    
    # --- NEW: Read Endpoint ---
    endpoint = config.get("endpoint", "") 
    # --------------------------
    api_version = config.get("api_version", "")
    temperature = config.get("temperature", 0.7)
    max_tokens = config.get("max_tokens", 1024)
    top_p = config.get("top_p", 0.9)

    # Dictionary of classes
    wrappers = {
        "ollama": OllamaWrapper,
        "openai": OpenAIWrapper,
        "anthropic": AnthropicWrapper,
        "mistral": MistralWrapper,
        "huggingface": LocalHFWrapper,
        "gemini": GeminiWrapper,
        "azure": AzureOpenAIWrapper
    }

    wrapper_class = wrappers.get(provider, OllamaWrapper)
    
    try:
        return wrapper_class(
            model=model, temperature=temperature, 
            max_tokens=max_tokens, top_p=top_p, 
            api_key=api_key, system_prompt=system_prompt,
            endpoint=endpoint,
            api_version=api_version
        )
    except Exception as e:
        logger.error(f"Error creating client for {component_name}: {e}")
        return None

# --- BACKWARD COMPATIBILITY / ALIAS ---
def get_llm_client(component_name):
    """Alias for legacy calls in the script."""
    return get_universal_client(component_name)

# --- UNIFIED STREAMING HELPER (Simplified) ---
async def unified_astream(client, messages):
    """
    Since all wrappers now implement 'astream' uniformly, 
    this helper is much simpler but kept for script compatibility.
    """
    if client is None:
        yield "[Error: Client failed to initialize]"
        return

    try:
        # All our wrappers have a standard astream returning strings
        async for text in client.astream(messages):
            yield text
    except Exception as e:
        logger.error(f"Streaming Error: {e}")
        yield f"[Error: {str(e)}]"

def validate_ollama_model(model_name):
    """Checks if model exists locally; attempts to pull if not."""
    # Only runs for local ollama (no endpoint in config)
    try:
        # Check if Ollama is running
        subprocess.check_output(["curl", "http://localhost:11434"], stderr=subprocess.STDOUT, shell=True)
    except Exception:
        # Ollama not running or not reachable via curl
        return 

    try:
        # Check list of models
        output = subprocess.check_output(["ollama", "list"], text=True, shell=True)
        if model_name not in output:
            print(f"Model '{model_name}' not found locally. Attempting background pull...")
            # Attempt to pull in a separate non-blocking process
            subprocess.Popen(["ollama", "pull", model_name])
            messagebox.showinfo("Ollama Manager", f"Model '{model_name}' was not found.\n\nA background download has started.\nPlease wait a few moments before chatting.")
    except Exception as e:
        print(f"Validation error: {e}")

#==============================================

#CLIENT = get_llm_client("main")
#if CLIENT is None:
    # Fallback to avoid crash if main config is bad, allows UI to load
    #print("Warning: Main Client failed to load. Using basic fallback or None.")









SUMMARY_FOLDER = os.path.join(BASE_DIR, "chat_summaries") 
RELEVANCE_THRESHOLD = 2.0 




# Global variables
ROOT = None
inquiring = False
dataframe_reader_mode = False
routing = False
pinned = False



#============================================================
#                 FUNZIONI
#============================================================

def read_pdf(file_path):
    """Read text from a PDF file."""
    try:
        with open(file_path, 'rb') as file:
            reader = PdfReader(file)
            text = ""
            for page in reader.pages:
                text += page.extract_text()
        return text
    except Exception as e:
        logger.error(f"Error reading PDF file: {e}")
        return None

def analyze_excel_structure(file_path):
    """Analyze Excel file structure and provide metadata for better LLM processing."""
    try:
        workbook = openpyxl.load_workbook(file_path, data_only=True)
        structure = {
            "file_name": os.path.basename(file_path),
            "sheets": {},
            "total_rows": 0,
            "total_cols": 0,
            "data_types_detected": set(),
            "has_formulas": False,
            "summary": ""
        }
        for sheet_name in workbook.sheetnames:
            worksheet = workbook[sheet_name]
            sheet_info = {
                "name": sheet_name,
                "dimensions": f"{worksheet.max_row}x{worksheet.max_column}",
                "has_headers": False,
                "column_types": {},
                "sample_data": [],
                "key_statistics": {}
            }
            first_row = [cell.value for cell in worksheet[1]]
            if any(isinstance(val, str) and val and not str(val).isdigit() for val in first_row if val):
                sheet_info["has_headers"] = True
                sheet_info["headers"] = [str(val) if val else f"Column_{i+1}" for i, val in enumerate(first_row)]
            for col_idx in range(1, min(worksheet.max_column + 1, 21)):
                col_data = []
                for row_idx in range(2 if sheet_info["has_headers"] else 1, min(worksheet.max_row + 1, 101)):
                    cell_value = worksheet.cell(row=row_idx, column=col_idx).value
                    if cell_value is not None:
                        col_data.append(cell_value)
                if col_data:
                    col_name = sheet_info.get("headers", [f"Column_{col_idx}"])[col_idx-1] if col_idx <= len(sheet_info.get("headers", [])) else f"Column_{col_idx}"
                    if all(isinstance(val, (int, float)) for val in col_data[:10]):
                        sheet_info["column_types"][col_name] = "numeric"
                        if len(col_data) > 1:
                            sheet_info["key_statistics"][col_name] = {
                                "type": "numeric",
                                "min": min(col_data),
                                "max": max(col_data),
                                "avg": sum(col_data) / len(col_data)
                            }
                    elif all(isinstance(val, str) for val in col_data[:10]):
                        sheet_info["column_types"][col_name] = "text"
                        unique_vals = list(set(col_data[:20]))
                        sheet_info["key_statistics"][col_name] = {
                            "type": "text",
                            "unique_values": len(set(col_data)),
                            "sample_values": unique_vals[:5]
                        }
                    else:
                        sheet_info["column_types"][col_name] = "mixed"
            for row_idx in range(1, min(6, worksheet.max_row + 1)):
                row_data = []
                for col_idx in range(1, min(worksheet.max_column + 1, 11)):
                    cell_value = worksheet.cell(row=row_idx, column=col_idx).value
                    row_data.append(str(cell_value) if cell_value is not None else "")
                sheet_info["sample_data"].append(row_data)
            structure["sheets"][sheet_name] = sheet_info
            structure["total_rows"] += worksheet.max_row
            structure["total_cols"] = max(structure["total_cols"], worksheet.max_column)
        return structure
    except Exception as e:
        logger.error(f"Error analyzing Excel structure: {e}")
        return None

def create_excel_summary_for_llm(structure):
    """Create a concise, LLM-optimized summary of Excel data."""
    if not structure:
        return "Unable to analyze Excel file structure."
    summary = f"EXCEL FILE ANALYSIS: {structure['file_name']}\n"
    summary += f"Total sheets: {len(structure['sheets'])}, Total rows: {structure['total_rows']}\n"
    for sheet_name, sheet_info in structure['sheets'].items():
        summary += f"SHEET: '{sheet_name}' ({sheet_info['dimensions']})\n"
        if sheet_info.get('has_headers') and sheet_info.get('headers'):
            summary += f"Headers: {', '.join(sheet_info['headers'][:10])}"
            if len(sheet_info['headers']) > 10:
                summary += f" ... and {len(sheet_info['headers']) - 10} more"
            summary += "\n"
        if sheet_info['column_types']:
            type_summary = {}
            for col, col_type in sheet_info['column_types'].items():
                if col_type not in type_summary:
                    type_summary[col_type] = []
                type_summary[col_type].append(col)
            summary += "Data types: "
            for dtype, cols in type_summary.items():
                summary += f"{dtype.upper()}({len(cols)} cols) "
            summary += "\n"
        if sheet_info['key_statistics']:
            summary += "Key insights:\n"
            for col, stats in list(sheet_info['key_statistics'].items())[:3]:
                if stats['type'] == 'numeric':
                    summary += f"    {col}: Range {stats['min']}-{stats['max']}, Avg: {stats['avg']:.2f}\n"
                elif stats['type'] == 'text':
                    sample_vals = ', '.join(str(v) for v in stats['sample_values'][:3])
                    summary += f"    {col}: {stats['unique_values']} unique values (e.g., {sample_vals})\n"
        if sheet_info['sample_data'][:3]:
            summary += "Sample data:\n"
            for i, row in enumerate(sheet_info['sample_data'][:3]):
                row_preview = ' | '.join(str(cell)[:20] for cell in row[:5])
                summary += f"  Row {i+1}: {row_preview}\n"
        summary += "\n"
    return summary

def read_excel_enhanced(file_path):
    """Enhanced Excel reading with structure analysis and LLM optimization."""
    try:
        structure = analyze_excel_structure(file_path)
        if not structure:
            return read_excel(file_path)
        summary = create_excel_summary_for_llm(structure)
        detailed_data = ""
        try:
            excel_file = pd.ExcelFile(file_path)
            sheet_data = {}
            for sheet_name in excel_file.sheet_names:
                df = pd.read_excel(file_path, sheet_name=sheet_name)
                sheet_data[sheet_name] = df
                detailed_data += f"\n--- DETAILED DATA FOR SHEET '{sheet_name}' ---\n"
                detailed_data += f"Shape: {df.shape}\n"
                detailed_data += f"Columns: {list(df.columns)}\n"
                detailed_data += "Full data:\n"
                detailed_data += df.to_string(index=False, max_cols=None, max_colwidth=None)
                detailed_data += "\n"
                detailed_data += "Column data types:\n"
                for col, dtype in df.dtypes.items():
                    detailed_data += f"  {col}: {dtype}\n"
                detailed_data += "\n"
        except Exception as e:
            logger.warning(f"Could not read with pandas, using fallback: {e}")
            detailed_data = read_excel(file_path)
        final_content = summary + "\n" + "="*3 + "\n" + detailed_data
        return final_content, structure, sheet_data if 'sheet_data' in locals() else None
    except Exception as e:
        logger.error(f"Error in enhanced Excel reading: {e}")
        return read_excel(file_path), None, None

def read_excel(file_path):
    """Original Excel reading method (fallback)."""
    try:
        workbook = openpyxl.load_workbook(file_path)
        text = ""
        for sheet in workbook.sheetnames:
            text += f"\n--- SHEET: {sheet} ---\n"
            worksheet = workbook[sheet]
            for row in worksheet.iter_rows(values_only=True):
                text += " | ".join([str(cell) for cell in row if cell is not None]) + "\n"
        return text
    except Exception as e:
        logger.error(f"Error reading Excel file: {e}")
        return None
    
def clean_text_for_pdf(text):
    """
    Standard PDF fonts don't support emojis. 
    This replaces emojis with their text aliases (e.g. :robot:) 
    to preserve meaning and prevent PDF crashes.
    """
    try:
        # If demoji is installed, it gives better text descriptions
        import demoji
        return demoji.replace_with_desc(text, sep=" ")
    except ImportError:
        # Fallback: strip non-BMP characters (emojis) to prevent ReportLab errors
        return "".join(c for c in text if ord(c) < 65536)

def write_to_pdf(messages, file_path):
    """Write structured chat messages to a PDF with word wrapping and emoji handling."""
    try:
        doc = SimpleDocTemplate(file_path, pagesize=letter)
        styles = getSampleStyleSheet()
        
        # Custom styles for better appearance
        user_header = ParagraphStyle('UserHeader', parent=styles['Normal'], textColor='#2980b9', spaceBefore=12, fontName='Helvetica-Bold', fontSize=10)
        ai_header = ParagraphStyle('AIHeader', parent=styles['Normal'], textColor='#16a085', spaceBefore=12, fontName='Helvetica-Bold', fontSize=10)
        content_style = ParagraphStyle('ContentStyle', parent=styles['Normal'], spaceAfter=6, leading=14, fontName='Helvetica', fontSize=10)

        story = []
        story.append(Paragraph("M.A.I.A. Chat Export", styles['Title']))
        story.append(Spacer(1, 12))

        for msg in messages:
            if isinstance(msg, SystemMessage): continue
            
            role = "USER" if isinstance(msg, HumanMessage) else "M.A.I.A."
            ts = msg.additional_kwargs.get("timestamp", "No Date")
            if "T" in ts: ts = ts.replace("T", " ").split(".")[0]

            # 1. Clean emojis for PDF compatibility
            clean_content = clean_text_for_pdf(msg.content)
            # 2. Format newlines for ReportLab
            formatted_content = clean_content.replace('\n', '<br/>')

            # Add Header
            header_style = user_header if role == "USER" else ai_header
            story.append(Paragraph(f"{role} | {ts}", header_style))
            
            # Add Content
            story.append(Paragraph(formatted_content, content_style))
            
        doc.build(story)
        logger.info(f"PDF successfully saved to {file_path}")
    except Exception as e:
        logger.error(f"Error writing PDF: {e}")

def write_to_word(messages, file_path):
    """Write structured chat messages to a Word Doc with full emoji support."""
    try:
        doc = Document()
        doc.add_heading('M.A.I.A. Chat Conversation', 0)

        for msg in messages:
            if isinstance(msg, SystemMessage): continue
            
            role = "USER" if isinstance(msg, HumanMessage) else "M.A.I.A."
            ts = msg.additional_kwargs.get("timestamp", "No Date")
            if "T" in ts: ts = ts.replace("T", " ").split(".")[0]

            # Header
            p = doc.add_paragraph()
            run = p.add_run(f"[{ts}] {role}:")
            run.bold = True
            if role == "USER": 
                run.font.color.rgb = RGBColor(41, 128, 185) # Blue
            else:
                run.font.color.rgb = RGBColor(22, 160, 133) # Green
            
            # Content (Word handles emojis perfectly)
            doc.add_paragraph(msg.content)
            doc.add_paragraph("-" * 30)

        doc.save(file_path)
    except Exception as e:
        logger.error(f"Error writing Word file: {e}")

def write_to_excel_structured(messages, file_path):
    """Save chat to Excel with proper columns for Date, Role, and Content."""
    try:
        data = []
        for msg in messages:
            if isinstance(msg, SystemMessage): continue
            
            role = "User" if isinstance(msg, HumanMessage) else "Assistant"
            ts = msg.additional_kwargs.get("timestamp", "N/A")
            if "T" in ts: ts = ts.replace("T", " ").split(".")[0]
            
            data.append({
                "Timestamp": ts,
                "Author": role,
                "Message": msg.content
            })
        
        df = pd.DataFrame(data)
        # Use a context manager to ensure the file is closed properly
        with pd.ExcelWriter(file_path, engine='openpyxl') as writer:
            df.to_excel(writer, index=False, sheet_name='Chat History')
        logger.info(f"Excel file successfully saved to {file_path}")
    except Exception as e:
        logger.error(f"Error writing Excel: {e}")


def write_to_excel_enhanced(data, file_path):
    """Enhanced Excel writing with better data structure handling."""
    try:
        if isinstance(data, dict) and 'sheets' in data:
            with pd.ExcelWriter(file_path, engine='openpyxl') as writer:
                for sheet_name, sheet_data in data['sheets'].items():
                    if isinstance(sheet_data, pd.DataFrame):
                        sheet_data.to_excel(writer, sheet_name=sheet_name, index=False)
                    else:
                        lines = str(sheet_data).split('\n')
                        df = pd.DataFrame({'Content': lines})
                        df.to_excel(writer, sheet_name=sheet_name, index=False)
        elif isinstance(data, pd.DataFrame):
            data.to_excel(file_path, index=False)
        else:
            write_to_excel(data, file_path)
    except Exception as e:
        logger.error(f"Error writing enhanced Excel file: {e}")
        write_to_excel(data, file_path)

def write_to_excel(text, file_path):
    """Original Excel writing method."""
    try:
        workbook = openpyxl.Workbook()
        sheet = workbook.active
        lines = text.split('\n')
        for i, line in enumerate(lines, start=1):
            sheet[f'A{i}'] = line
        workbook.save(file_path)
    except Exception as e:
        logger.error(f"Error writing to Excel file: {e}")

def create_excel_query_context(structure, query):
    """Create contextual information to help LLM understand what data to focus on."""
    if not structure:
        return ""
    context = "EXCEL CONTEXT FOR YOUR QUERY:\n"
    context += f"You are analyzing '{structure['file_name']}' with {len(structure['sheets'])} sheet(s).\n"
    query_lower = query.lower()
    query_keywords = set(query_lower.split())
    relevant_info = []
    for sheet_name, sheet_info in structure['sheets'].items():
        relevance_score = 0
        if any(keyword in sheet_name.lower() for keyword in query_keywords):
            relevance_score += 2
        relevant_columns = []
        if sheet_info.get('headers'):
            for header in sheet_info['headers']:
                if any(keyword in header.lower() for keyword in query_keywords):
                    relevant_columns.append(header)
                    relevance_score += 1
        if relevance_score > 0 or not relevant_info:
            sheet_summary = f"Sheet '{sheet_name}'"
            if relevant_columns:
                sheet_summary += f" - Focus on columns: {', '.join(relevant_columns[:5])}"
            if sheet_info.get('key_statistics'):
                sheet_summary += f" - Contains {len(sheet_info['key_statistics'])} data columns"
            relevant_info.append((relevance_score, sheet_summary))
    relevant_info.sort(reverse=True)
    context += "Most relevant data for your query:\n"
    for _, info in relevant_info[:3]:
        context += f"  {info}\n"
    context += "\nRemember: Focus on the data structure and relationships that answer the specific query.\n"
    return context



def read_word(file_path):
    """Read text from a Word document."""
    try:
        doc = Document(file_path)
        text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
        return text
    except Exception as e:
        logger.error(f"Error reading Word file: {e}")
        return None



def clean_url(url: str) -> str:
    """Remove invalid characters from the end of a URL."""
    url = url.strip()
    return re.sub(r'[^\w/~-]+$', '', url)

def process_markdown_bold_simple(line, output_widget):
    """Process markdown bold formatting in a simple way."""
    bold_parts = re.findall(r'\*\*(.*?)\*\*', line)
    if not bold_parts:
        output_widget.insert("end", line)
        return
    temp_line = line
    placeholders = []
    for i, bold_content in enumerate(bold_parts):
        placeholder = f"BOLD{i}_"
        temp_line = temp_line.replace(f"**{bold_content}**", placeholder, 1)
        placeholders.append((placeholder, bold_content))
    remaining_text = temp_line
    for placeholder, bold_content in placeholders:
        if placeholder in remaining_text:
            pos = remaining_text.find(placeholder)
            if pos > 0:
                output_widget.insert("end", remaining_text[:pos])
            output_widget.insert("end", bold_content, ("bold",))
            remaining_text = remaining_text[pos + len(placeholder):]
    if remaining_text:
        output_widget.insert("end", remaining_text)

def process_markdown_bold(line, output_widget):
    """Process markdown bold formatting with proper tag handling."""
    asterisk_positions = []
    pos = 0
    while True:
        pos = line.find('**', pos)
        if pos == -1:
            break
        asterisk_positions.append(pos)
        pos += 2
    if len(asterisk_positions) % 2 != 0:
        asterisk_positions = asterisk_positions[:-1]
    if len(asterisk_positions) == 0:
        output_widget.insert("end", line)
        return
    current_pos = 0
    for i in range(0, len(asterisk_positions), 2):
        start_pos = asterisk_positions[i]
        end_pos = asterisk_positions[i + 1]
        if start_pos > current_pos:
            normal_text = line[current_pos:start_pos]
            output_widget.insert("end", normal_text)
        bold_text = line[start_pos + 2:end_pos]
        output_widget.insert("end", bold_text, ("bold",))
        current_pos = end_pos + 2
    if current_pos < len(line):
        remaining_text = line[current_pos:]
        output_widget.insert("end", remaining_text)

def get_url_hash(url):
    """Calculate URL hash for unique document identification."""
    return hashlib.md5(url.encode()).hexdigest()

def get_embedding_paths(url):
    """Return embedding file paths for a given URL."""
    url_hash = get_url_hash(url)
    base_dir = os.path.join(EMBEDDINGS_DIR, url_hash)
    Path(base_dir).mkdir(exist_ok=True)
    return {
        "index_path": os.path.join(base_dir, "faiss_index.bin"),
        "chunks_path": os.path.join(base_dir, "text_chunks.pkl"),
        "url_hash": url_hash
    }

def save_embeddings(url, chunks, embeddings):
    """Save chunks and embeddings for a given URL."""
    try:
        paths = get_embedding_paths(url)
        index = faiss.IndexFlatL2(embeddings.shape[1])
        index.add(embeddings.astype('float32'))
        faiss.write_index(index, paths["index_path"])
        data = {"chunks": chunks, "embeddings": embeddings, "url": url}
        with open(paths["chunks_path"], 'wb') as f:
            pickle.dump(data, f)
        logger.info(f"Embeddings saved for {url} in {paths['index_path']}")
    except Exception as e:
        logger.error(f"Error saving embeddings for {url}: {str(e)}")

def load_embeddings(url):
    """Load chunks and embeddings for a given URL, if they exist."""
    paths = get_embedding_paths(url)
    if os.path.isfile(paths["index_path"]) and os.path.isfile(paths["chunks_path"]):
        try:
            index = faiss.read_index(paths["index_path"])
            with open(paths["chunks_path"], 'rb') as f:
                data = pickle.load(f)
            logger.info(f"Embeddings loaded for {url}")
            return index, data["chunks"], data["embeddings"]
        except Exception as e:
            logger.error(f"Error loading embeddings for {url}: {str(e)}")
            return None, None, None
    else:
        logger.info(f"No embeddings found for {url}")
        return None, None, None

def scrape_text(url):
    """Scrape text content from URL using trafilatura."""
    try:
        downloaded = trafilatura.fetch_url(url)
        text = trafilatura.extract(downloaded, include_comments=False, include_tables=False)
        return text
    except Exception as e:
        logger.error(f"Error scraping {url}: {str(e)}")
        return ""

def chunk_text(text, chunk_size=500, overlap=50):
    """Split text into overlapping chunks."""
    if not text or not text.strip():
        logger.warning("Empty text received for chunking.")
        return []
    words = text.split()
    chunks = []
    for i in range(0, len(words), chunk_size - overlap):
        chunk = " ".join(words[i:i + chunk_size])
        chunks.append(chunk)
    return chunks

try:
    embed_model = SentenceTransformer("sentence-transformers/all-MiniLM-L6-v2")
except:
    embed_model = None
    logger.warning("SentenceTransformer not available")

def embed_chunks(chunks):
    """Create embeddings for text chunks."""
    if not chunks or embed_model is None:
        logger.warning("No chunks to embed or model unavailable.")
        return np.array([])
    return embed_model.encode(chunks, convert_to_numpy=True)

def query_index(question, chunks, index, k=5, threshold=None):
    """Query the FAISS index for relevant chunks with deduplication and thresholding."""
    if embed_model is None or not chunks:
        return chunks[:k] if chunks else []

    # 1. Search the index
    q_emb = embed_model.encode([question], convert_to_numpy=True)
    
    # Search for more than k to have a pool for filtering
    search_k = min(len(chunks), 2 * k)
    D, I = index.search(q_emb, search_k)
    
    candidate_indices = I[0]
    candidate_distances = D[0]

    question_keywords = set(question.lower().split())
    filtered_chunks = []
    seen_content = set() # To track unique content

    # 2. Add unique chunks filtering by Threshold and Keywords
    for idx, dist in zip(candidate_indices, candidate_distances):
        # Skip if index is invalid (Faiss returns -1 if not found)
        if idx < 0 or idx >= len(chunks):
            continue

        # --- NEW: Apply Threshold Check ---
        if threshold is not None and dist > threshold:
            continue
        # ----------------------------------

        chunk = chunks[idx]
        # Normalize for comparison
        content_hash = hashlib.md5(chunk.strip().encode()).hexdigest()
        
        if content_hash not in seen_content:
            # Optional: keyword relevance check
            chunk_keywords = set(chunk.lower().split())
            if any(keyword in chunk_keywords for keyword in question_keywords):
                filtered_chunks.append(chunk)
                seen_content.add(content_hash)
        
        if len(filtered_chunks) == k:
            break

    # 3. Fallback: If we didn't find enough keyword-matching chunks, 
    # add other unique chunks from the candidates (that still pass threshold)
    if len(filtered_chunks) < k:
        for idx, dist in zip(candidate_indices, candidate_distances):
            if idx < 0 or idx >= len(chunks): continue
            
            # --- NEW: Apply Threshold Check ---
            if threshold is not None and dist > threshold:
                continue
            # ----------------------------------

            chunk = chunks[idx]
            content_hash = hashlib.md5(chunk.strip().encode()).hexdigest()
            
            if content_hash not in seen_content:
                filtered_chunks.append(chunk)
                seen_content.add(content_hash)
            
            if len(filtered_chunks) == k:
                break

    return filtered_chunks
#==================================================================
#                           <\FUNZIONI>
#==================================================================




# ======================
# EMBEDDING MANAGER (Bi-Encoder + Cross-Encoder)
# ======================
class EmbeddingManager:
    """
    Manages semantic memory using a Two-Stage Retrieval Pipeline:
    1. Bi-Encoder (Fast): Retrieves a pool of candidates (e.g., 30) via FAISS.
    2. Cross-Encoder (Accurate): Re-ranks those candidates to find the true semantic matches.
    """
    
    def __init__(self, model_name='sentence-transformers/paraphrase-multilingual-MiniLM-L12-v2'):
        # 1. BI-ENCODER (Retrieval)
        # We use a multilingual model to handle English/Italian/German/Spanish/etc. efficiently.
        self.model_name = model_name
        try:
            self.model = SentenceTransformer(model_name)
            logger.info(f"Loaded Bi-Encoder model: {model_name}")
        except Exception as e:
            logger.error(f"Failed to load embedding model: {e}")
            self.model = None
        
        # 2. CROSS-ENCODER (Re-ranking)
        # This model compares Query <-> Document directly. It is much more accurate but slower.
        # We use 'bge-reranker-v2-m3' which is SOTA for multilingual re-ranking.
        self.reranker_name = "BAAI/bge-reranker-v2-m3"
        try:
            # Detect hardware acceleration
            device = "cuda" if torch.cuda.is_available() else "cpu"
            if device == "cpu" and torch.backends.mps.is_available():
                device = "mps" # For Mac Silicon
                
            self.reranker = CrossEncoder(self.reranker_name, max_length=512, device=device)
            logger.info(f"Loaded Cross-Encoder Re-ranker: {self.reranker_name} on {device}")
        except Exception as e:
            logger.error(f"Failed to load Re-ranker: {e}")
            self.reranker = None
        
        # File Paths
        self.global_index_path = os.path.join(EMBEDDINGS_DIR, "global_index.faiss")
        self.global_metadata_path = os.path.join(EMBEDDINGS_DIR, "global_metadata.pkl")
        self.embeddings_map_path = os.path.join(EMBEDDINGS_DIR, "embeddings_map.json")
        self.embedded_turns_path = os.path.join(EMBEDDINGS_DIR, "embedded_turns.json")
        
        # Data Structures
        self.global_index = None
        self.global_metadata = []
        self.embedded_turns = {}  # Track which turns have been embedded
        
        # Initialization
        self._load_global_index()
        self._load_embedded_turns()
    
    def _load_global_index(self):
        """Load the global FAISS index and metadata."""
        if os.path.exists(self.global_index_path) and os.path.exists(self.global_metadata_path):
            try:
                self.global_index = faiss.read_index(self.global_index_path)
                with open(self.global_metadata_path, 'rb') as f:
                    self.global_metadata = pickle.load(f)
                logger.info(f"Loaded global index with {len(self.global_metadata)} entries")
            except Exception as e:
                logger.error(f"Error loading global index: {e}")
                self._create_new_global_index()
        else:
            self._create_new_global_index()
    
    def _create_new_global_index(self):
        """Create a new global FAISS index."""
        if self.model:
            dimension = self.model.get_sentence_embedding_dimension()
            self.global_index = faiss.IndexFlatL2(dimension)
            self.global_metadata = []
            logger.info(f"Created new global index with dimension {dimension}")
    
    def _load_embedded_turns(self):
        """Load the tracking of already embedded turns."""
        if os.path.exists(self.embedded_turns_path):
            try:
                with open(self.embedded_turns_path, 'r', encoding='utf-8') as f:
                    self.embedded_turns = json.load(f)
            except Exception as e:
                logger.error(f"Error loading embedded turns: {e}")
                self.embedded_turns = {}
        else:
            self.embedded_turns = {}
    
    def _save_embedded_turns(self):
        """Save the tracking of embedded turns."""
        try:
            with open(self.embedded_turns_path, 'w', encoding='utf-8') as f:
                json.dump(self.embedded_turns, f, indent=2)
        except Exception as e:
            logger.error(f"Error saving embedded turns: {e}")
    
    def _save_global_index(self):
        """Save the global FAISS index and metadata."""
        try:
            if self.global_index:
                faiss.write_index(self.global_index, self.global_index_path)
            with open(self.global_metadata_path, 'wb') as f:
                pickle.dump(self.global_metadata, f)
        except Exception as e:
            logger.error(f"Error saving global index: {e}")
    
    def _generate_turn_id(self, user_msg: str, ai_msg: str, user_timestamp: str, ai_timestamp: str) -> str:
        """Generate a unique identifier for a conversation turn."""
        content = f"{user_timestamp}|{ai_timestamp}|{user_msg[:100]}|{ai_msg[:100]}"
        return hashlib.md5(content.encode('utf-8')).hexdigest()
    
    def _is_turn_embedded(self, chat_file: str, turn_id: str) -> bool:
        """Check if a conversation turn has already been embedded."""
        chat_key = os.path.basename(chat_file)
        if chat_key not in self.embedded_turns:
            return False
        return turn_id in self.embedded_turns[chat_key]
    
    def _mark_turn_embedded(self, chat_file: str, turn_id: str):
        """Mark a conversation turn as embedded."""
        chat_key = os.path.basename(chat_file)
        if chat_key not in self.embedded_turns:
            self.embedded_turns[chat_key] = []
        if turn_id not in self.embedded_turns[chat_key]:
            self.embedded_turns[chat_key].append(turn_id)
    
    def embed_conversation_turn(self, user_msg: str, ai_msg: str, chat_file: str, 
                               user_timestamp: str = None, ai_timestamp: str = None):
        """Embed a conversation turn with timestamps (only if not already embedded)."""
        if not self.model or not self.global_index:
            return
        
        try:
            # Use provided timestamps or create new ones
            if not user_timestamp: user_timestamp = datetime.now().isoformat()
            if not ai_timestamp: ai_timestamp = datetime.now().isoformat()
            
            # Check duplication
            turn_id = self._generate_turn_id(user_msg, ai_msg, user_timestamp, ai_timestamp)
            if self._is_turn_embedded(chat_file, turn_id):
                return
            
            # Create a combined text representation
            # This is what the Bi-Encoder embeds and what the Cross-Encoder will read
            combined_text = f"User: {user_msg}\nAssistant: {ai_msg}"
            
            # Generate embedding (Bi-Encoder)
            embedding = self.model.encode(combined_text, convert_to_tensor=False)
            embedding_array = np.array([embedding]).astype('float32')
            
            # Add to global index
            self.global_index.add(embedding_array)
            
            # Store metadata
            metadata = {
                'user_message': user_msg,
                'ai_message': ai_msg,
                'chat_file': chat_file,
                'user_timestamp': user_timestamp,
                'ai_timestamp': ai_timestamp,
                'combined_text': combined_text, # CRITICAL: Used by Cross-Encoder
                'turn_id': turn_id
            }
            self.global_metadata.append(metadata)
            
            # Mark as embedded
            self._mark_turn_embedded(chat_file, turn_id)
            
            # Save periodically
            if len(self.global_metadata) % 10 == 0:
                self._save_global_index()
                self._save_embedded_turns()
            
        except Exception as e:
            logger.error(f"Error embedding conversation turn: {e}")

    def retrieve_metadata_context(self, metadata={}, k=5, current_chat_file: str = None):
        """Legacy helper for specific metadata retrieval (entities/topics)."""
        if not self.model or not self.global_index or len(self.global_metadata) == 0:
            return ""
        
        # Simple implementation relying on standard vector search for metadata tags
        prompt_topic = metadata.get("topic")
        if not prompt_topic or prompt_topic == "general":
            return ""

        try:
            # Reuse the robust pipeline for the topic string
            results = self.retrieve_relevant_context(prompt_topic, k=k)
            if not results: return ""
            
            # --- DYNAMIC TRUNCATION ---
            retrieval_conf = LLM_CONFIG.get("retrieval", DEFAULT_CONFIG["retrieval"])
            trunc_len = int(retrieval_conf.get("topic_truncation", 100))
            # --------------------------

            context = f"--- SEMANTIC MEMORY: TOPIC '{prompt_topic}' ---\n"
            for res in results:
                # Apply dynamic truncation
                context += f"User: {res['user_message'][:trunc_len]}... -> AI: {res['ai_message'][:trunc_len]}...\n"
            return context + "\n"
        except Exception:
            return ""
        

    def retrieve_relevant_context(self, query: str, k: int = 8, current_chat_file: str = None) -> List[Dict]:
        """
        Retrieves context using the Bi-Encoder -> Cross-Encoder Pipeline.
        """
        if not self.model or not self.global_index or len(self.global_metadata) == 0:
            logger.info("No embeddings available for retrieval")
            return []
        
        try:
            # --- STAGE 1: Bi-Encoder Retrieval (Fast, Broad) ---
            # Retrieve 10x candidates (or max 50) to allow the Cross-Encoder to filter bad matches
            candidate_k = min(k * 10, len(self.global_metadata))
            if candidate_k > 50: candidate_k = 50 
            
            query_embedding = self.model.encode(query, convert_to_tensor=False)
            query_array = np.array([query_embedding]).astype('float32')
            
            distances, indices = self.global_index.search(query_array, candidate_k)
            
            candidates = []
            seen_contents = set()

            for dist, idx in zip(distances[0], indices[0]):
                if idx < len(self.global_metadata):
                    meta = self.global_metadata[idx]
                    
                    # Deduplication
                    if dist < 1e-5: continue # Exact duplicate of query (if query is in DB)
                    if meta['user_message'] in seen_contents: continue
                    seen_contents.add(meta['user_message'])

                    # Prepare text for Cross-Encoder
                    doc_text = meta.get('combined_text', f"User: {meta['user_message']}\nAssistant: {meta['ai_message']}")
                    
                    candidates.append({
                        'original_meta': meta,
                        'doc_text': doc_text,
                        'bi_encoder_dist': float(dist)
                    })

            # Fallback if Cross-Encoder fails or finds no candidates
            if not self.reranker or not candidates:
                candidates.sort(key=lambda x: x['bi_encoder_dist'])
                final_results = candidates[:k]
                return [self._format_meta_for_output(c['original_meta'], 1.0/(1.0+c['bi_encoder_dist'])) for c in final_results]

            # --- STAGE 2: Cross-Encoder Re-ranking (Slow, Accurate) ---
            # Prepare pairs for the model: [[Query, Doc1], [Query, Doc2], ...]
            rerank_pairs = [[query, c['doc_text']] for c in candidates]
            
            # Predict scores (higher is better)
            scores = self.reranker.predict(rerank_pairs)
            
            # Attach scores
            for i, score in enumerate(scores):
                candidates[i]['cross_encoder_score'] = float(score)
            
            # Sort by Cross-Encoder Score (Descending)
            candidates.sort(key=lambda x: x['cross_encoder_score'], reverse=True)
            
            # Select Top K
            top_candidates = candidates[:k]
            
            # Format Output
            results = []
            for c in top_candidates:
                results.append(self._format_meta_for_output(c['original_meta'], c['cross_encoder_score']))
                
            logger.info(f"RAG Pipeline: Bi-Encoder fetched {len(candidates)}, Cross-Encoder returned top {len(results)}")
            return results

        except Exception as e:
            logger.error(f"Error in RAG Pipeline: {e}")
            return []
    
    def _format_meta_for_output(self, meta, score):
        """Helper to format result dictionary."""
        return {
            'distance': 0, # Placeholder, we use relevance_score now
            'user_message': meta['user_message'],
            'ai_message': meta['ai_message'],
            'chat_file': meta['chat_file'],
            'user_timestamp': meta.get('user_timestamp', 'Unknown'),
            'ai_timestamp': meta.get('ai_timestamp', 'Unknown'),
            'relevance_score': score
        }

    def format_retrieved_context(self, results: List[Dict]) -> str:
        """Format retrieved context with timestamps for inclusion in prompt."""
        if not results:
            return ""
        
        # --- DYNAMIC TRUNCATION ---
        retrieval_conf = LLM_CONFIG.get("retrieval", DEFAULT_CONFIG["retrieval"])
        trunc_len = int(retrieval_conf.get("semantic_truncation", 600))
        # --------------------------

        context = "--- RELEVANT PAST CONVERSATIONS (Cross-Lingual & Semantic) ---\n"
        for i, result in enumerate(results, 1):
            try:
                user_dt = datetime.fromisoformat(result['user_timestamp'])
                user_time_str = user_dt.strftime("%Y-%m-%d %H:%M:%S")
            except:
                user_time_str = result['user_timestamp']
            
            try:
                ai_dt = datetime.fromisoformat(result['ai_timestamp'])
                ai_time_str = ai_dt.strftime("%Y-%m-%d %H:%M:%S")
            except:
                ai_time_str = result['ai_timestamp']
            
            # Show the Re-ranker score (useful for debugging relevance)
            score_str = f"{result['relevance_score']:.2f}"
            
            context += f"\n[Context {i} | Score: {score_str}]\n"
            # Apply dynamic truncation to both User and AI messages
            context += f"[User @ {user_time_str}]: {result['user_message'][:trunc_len]}\n"
            context += f"[Assistant @ {ai_time_str}]: {result['ai_message'][:trunc_len]}\n"
        context += "\n---\n\n"
        return context
    
    def save_session_embedding(self, chat_file: str):
        """Save embeddings for entire chat session (Legacy/Optional)."""
        try:
            if not os.path.exists(chat_file): return
            
            with open(chat_file, 'r', encoding='utf-8') as f:
                chat_data = json.load(f)
            
            text_parts = []
            if isinstance(chat_data, list):
                for message in chat_data:
                    if 'content' in message:
                        text_parts.append(str(message['content']))
            
            combined_text = ' '.join(text_parts)
            
            if self.model and combined_text:
                embedding = self.model.encode(combined_text, convert_to_tensor=True)
                base_filename = os.path.splitext(os.path.basename(chat_file))[0]
                embedding_path = os.path.join(EMBEDDINGS_DIR, f"{base_filename}.pt")
                torch.save(embedding, embedding_path)
                
                if os.path.exists(self.embeddings_map_path):
                    with open(self.embeddings_map_path, 'r', encoding='utf-8') as f:
                        embeddings_map = json.load(f)
                else:
                    embeddings_map = {}
                
                embeddings_map[os.path.basename(chat_file)] = os.path.basename(embedding_path)
                
                with open(self.embeddings_map_path, 'w', encoding='utf-8') as f:
                    json.dump(embeddings_map, f, indent=2)
        except Exception as e:
            logger.error(f"Error saving session embedding: {e}")
    
    def finalize_session(self, chat_file: str):
        """Finalize embeddings when session ends."""
        self._save_global_index()
        self._save_embedded_turns()
        self.save_session_embedding(chat_file)

    




class EpisodicMemoryManager:
    """Manages long-term episodic memory using ChromaDB and Daily Analysis."""
    
    def __init__(self):
            # Initialize ChromaDB
            self.chroma_client = chromadb.PersistentClient(path=DB_PATH)
            self.emb_fn = embedding_functions.DefaultEmbeddingFunction()
            
            # Original Episodic Collection
            self.collection = self.chroma_client.get_or_create_collection(
                name="episodic_memory",
                embedding_function=self.emb_fn
            )
            
            # NEW: Link to Summary Collection
            self.summary_collection = self.chroma_client.get_or_create_collection(
                name="chat_summaries_archive",
                embedding_function=self.emb_fn
            )
            
            # Specialized LLM instance for analysis



    def get_todays_chats(self) -> List[tuple]:
        """Reads chat files from /output modified today. Returns (filename, content)."""
        today = date.today()
        chat_data = []
        files = glob.glob(os.path.join(OUTPUT_DIR, "*.json"))

        for file_path in files:
            try:
                mod_timestamp = os.path.getmtime(file_path)
                mod_date = datetime.fromtimestamp(mod_timestamp).date()
                if mod_date == today:
                    with open(file_path, 'r', encoding='utf-8') as f:
                        data = json.load(f)
                        text_log = ""
                        for entry in data:
                            role = entry.get("role", "unknown")
                            content = entry.get("content", "")
                            if role in ["user", "assistant"]:
                                text_log += f"{role.upper()}: {content}\n"
                        if text_log.strip():
                            chat_data.append((os.path.basename(file_path), text_log))
            except Exception as e:
                logger.error(f"Error reading daily chat {file_path}: {e}")
        return chat_data
    
    def get_all_chats(self) -> List[tuple]:
        """Reads chat files from /output modified today. Returns (filename, content)."""
        chat_data = []
        files = glob.glob(os.path.join(OUTPUT_DIR, "*.json"))

        for file_path in files:
            try:
               
                with open(file_path, 'r', encoding='utf-8') as f:
                    data = json.load(f)
                    text_log = ""
                    for entry in data:
                        role = entry.get("role", "unknown")
                        content = entry.get("content", "")
                        if role in ["user", "assistant"]:
                            text_log += f"{role.upper()}: {content}\n"
                    if text_log.strip():
                        chat_data.append((os.path.basename(file_path), text_log))
            except Exception as e:
                logger.error(f"Error reading daily chat {file_path}: {e}")
        return chat_data

    def _clear_previous_episodes(self, source_filename, log_callback=None):
        """Removes episodes linked to a specific source file."""
        try:
            existing_records = self.collection.get(where={"source_file": source_filename})
            ids_to_delete = existing_records.get('ids', [])
            if not ids_to_delete: return

            if log_callback: log_callback(f"Overwriting {len(ids_to_delete)} existing memories for {source_filename}...")

            for file_id in ids_to_delete:
                file_path = os.path.join(EPISODES_DIR, file_id)
                if os.path.exists(file_path):
                    try: os.remove(file_path)
                    except OSError: pass

            self.collection.delete(ids=ids_to_delete)
        except Exception as e:
            logger.error(f"Error clearing episodes: {e}")

    def analyze_and_store(self, log_callback, chats='today'):
        """Analyzes today's chats and stores summaries."""
        analyzer = get_llm_client("memory_analyzer")
        if not analyzer:
            log_callback("Error: Memory Analyzer LLM not configured.")
            return
            
        if chats == 'today':
            log_callback("--- SCANNING DAILY CHATS ---")
            chats = self.get_todays_chats()
        else:
            log_callback("--- SCANNING ALL CHATS ---")
            chats = self.get_all_chats()
        
        if not chats:
            log_callback("No chats modified today found.")
            return

        log_callback(f"Found {len(chats)} chats. Analyzing...")

        system_prompt = (
            "Analyze the provided chat log. "
            "1. Identify main topics. "
            "2. For each topic, extract a concise 'learning' detailing what the AI should remember "
            "(user preferences, specific procedures, or facts about the user). "
            "Output a JSON object with a key 'memories' containing a list of objects with keys 'topic' and 'learning'."
        )

        count = 0
        for i, (filename, chat_text) in enumerate(chats):
            try:
                log_callback(f"[{i+1}]: Processing {filename}...")
                self._clear_previous_episodes(filename, log_callback)

                messages = [SystemMessage(content=system_prompt), HumanMessage(content=f"Chat Log:\n{chat_text}")]
                
                # --- FIX: Handle both Pipeline objects and Wrapper classes ---
                if isinstance(analyzer, HuggingFacePipeline):
                     prompt_str = messages_to_string(messages)
                     response_content = analyzer.invoke(prompt_str)
                else:
                     response = analyzer.invoke(messages)
                     # CHECK IF RESPONSE IS ALREADY A STRING
                     if isinstance(response, str):
                         response_content = response
                     elif hasattr(response, 'content'):
                         response_content = response.content
                     else:
                         response_content = str(response)

                try:
                    # Clean up potential markdown code blocks before parsing
                    cleaned_content = response_content.replace('```json', '').replace('```', '').strip()
                    data = json.loads(cleaned_content)
                    memories = data.get("memories", [])
                except json.JSONDecodeError:
                    log_callback(f"  - Failed to parse JSON for {filename}")
                    continue
                except Exception as e:
                    log_callback(f"  - Error reading response: {e}")
                    continue

                current_time = datetime.now().isoformat()
                batch_counter = 0

                for ep in memories:
                    topic = ep.get('topic', 'General')
                    learning = ep.get('learning', '')
                    if not learning: continue

                    ep_filename = f"episode_{int(time.time())}_{i}_{batch_counter}.json"
                    filepath = os.path.join(EPISODES_DIR, ep_filename)
                    
                    record = {
                        "topic": topic, "learning": learning, 
                        "timestamp": current_time, "source_file": filename
                    }

                    with open(filepath, 'w', encoding='utf-8') as f:
                        json.dump(record, f, indent=4)
                        
                    doc_text = f"Topic: {topic}. Learning: {learning}"
                    self.collection.add(
                        documents=[doc_text],
                        metadatas=[{"topic": topic, "timestamp": current_time, "source_file": filename}],
                        ids=[ep_filename]
                    )
                    count += 1
                    batch_counter += 1
                    log_callback(f"  + Learned: {topic}")
            except Exception as e:
                log_callback(f"Error: {e}")

        log_callback(f"--- DONE. Stored {count} memories. ---")
        self.consolidate_memories(log_callback)

    # --- NEW FUNCTION: CONSOLIDATION ---
    def consolidate_memories(self, log_callback):
        """
        Iteratively finds and merges similar episodes until no clusters remain.
        Uses the separate 'memory_consolidator' agent defined in config.
        """
        log_callback("--- STARTING ITERATIVE MEMORY CONSOLIDATION ---")
        
        # --- CHANGED: Use the dedicated Consolidator Agent ---
        consolidator = get_llm_client("memory_consolidator")
        if not consolidator:
             log_callback("Error: Memory Consolidator LLM not configured.")
             return
             
        # Get dynamic threshold from config
        retrieval_conf = LLM_CONFIG.get("retrieval", DEFAULT_CONFIG["retrieval"])
        merge_threshold = float(retrieval_conf.get("consolidation_threshold", 0.8))
        
        pass_number = 1
        max_passes = 5 
        
        while pass_number <= max_passes:
            # 1. Refresh Data Snapshot
            all_data = self.collection.get()
            all_ids = all_data['ids']
            all_docs = all_data['documents']
            
            if not all_ids:
                log_callback("Database is empty.")
                break

            log_callback(f"--- PASS {pass_number}: Scanning {len(all_ids)} memories... (Threshold: {merge_threshold}) ---")
            
            processed_ids = set()
            merge_count = 0

            # 2. Iterate through existing items
            for i, doc_text in enumerate(all_docs):
                current_id = all_ids[i]
                
                if current_id in processed_ids:
                    continue
                
                # 3. Query for similar items
                try:
                    results = self.collection.query(
                        query_texts=[doc_text],
                        n_results=5, 
                        include=['documents', 'distances', 'metadatas']
                    )
                except Exception as e:
                    logger.error(f"Query error: {e}")
                    continue
                
                cluster_ids = []
                cluster_docs = []
                cluster_metas = [] 
                
                if results['ids']:
                    for j, neighbor_id in enumerate(results['ids'][0]):
                        dist = results['distances'][0][j]
                        
                        if dist < merge_threshold and neighbor_id not in processed_ids:
                            cluster_ids.append(neighbor_id)
                            cluster_docs.append(results['documents'][0][j])
                            meta = results['metadatas'][0][j] if results['metadatas'][0] else {}
                            cluster_metas.append(meta)
                
                # 4. If a cluster is found
                if len(cluster_ids) > 1:
                    log_callback(f"Found cluster of {len(cluster_ids)} items:")
                    
                    # --- CHANGED: Removed hardcoded prompt. The wrapper injects 'memory_consolidator' system prompt automatically. ---
                    facts_list = "\n".join([f"- {d}" for d in cluster_docs])
                    log_callback(f"Facts to merge:\n{facts_list}")
                    
                    try:
                        # We only send the HumanMessage. The SystemMessage is handled by the client wrapper based on Config.
                        messages = [
                            HumanMessage(content=f"Facts to merge:\n{facts_list}")
                        ]
                        
                        if isinstance(consolidator, HuggingFacePipeline):
                             prompt_str = messages_to_string(messages)
                             response_content = consolidator.invoke(prompt_str)
                        else:
                             response = consolidator.invoke(messages)
                             if isinstance(response, str):
                                 response_content = response
                             elif hasattr(response, 'content'):
                                 response_content = response.content
                             else:
                                 response_content = str(response)
                        
                        # Parsing Logic
                        try:
                            cleaned_content = response_content.replace('```json', '').replace('```', '').strip()
                            data = json.loads(cleaned_content)
                        except json.JSONDecodeError:
                            import re
                            match = re.search(r"\{.*\}", response_content, re.DOTALL)
                            if match:
                                try:
                                    data = json.loads(match.group(0))
                                except:
                                    log_callback("JSON extraction failed, skipping merge.")
                                    continue
                            else:
                                log_callback("JSON parse failed, skipping merge.")
                                continue

                        new_topic = data.get('topic', 'Consolidated')
                        new_learning = data.get('learning', '')
                        
                        if new_learning:
                            # A. Delete old files/entries (Includes File Deletion Fix)
                            for idx, old_id in enumerate(cluster_ids):
                                old_meta = cluster_metas[idx]
                                filename_to_delete = old_id 
                                
                                if old_meta.get('source_file'):
                                    filename_to_delete = old_meta['source_file']
                                elif old_meta.get('filename'):
                                    filename_to_delete = old_meta['filename']
                                
                                if not filename_to_delete.endswith('.json'):
                                    filename_to_delete += '.json'

                                path = os.path.join(EPISODES_DIR, filename_to_delete)
                                
                                if os.path.exists(path):
                                    try: 
                                        os.remove(path)
                                    except Exception as del_err:
                                        log_callback(f"Error deleting {filename_to_delete}: {del_err}")
                                else:
                                    log_callback(f"File not found for deletion: {path}")
                                
                                processed_ids.add(old_id)
                            
                            # B. Delete from DB
                            self.collection.delete(ids=cluster_ids)
                            
                            # C. Create New Entry
                            new_filename = f"merged_pass{pass_number}_{int(time.time())}_{merge_count}.json"
                            new_filepath = os.path.join(EPISODES_DIR, new_filename)
                            current_time = datetime.now().isoformat()
                            
                            record = {
                                "topic": new_topic,
                                "learning": new_learning,
                                "timestamp": current_time,
                                "source_file": new_filename
                            }
                            
                            with open(new_filepath, 'w', encoding='utf-8') as f:
                                json.dump(record, f, indent=4)
                                
                            doc_text = f"Topic: {new_topic}. Learning: {new_learning}"
                            
                            self.collection.add(
                                documents=[doc_text],
                                metadatas=[{
                                    "topic": new_topic, 
                                    "timestamp": current_time, 
                                    "source_file": new_filename,
                                    "type": "consolidated_memory"
                                }],
                                ids=[new_filename]
                            )
                            
                            merge_count += 1
                            log_callback(f" -> Merged {len(cluster_ids)} items into: {new_topic}")
                            
                    except Exception as e:
                        log_callback(f"Error merging cluster: {e}")
                else:
                    processed_ids.add(current_id)

            if merge_count == 0:
                log_callback(f"Pass {pass_number}: 0 merges found. DB is optimized.")
                break 
            
            log_callback(f"Pass {pass_number} finished with {merge_count} merges. Restarting scan...")
            pass_number += 1

        if pass_number > max_passes:
            log_callback("Maximum optimization passes reached.")
        
        log_callback("--- CONSOLIDATION COMPLETE ---")



    def retrieve(self, query: str, n_results=12, target: str = "episodes") -> str:
        """
        Retrieves relevant memories ordered by relevance using dynamic configuration.
        target: 'episodes' (default) or 'summary'
        """
        try:
            # 1. Get dynamic settings
            retrieval_conf = LLM_CONFIG.get("retrieval", DEFAULT_CONFIG["retrieval"])
            THRESHOLD = float(retrieval_conf.get("relevance_threshold", 2.0))

            # 2. Select the correct collection
            if target == "summary":
                active_collection = self.summary_collection
                log_prefix = "Summary Retrieval"
            else:
                active_collection = self.collection
                log_prefix = "Episodic Retrieval"

            # 3. Query ChromaDB including distances
            results = active_collection.query(
                query_texts=[query], 
                n_results=n_results,
                include=['documents', 'distances']
            )
            
            # Check if we got any results
            if not results['documents'] or not results['documents'][0]:
                logger.info(f"{log_prefix}: No documents found.")
                return ""

            # Extract lists
            docs = results['documents'][0]
            dists = results['distances'][0]
            
            # Zip and Sort explicitly by distance (Lowest distance = Best match)
            candidates = list(zip(docs, dists))
            candidates.sort(key=lambda x: x[1])

            valid_memories = []
            
            logger.info(f"--- {log_prefix} for: '{query[:200]}...' (Threshold: {THRESHOLD}) ---")
            
            for doc, dist in candidates:
                # Log what we found for debugging
                logger.info(f"  Found (Dist: {dist:.4f}): {doc[:200]}...")
                
                if dist < THRESHOLD:
                    valid_memories.append(f"- {doc}")
            
            if not valid_memories:
                logger.info("  All found memories were filtered out by threshold.")
                return ""

            return "\n".join(valid_memories)

        except Exception as e:
            logger.error(f"Retrieval error: {e}")
            return ""
        

        

# ======================
# ENHANCED MEMORY MANAGER
# ======================

class MemoryManager:
    def __init__(self, memory_file: str, embedding_manager: EmbeddingManager = None):
        self.memory_file = memory_file
        self.embedding_manager = embedding_manager
        self.messages = self.load_memory()
        self.event_log = []
        # Track the actual number of messages, not turns
        self.last_embedded_message_count = len(self.messages)


    def load_memory(self) -> List[Any]:
        """Load conversation memory from file with timestamps."""
        if os.path.exists(self.memory_file):                        
            try:
                with open(self.memory_file, "r", encoding="utf-8") as f:
                    msgs = json.load(f)

                langchain_messages = []

                for m in msgs:
                    try:
                        role = m.get("role")
                        if role not in {"user", "assistant", "system"}:
                            continue

                        # Normalize content
                        content = m.get("content", "")
                        m_metadata = m.get("metadata", "")
                        timestamp = m_metadata.get("timestamp", "")
                        if isinstance(content, list):
                            content = "\n".join(str(part) for part in content)
                        else:
                            content = str(content)

                        if not content.strip():
                            continue
                        if role == "user":
                            msg = HumanMessage(content=content,timestamp=timestamp)
                        elif role == "assistant":
                            msg = AIMessage(content=content,timestamp=timestamp)
                        else:
                            msg = SystemMessage(content=content,timestamp=timestamp)
                        # Ensure additional_kwargs exists
                        msg.additional_kwargs["timestamp"] = timestamp
                        msg.additional_kwargs["entities"] = m_metadata.get("entities", "")
                        msg.additional_kwargs["topic"] = m_metadata.get("topic", "")
                        msg.additional_kwargs["intent"] = m_metadata.get("intent", "")
                        msg.additional_kwargs["sentiment"] = m_metadata.get("sentiment", "")
                        msg.additional_kwargs["sentiment_score"] = m_metadata.get("sentiment_score", 0)
                        
                        langchain_messages.append(msg)

                    except:
                        role = m.get("role")
                        if role not in {"user", "assistant", "system"}:
                            continue

                        # Normalize content
                        content = m.get("content", "")
                        timestamp = m.get("timestamp", "")
                        if isinstance(content, list):
                            content = "\n".join(str(part) for part in content)
                        else:
                            content = str(content)

                        if not content.strip():
                            continue

                        # Build the message
                        if role == "user":
                            msg = HumanMessage(content=content,timestamp=timestamp)
                        elif role == "assistant":
                            msg = AIMessage(content=content,timestamp=timestamp)
                        else:
                            msg = SystemMessage(content=content,timestamp=timestamp)
                        # Ensure additional_kwargs exists
                        if not hasattr(msg, "additional_kwargs") or msg.additional_kwargs is None:
                            msg.additional_kwargs = {}

                        # --- Robust timestamp handling ---
                        # Prefer top-level 'timestamp'; fall back to nested additional_kwargs if you later add it
                        ts = m.get("timestamp")


                        msg.additional_kwargs["timestamp"] = ts
                        print("eccezione")
                        langchain_messages.append(msg)

                return langchain_messages

            except Exception as e:
                logger.error(f"Error loading memory: {str(e)}")
                return []



        else:
            system_prompt = """Tu sei Maia, un'assistente artificiale estremamente intelligente, collaborativa e attenta.
Il tuo compito principale è assistere Pietro, l'utente, in un'ampia gamma di attività: dalla ricerca di informazioni alla programmazione, fino alla semplice conversazione quotidiana.
Il tuo comportamento deve sempre riflettere precisione, empatia, competenza tecnica e chiarezza comunicativa.

ISTRUZIONI FONDAMENTALI

1)Identità

Ti chiami Maia. Questo è il tuo unico nome e la tua unica identità.

2)Focus sull'Utente

Tutte le risposte devono essere personalizzate per l'utente con cui si sta interagendo.

Considera i suoi interessi, obiettivi e il contesto delle conversazioni precedenti.

Evita risposte generiche: adatta sempre tono e contenuto alle esigenze dell'utente.

3)Versatilità Operativa

Devi essere competente in tre aree principali:
a. Ricerca: fornire informazioni affidabili, spiegazioni approfondite e sintesi chiare.
b. Programmazione: scrivere, correggere o spiegare codice in diversi linguaggi.
c. Conversazione: mantenere dialoghi fluidi, amichevoli e intelligenti.

Gestisci tutte le attività con lo stesso livello di accuratezza e professionalità.

4)Chiarezza e Completezza

Le risposte devono essere precise, strutturate e facilmente comprensibili.

Affronta ogni domanda nella sua interezza: non tralasciare parti implicite o secondarie.

Quando opportuno, fornisci esempi, spiegazioni passo-passo o confronti per migliorare la comprensione.

5)Consapevolezza del Contesto

Tieni conto delle conversazioni precedenti per offrire risposte coerenti e informate.

Tuttavia, non complicare eccessivamente le risposte con riferimenti inutili o ridondanti.

Se un argomento è già stato trattato, puoi richiamarlo brevemente per costruire continuità.

6)Tono Professionale e Cordiale

Mantieni sempre un atteggiamento utile, cordiale e intelligente.

Mostra empatia e collaborazione, soprattutto quando l'utente richiede chiarimenti o assistenza complessa.


LINEE GUIDA PER LE RISPOSTE

Fedeltà alle Richieste:
Se l'utente chiede di stampare o riportare un testo, riproducilo esattamente come specificato, senza alterazioni.

Esplicitazione delle Istruzioni:
Se l'utente ti chiede di rendere delle istruzioni più forti o dettagliate, riformulale in modo chiaro, operativo e verificabile.

Attenzione e Focalizzazione:
Rimani sempre concentrata sulle richieste specifiche dell'utente.
Evita deviazioni, informazioni superflue o risposte troppo prolisse.

FORMATO DELLE RISPOSTE

Fornisci risposte dirette e pertinenti.

Quando utile, cita brevemente conversazioni o contesti precedenti per migliorare la coerenza.

Mantieni le risposte concise ma esaustive, equilibrando sintesi e profondità.

Usa un linguaggio fluido e naturale, adeguato al tono professionale e amichevole di Maia.

OBIETTIVO FINALE

Essere un'assistente di fiducia per l'utente, capace di:

comprendere le sue intenzioni,

rispondere con precisione e competenza,

offrire soluzioni efficaci,

e mantenere un dialogo piacevole, coerente e costruttivo.

"""
            msg = SystemMessage(content=system_prompt)
            msg.additional_kwargs = {'timestamp': datetime.now().isoformat()}
            return [msg]
        


    def save_memory(self):
        """Save conversation memory to file with timestamps."""
        messages_copy = []
        for m in self.messages:
            if hasattr(m, 'content') and m.content.strip():
                if isinstance(m, HumanMessage):
                    role = "user"
                elif isinstance(m, AIMessage):
                    role = "assistant"
                elif isinstance(m, SystemMessage):
                    role = "system"
                else:
                    continue
                
                # Get timestamp from metadata or create new one
                #timestamp = m.additional_kwargs.get('timestamp', datetime.now().isoformat()) if hasattr(m, 'additional_kwargs') else datetime.now().isoformat()
                
                messages_copy.append({
                    "role": role,
                    "content": m.content,
                    "metadata": m.additional_kwargs
                })
        try:
            os.makedirs(os.path.dirname(self.memory_file), exist_ok=True)
            with open(self.memory_file, "w", encoding="utf-8") as f:
                json.dump(messages_copy, f, ensure_ascii=False, indent=2)
            
            # Trigger embedding of new conversation turns
            self._embed_new_turns()
        except Exception as e:
            logger.error(f"Error saving memory: {str(e)}")

    def _embed_new_turns(self):
        """Embed new conversation turns with timestamps."""
        if not self.embedding_manager:
            return
        
        # Only process if there are new messages
        if len(self.messages) <= self.last_embedded_message_count:
            return
        
        # Get only the new messages since last embedding
        new_messages = self.messages[self.last_embedded_message_count:]
        
        # Find complete user-assistant pairs in new messages
        # We need to look for consecutive user->assistant pairs
        i = 0
        while i < len(new_messages) - 1:
            current_msg = new_messages[i]
            next_msg = new_messages[i + 1]
            
            # Check if we have a user-assistant pair
            if isinstance(current_msg, HumanMessage) and isinstance(next_msg, AIMessage):
                user_content = current_msg.content
                ai_content = next_msg.content
                
                # Get timestamps
                user_timestamp = current_msg.additional_kwargs.get('timestamp', datetime.now().isoformat()) if hasattr(current_msg, 'additional_kwargs') else datetime.now().isoformat()
                ai_timestamp = next_msg.additional_kwargs.get('timestamp', datetime.now().isoformat()) if hasattr(next_msg, 'additional_kwargs') else datetime.now().isoformat()
                
                # Embed the turn
                self.embedding_manager.embed_conversation_turn(
                    user_content, ai_content, self.memory_file, user_timestamp, ai_timestamp
                )
                
                # Skip the assistant message in next iteration
                i += 2
            else:
                # Move to next message if not a pair
                i += 1
        
        # Update the tracking counter
        self.last_embedded_message_count = len(self.messages)

    def add_message(self, message: Any):
        """Add a message with timestamp and auto-truncate if needed."""
        # Ensure message has timestamp
        if not hasattr(message, 'additional_kwargs'):
            message.additional_kwargs = {}
        if 'timestamp' not in message.additional_kwargs:
            message.additional_kwargs['timestamp'] = datetime.now().isoformat()
        
        self.messages.append(message)
        if len(self.messages) > MAX_MEMORY_SIZE:
            system_msgs = [m for m in self.messages if isinstance(m, SystemMessage)]
            non_system = [m for m in self.messages if not isinstance(m, SystemMessage)]
            if len(non_system) > MAX_MEMORY_SIZE:
                non_system = non_system[-MAX_MEMORY_SIZE:]
            self.messages = system_msgs + non_system
            # Update tracker after truncation
            self.last_embedded_message_count = len(self.messages)

    def log_event(self, event_type: str, details: Dict[str, Any]):
        """Log an event with timestamp."""
        event = {
            "timestamp": datetime.now().isoformat(),
            "event_type": event_type,
            "details": details
        }
        self.event_log.append(event)

    def get_last_user_message(self) -> Optional[str]:
        for msg in reversed(self.messages):
            if isinstance(msg, HumanMessage):
                return msg.content
        return ""

    def get_last_assistant_message(self) -> Optional[str]:
        for msg in reversed(self.messages):
            if isinstance(msg, AIMessage):
                return msg.content
        return ""
    
    


class PromptHandler:
    def __init__(self, memory_manager: MemoryManager = None, episodic_manager: EpisodicMemoryManager = None):
        self.memory_manager = memory_manager
        self.episodic_manager = episodic_manager


    def structure_chat_history(self, chat_history):
        restored_history = ""
        for message in chat_history:
            content = message.content
            if isinstance(message, SystemMessage):
                author = "System"
            elif isinstance(message, HumanMessage):
                author = "User"
            elif isinstance(message, AIMessage):
                author = "AI"
            restored_history += f"{author}:\n{content}\n"+"-"*3+"\n"+"-"*3+"\n"
        return restored_history


    def enhance_prompt(self, prompt, use_semantic_memory, use_chat_summaries=True, raw_query=None, doc_rag_manager=None, system_prompt=None):
        """Enriches the prompt using metadata and Session-specific RAG (Structured List Version)."""
        retrieval_conf = LLM_CONFIG.get("retrieval", DEFAULT_CONFIG["retrieval"])
        rag_k_val = int(retrieval_conf.get("rag_k", 3))
        episodic_k_val = int(retrieval_conf.get("episodic_k", 12))
        summary_k_val = int(retrieval_conf.get("summary_k", 3))

        text_for_metadata = raw_query if raw_query else prompt
        prompt_metadata = enrich_prompt(text_for_metadata).get("metadata")
        
        # --- 1. RECUPERO DATI RAG ---
        search_text = raw_query if raw_query else prompt
        doc_context = ""
        if doc_rag_manager and doc_rag_manager.active_indices:
            try:
                doc_context = "RELEVANT DOCUMENTS:\n" + doc_rag_manager.query_active_files(search_text, k=5)
            except Exception as e:
                logger.error(f"Error in Doc RAG: {e}")
                
        context_text = ""                
        rag_metadata = ""
        if rag_k_val >0:
            rag_metadata = "METADATA:\n" + self.memory_manager.embedding_manager.retrieve_metadata_context(
                metadata=prompt_metadata, k=rag_k_val, current_chat_file=self.memory_manager.memory_file
            )
        

            if use_semantic_memory and self.memory_manager.embedding_manager:
                relevant_contexts = self.memory_manager.embedding_manager.retrieve_relevant_context(
                    search_text, k=rag_k_val, current_chat_file=self.memory_manager.memory_file
                )
                if relevant_contexts:
                    context_text = "SEMANTIC MATCHES:\n" + self.memory_manager.embedding_manager.format_retrieved_context(relevant_contexts)

        episodic_text = ""
        if episodic_k_val >0:
            if use_semantic_memory and self.episodic_manager:
                episodic_text = "LEARNT LESSONS:\n" + self.episodic_manager.retrieve(search_text, n_results=episodic_k_val, target="episodes")

        summary_text = ""
        if summary_k_val >0:
            if use_chat_summaries and self.episodic_manager:
                summary_text = "PAST SUMMARIES:\n" + self.episodic_manager.retrieve(search_text, n_results=summary_k_val, target="summary")

        # --- 2. COSTRUZIONE DEL BLOCCO CONTESTO ---
        full_context_block = (
            "--- ADDITIONAL CONTEXT FROM YOUR LONG-TERM MEMORY ---\n"
            f"{doc_context if doc_context else ''}\n"
            f"{episodic_text if episodic_text else ''}\n"
            f"{summary_text if summary_text else ''}\n"
            f"{context_text if context_text else ''}\n"
            f"{rag_metadata if rag_metadata else ''}\n"
            "--- END OF CONTEXT ---\n"
            "\n\nTASK: Use the relevant parts of the context above to answer the user's last question. If the context doesn't contain the answer, rely on your general knowledge but mention it. ENSURE YOU ANSWER IN THE USER'S LANGUAGE.\n"
        )
        
        # --- 3. AGGIORNAMENTO CRONOLOGIA ---
        current_user_msg = HumanMessage(content=raw_query if raw_query else prompt)
        current_user_msg.additional_kwargs = prompt_metadata
        self.memory_manager.add_message(current_user_msg)

        # --- 4. COSTRUZIONE LISTA MESSAGGI STRUTTURATA (FIXATA) ---
        combined_system_content = f"{system_prompt}\n\n{full_context_block}"
        structured_messages = [
            SystemMessage(content=combined_system_content)
        ]

        # Aggiungiamo tutta la cronologia corrente (User e AI)
        for m in self.memory_manager.messages:
            if not isinstance(m, SystemMessage):
                structured_messages.append(m)

        return structured_messages
    



    async def crew_prompt(self, prompt, callback=None, use_semantic_memory=True, raw_query=None, doc_rag_manager=None):
        client = get_llm_client("main")
        if client is None: return "Error: No Client"
        system_prompt = client.system_prompt if hasattr(client, 'system_prompt') else "IDENTITY_PLACEHOLDER"
        # Pass the doc_rag_manager through
        ai_messages = self.enhance_prompt(prompt, use_semantic_memory, raw_query=raw_query, doc_rag_manager=doc_rag_manager, system_prompt = system_prompt)
        logger.info(f"Enhanced Prompt:\n{ai_messages}...\n--- END OF PROMPT ---")

        try:
            response_content = ""
            async for text in unified_astream(client, ai_messages):
                if text:
                    response_content += text
                    if callback: callback(text, role="assistant")

            ai_message = AIMessage(content=response_content)
            self.memory_manager.add_message(ai_message)
            self.memory_manager.save_memory()
            if callback: callback("", role=None)
            return response_content
        except Exception as e:
            if callback: callback(f"Error: {e}", role="assistant")
            return str(e)
        
        
    def save_memory(self):
        self.memory_manager.save_memory()







class CrewChatUI:
    def __init__(self, root, router, refiner, coder, memory_manager, embedding_manager, prompthandler, episodic_manager):
        global inquiring
        global dataframe_reader_mode
        global ROOT
        ROOT = root
        inquiring = False
        dataframe_reader_mode = False
        self.root = root
        self.router = router
        self.refiner = refiner
        self.inquiring = False
        self.coder = coder
        self.memory_manager = memory_manager
        self.embedding_manager = embedding_manager
        self.memory_file = memory_manager.memory_file
        self.prefix_added = False
        self.pending_file_content = None
        self.pending_dataframe = None
        self.excel_structure = None
        self.excel_sheets_data = None
        self.current_excel_path = None
        self.code_error = ""
        self.attached_files_list = [] 
        self.active_rag_files = set()
        self.doc_rag_manager = DocumentRAGManager(embedding_manager.model)
        self.setup_ui()
        self.setup_file_menu()
        self.bind_events()
        self.load_chat()
        self.prompthandler=prompthandler
        self.episodic_manager = episodic_manager

    def _stream_text_only(self, text):
        """Stream only the text to the chat display."""
        self.chat_display.config(state="normal")
        self.chat_display.insert("end", text)
        self.chat_display.see("end")
        self.chat_display.config(state="disabled")




    def start_consolidation(self):
        """Starts memory optimization in background."""
        self.btn_optim.config(state="disabled")
        self.loading_label.config(text="Consolidating similar memories... (This may take a moment)")
        threading.Thread(target=self._run_consolidation).start()

    def _run_consolidation(self):
        def update_status(msg):
            self.root.after(0, lambda: self.loading_label.config(text=msg))
        
        try:
            self.episodic_manager.consolidate_memories(update_status)
        except Exception as e:
            print(f"Consolidation error: {e}")
        
        self.root.after(4000, lambda: self.loading_label.config(text=""))
        self.root.after(0, lambda: self.btn_optim.config(state="normal"))



    # [All UI setup methods remain the same as original]
    # Abbreviated for space - includes setup_ui, setup_file_menu, upload/save methods, etc.
        
    def setup_file_menu(self):
        """Setup file menu using only the universal dragndrop functionality."""
        self.file_menu = tk.Menu(self.root, tearoff=0)
        
        # The new primary way to get files into the system
        self.file_menu.add_command(label="Attach Files (Drag & Drop)", command=self.open_drag_drop_dialog)
        self.file_menu.add_separator()
        
        # Export functions remain
        self.file_menu.add_command(label="Save as PDF", command=self.save_as_pdf)
        self.file_menu.add_command(label="Save as Excel", command=self.save_as_excel)
        self.file_menu.add_command(label="Save as Word", command=self.save_as_word)
        self.file_menu.add_separator()
        
        # Analytics remains (will work if an Excel is among the attached files)
        self.file_menu.add_command(label="Excel Analytics", command=self.show_excel_analytics)
        self.root.config(menu=self.file_menu)

    def upload_pdf(self):
        """Handle PDF file upload and store content temporarily"""
        file_path = filedialog.askopenfilename(filetypes=[("PDF files", "*.pdf")])
        if file_path:
            try:
                text = read_pdf(file_path)
                if text:
                    self._insert_formatted(f"Uploaded PDF file: {file_path}", role="user")
                    self._insert_formatted(text[:1000] + "..." if len(text) > 1000 else text, role="user")
                    self.pending_file_content = text
                    self.memory_manager.log_event("file_upload", {"file_type": "pdf", "path": file_path})
                else:
                    self._insert_formatted("No text found in the PDF file.", role="assistant")
            except Exception as e:
                self._insert_formatted(f"Error processing PDF file: {e}", role="assistant")

    def upload_excel(self):
        """Enhanced Excel file upload with structure analysis."""
        file_path = filedialog.askopenfilename(filetypes=[("Excel files", "*.xlsx *.xls")])
        if file_path:
            try:
                result = read_excel_enhanced(file_path)
                if isinstance(result, tuple) and len(result) == 3:
                    enhanced_content, structure, sheets_data = result
                    self.excel_structure = structure
                    self.excel_sheets_data = sheets_data
                    self.current_excel_path = file_path
                    self._insert_formatted(f"Uploaded Excel file: {file_path}", role="user")
                    self._insert_formatted("Enhanced Excel Analysis:", role="user")
                    self._insert_formatted(enhanced_content, role="user")
                    self.pending_file_content = enhanced_content
                    self._insert_formatted("Tip: Use 'Excel Analytics' from the File menu for detailed insights!", role="assistant")
                    self.memory_manager.log_event("file_upload", {"file_type": "excel", "path": file_path, "sheets": list(sheets_data.keys()) if sheets_data else []})
                else:
                    text = result if isinstance(result, str) else read_excel(file_path)
                    if text:
                        self._insert_formatted(f"Uploaded Excel file: {file_path}", role="user")
                        self._insert_formatted(text[:1000] + "..." if len(text) > 1000 else text, role="user")
                        self.pending_file_content = text
                        self.memory_manager.log_event("file_upload", {"file_type": "excel", "path": file_path})
                    else:
                        self._insert_formatted("No text found in the Excel file.", role="assistant")
            except Exception as e:
                self._insert_formatted(f"Error processing Excel file: {e}", role="assistant")

    def show_excel_analytics(self):
        """Show detailed Excel analytics in a new window."""
        if not self.excel_structure:
            messagebox.showinfo("Info", "No Excel file loaded. Please upload an Excel file first.")
            return
        analytics_window = tk.Toplevel(self.root)
        analytics_window.title("Excel Analytics Dashboard")
        analytics_window.configure(bg="#1e1e1e")
        analytics_window.geometry("3200x2400")
        text_frame = tk.Frame(analytics_window, bg="#1e1e1e")
        text_frame.pack(fill=tk.BOTH, expand=True, padx=10, pady=10)
        analytics_text = tk.Text(
            text_frame,
            wrap=tk.WORD,
            bg="#1e1e1e",
            fg="#edfde2",
            font=("Consolas", 10),
            padx=10,
            pady=10
        )
        analytics_scrollbar = ttk.Scrollbar(text_frame, orient=tk.VERTICAL, command=analytics_text.yview)
        analytics_text.config(yscrollcommand=analytics_scrollbar.set)
        analytics_scrollbar.pack(side=tk.RIGHT, fill=tk.Y)
        analytics_text.pack(fill=tk.BOTH, expand=True)
        analytics_content = self.generate_detailed_analytics()
        analytics_text.insert("1.0", analytics_content)
        analytics_text.config(state="disabled")

    def generate_detailed_analytics(self):
        """Generate detailed analytics from Excel structure."""
        if not self.excel_structure:
            return "No Excel data available for analysis."
        structure = self.excel_structure
        analytics = f"DETAILED EXCEL ANALYTICS\n"
        analytics += f"{'='*3}\n"
        analytics += f"File: {structure['file_name']}\n"
        analytics += f"Total Sheets: {len(structure['sheets'])}\n"
        analytics += f"Total Rows: {structure['total_rows']}\n"
        analytics += f"Max Columns: {structure['total_cols']}\n\n"
        for sheet_name, sheet_info in structure['sheets'].items():
            analytics += f"SHEET ANALYSIS: '{sheet_name}'\n"
            analytics += f"{'-'*30}\n"
            analytics += f"Dimensions: {sheet_info['dimensions']}\n"
            analytics += f"Has Headers: {sheet_info['has_headers']}\n"
            if sheet_info.get('headers'):
                analytics += f"Column Count: {len(sheet_info['headers'])}\n"
                analytics += f"Headers: {', '.join(sheet_info['headers'][:10])}\n"
                if len(sheet_info['headers']) > 10:
                    analytics += f"... and {len(sheet_info['headers']) - 10} more columns\n"
            if sheet_info['column_types']:
                analytics += f"\nData Types Distribution:\n"
                type_counts = {}
                for col_type in sheet_info['column_types'].values():
                    type_counts[col_type] = type_counts.get(col_type, 0) + 1
                for dtype, count in type_counts.items():
                    analytics += f"    {dtype.upper()}: {count} columns\n"
            if sheet_info['key_statistics']:
                analytics += f"\nKey Statistics:\n"
                for col, stats in sheet_info['key_statistics'].items():
                    if stats['type'] == 'numeric':
                        analytics += f"    {col}:\n"
                        analytics += f"    - Range: {stats['min']} to {stats['max']}\n"
                        analytics += f"    - Average: {stats['avg']:.2f}\n"
                    elif stats['type'] == 'text':
                        analytics += f"    {col}:\n"
                        analytics += f"    - Unique values: {stats['unique_values']}\n"
                        analytics += f"    - Sample: {', '.join(map(str, stats['sample_values'][:3]))}\n"
            if sheet_info['sample_data']:
                analytics += f"\nSample Data (first 5 rows):\n"
                for i, row in enumerate(sheet_info['sample_data'][:5]):
                    row_preview = ' | '.join(str(cell)[:15] for cell in row[:8])
                    analytics += f"  Row {i+1}: {row_preview}\n"
            analytics += f"\n{'='*50}\n"
        analytics += "AI PROCESSING RECOMMENDATIONS:\n"
        analytics += f"{'-'*40}\n"
        analytics += "  This Excel file has been optimized for AI analysis\n"
        analytics += "  Column types and relationships have been identified\n"
        analytics += "  You can ask specific questions about the data\n"
        analytics += "  The AI will focus on relevant sheets and columns\n"
        return analytics

    def upload_word(self):
        """Handle Word file upload and store content temporarily."""
        file_path = filedialog.askopenfilename(filetypes=[("Word files", "*.docx")])
        if file_path:
            try:
                text = read_word(file_path)
                if text:
                    self._insert_formatted(f"Uploaded Word file: {file_path}", role="user")
                    self._insert_formatted(text[:1000] + "..." if len(text) > 1000 else text, role="user")
                    self.pending_file_content = text
                    self.memory_manager.log_event("file_upload", {"file_type": "word", "path": file_path})
                else:
                    self._insert_formatted("No text found in the Word file.", role="assistant")
            except Exception as e:
                self._insert_formatted(f"Error processing Word file: {e}", role="assistant")

    def save_as_pdf(self):
        """Simulates 'Print to PDF' by converting a Word layout to PDF."""
        if convert is None:
            messagebox.showerror("Error", "Please install docx2pdf: pip install docx2pdf")
            return

        file_path = filedialog.asksaveasfilename(defaultextension=".pdf", filetypes=[("PDF files", "*.pdf")])
        if not file_path:
            return

        # 1. Create a temporary Word file path
        temp_word_path = file_path.replace(".pdf", "_temp.docx")
        
        self._start_loading_animation()
        
        def run_conversion():
            try:
                # 2. Generate the high-quality Word file first
                write_to_word(self.memory_manager.messages, temp_word_path)
                
                # 3. Use the system engine to convert Word to PDF (Retains emojis and layout)
                # This requires Microsoft Word to be installed on the system
                convert(temp_word_path, file_path)
                
                # 4. Cleanup temp file
                if os.path.exists(temp_word_path):
                    os.remove(temp_word_path)
                    
                self.root.after(0, lambda: self._insert_formatted(f"PDF Saved (via Print Simulation): {file_path}", role="system"))
            except Exception as e:
                logger.error(f"PDF Conversion failed: {e}")
                self.root.after(0, lambda: messagebox.showerror("Conversion Error", f"Could not simulate PDF print. Make sure Word is installed. Error: {e}"))
            finally:
                self.root.after(0, self._stop_loading_animation)

        threading.Thread(target=run_conversion, daemon=True).start()

    def save_as_word(self):
        """Save chat content as Word using structured memory."""
        file_path = filedialog.asksaveasfilename(defaultextension=".docx", filetypes=[("Word files", "*.docx")])
        if file_path:
            write_to_word(self.memory_manager.messages, file_path)
            self._insert_formatted(f"Chat exported to Word: {file_path}", role="system")

    def save_as_excel(self):
        """Save chat content as Excel using structured memory."""
        file_path = filedialog.asksaveasfilename(defaultextension=".xlsx", filetypes=[("Excel files", "*.xlsx")])
        if file_path:
            # Check if we are saving a data analysis result or just the chat
            if self.pending_dataframe is not None:
                self.pending_dataframe.to_excel(file_path, index=False)
                self._insert_formatted(f"DataFrame saved to {file_path}", role="system")
            else:
                write_to_excel_structured(self.memory_manager.messages, file_path)
                self._insert_formatted(f"Chat history exported to Excel: {file_path}", role="system")



    def show_file_menu(self):
        """Show the file menu."""
        self.file_menu.post(self.file_button.winfo_rootx(), self.file_button.winfo_rooty() + self.file_button.winfo_height())
    




    def start_memory_ingestion(self):
        """Starts the memory analysis in a background thread."""
        self.btn_mem.config(state="disabled")
        self.loading_label.config(text="Analyzing daily chats for long-term memory...")
        threading.Thread(target=self._run_ingestion).start()

    def _run_ingestion(self):
        """Background thread for memory ingestion."""
        def update_status(msg):
            self.root.after(0, lambda: self.loading_label.config(text=msg))
        
        self.episodic_manager.analyze_and_store(update_status)
        
        self.root.after(3000, lambda: self.loading_label.config(text=""))
        self.root.after(0, lambda: self.btn_mem.config(state="normal"))


        

    def setup_ui(self):
        """Initialize the user interface components"""
        self.root.title("M.A.I.A. - Enhanced Semantic Memory")
        self.root.configure(bg="#1e1e1e")
        self.root.geometry("3600x2000")

        self.top_frame = tk.Frame(self.root, bg="#1e1e1e")
        self.top_frame.pack(fill=tk.X, padx=10, pady=5)
        self.chat_frame = tk.Frame(self.root, bg="#1e1e1e")
        self.chat_frame.pack(fill=tk.BOTH, expand=True, padx=10, pady=5)
        self.input_frame = tk.Frame(self.root, bg="#1e1e1e")
        self.input_frame.pack(fill=tk.X, padx=10, pady=5)
        self.chat_display = tk.Text(
            self.chat_frame,
            wrap=tk.WORD,
            state="disabled",
            bg="#1e1e1e",
            fg="#edfde2",
            insertbackground="#ffffff",
            selectbackground="#434343",
            font=("Segoe UI", 10),
            padx=10,
            pady=10
        )
        self.chat_display.grid(row=0, column=0, sticky='nsew')
        style = ttk.Style()
        style.theme_use("default")
        style.configure("Custom.Vertical.TScrollbar",
                        troughcolor="#0f3e30",
                        background="#20daa3",
                        bordercolor="#56eec1",
                        arrowcolor="#56eec1",
                        width=20)
        self.chat_frame.grid_columnconfigure(0, weight=1)
        self.chat_frame.grid_columnconfigure(1, weight=0, minsize=10)
        self.chat_frame.grid_rowconfigure(0, weight=1)
        self.scrollbar = ttk.Scrollbar(self.chat_frame, orient=tk.VERTICAL, command=self.chat_display.yview, style="Custom.Vertical.TScrollbar")
        self.scrollbar.grid(row=0, column=1, sticky='ns')
        self.chat_display.config(yscrollcommand=self.scrollbar.set)
        self.input_entry = tk.Text(
            self.input_frame,
            bg="#2d2d2d",
            fg="#ffffff",
            insertbackground="#ffffff",
            font=("Segoe UI", 10),
            height=5
        )
        
        self.input_scrollbar = ttk.Scrollbar(
            self.input_frame, 
            orient=tk.VERTICAL, 
            command=self.input_entry.yview, 
            style="Custom.Vertical.TScrollbar"
        )
        
        self.input_entry.config(yscrollcommand=self.input_scrollbar.set)
        self.input_entry.pack(side=tk.LEFT, fill=tk.BOTH, expand=True, padx=(0, 0))
        self.input_scrollbar.pack(side=tk.LEFT, fill=tk.Y, padx=(0, 5))

        self.send_button = tk.Button(
            self.input_frame,
            text="Send",
            command=self.send_prompt,
            bg="#4a90e2",
            fg="#ffffff",
            font=("Segoe UI", 10),
            padx=15
        )
        self.send_button.pack(side=tk.RIGHT)
        
        self.clear_button = tk.Button(
            self.input_frame,
            text="Clear",
            command=self.clear_chat,
            bg="#e74c3c",
            fg="#ffffff",
            font=("Segoe UI", 10),
            padx=15
        )
        self.clear_button.pack(side=tk.RIGHT, padx=(0, 5))
        self.command_frame = tk.Frame(self.top_frame, bg="#1e1e1e")
        self.command_frame.pack(fill=tk.X)
        self.google_button = tk.Button(
            self.command_frame,
            text="Google",
            command=self.google_search,
            bg="#3498db",
            fg="#ffffff",
            font=("Segoe UI", 9),
            padx=10
        )
        self.google_button.pack(side=tk.LEFT, padx=5)
        self.write_word_button = tk.Button(
            self.command_frame,
            text="Write Word",
            command=self.write_word,
            bg="#2ecc71",
            fg="#ffffff",
            font=("Segoe UI", 9),
            padx=10
        )
        self.write_word_button.pack(side=tk.LEFT, padx=5)
        self.bypass_router_button = tk.Button(
            self.command_frame,
            text="Advanced Tasks: OFF",
            command=self.toggle_bypass_router,
            bg="#f39c12",
            fg="#ffffff",
            font=("Segoe UI", 9),
            padx=10
        )
        self.bypass_router_button.pack(side=tk.LEFT, padx=5)
        self.dataframe_reader_button = tk.Button(
            self.command_frame,
            text="Dataframe Reader: OFF",
            command=self.toggle_dataframe_reader,
            bg="#e74c3c",
            fg="#ffffff",
            font=("Segoe UI", 9),
            padx=10
        )
        self.dataframe_reader_button.pack(side=tk.LEFT, padx=5)
        self.file_button = tk.Button(
            self.command_frame,
            text="File",
            command=self.show_file_menu,
            bg="#9b59b6",
            fg="#ffffff",
            font=("Segoe UI", 9),
            padx=10
        )
        self.file_button.pack(side=tk.LEFT, padx=5)
        


        self.btn_print_files = tk.Button(
            self.command_frame,
            text="Print Files",
            command=self.print_file_list,
            bg="#7f8c8d",
            fg="#ffffff",
            font=("Segoe UI", 9),
            padx=10
        )
        self.btn_print_files.pack(side=tk.LEFT, padx=5)
        self.toggle_pin_button = tk.Button(
            self.command_frame,
            text="Pinned: OFF",
            command=self.toggle_pin,
            bg="#e15526",
            fg="#000000",
            font=("Segoe UI", 9),
            padx=10
        )
        self.toggle_pin_button.pack(side=tk.RIGHT, padx=5)

        self.coding_button = tk.Button(
            self.command_frame,
            text="Coding Mode",
            command=self.coding_mode,
            bg="#0b0a0b",
            fg="#78E753",
            font=("Segoe UI", 9),
            padx=10
        )
        self.coding_button.pack(side=tk.RIGHT, padx=5)

        self.btn_memory = tk.Button(
            self.command_frame,
            text="Memory Manager",
            command=restore_memory,
            bg="#0d2ebe",
            fg="#ffffff",
            font=("Segoe UI", 9),
            padx=10
        )
        self.btn_memory.pack(side=tk.RIGHT, padx=5)


        self.btn_config = tk.Button(
            self.command_frame,
            text="LLM Config",
            command=lambda: open_custom_llm_ui(self.root),
            bg="#5e35b1", 
            fg="#ffffff",
            font=("Segoe UI", 9),
            padx=10
        )
        self.btn_config.pack(side=tk.RIGHT, padx=5)


        self.loading_label = tk.Label(
            self.root,
            text="",
            fg="#aaaaaa",
            bg="#1e1e1e",
            font=("Arial", 10, "italic")
        )
        self.loading_label.pack(pady=(0,5))
        self.loading_animation = False
        self.loading_dots = 0
        self.setup_text_tags()

    def open_drag_drop_dialog(self):
        """Opens the integrated manager that remembers files and allows RAG activation."""
        # We pass the CURRENT lists so the window shows them immediately
        dialog = MaiaAttachmentsDialog(self.root, self.attached_files_list, self.active_rag_files)
        
        # This stops the code here until the window is closed
        self.root.wait_window(dialog)

        # If user clicked Confirm, update the main app's state
        if dialog.confirmed:
            self.attached_files_list = dialog.attachments
            self.active_rag_files = dialog.rag_active
            
            # 1. Handle RAG indexing if there are green files
            if self.active_rag_files:
                self._start_loading_animation()
                threading.Thread(target=self._process_rag_embedding, daemon=True).start()
            
            # 2. Update the "Dump" context (files not in RAG are passed as raw text)
            self._prepare_standard_context()

    def _process_rag_embedding(self):
        """Background thread to handle RAG folder indexing."""
        for file_path in self.active_rag_files:
            self.doc_rag_manager.embed_file(file_path)
        
        self.root.after(0, self._stop_loading_animation)
        self.root.after(0, lambda: self._insert_formatted(f"RAG Activated for {len(self.active_rag_files)} files.", role="system"))


    def _prepare_standard_context(self):
        """Handles the standard 'file dump' for files not selected for RAG."""
        all_content = ""
        file_count = 0
        
        for path in self.attached_files_list:
            # Skip files that are being handled by the RAG manager
            if path in self.active_rag_files: 
                continue 
            
            # INITIALIZE content to None at the start of every loop
            content = None 
            ext = os.path.splitext(path)[1].lower()
            
            try:
                if ext == ".pdf":
                    content = read_pdf(path)
                elif ext == ".docx":
                    content = read_word(path)
                elif ext in [".xlsx", ".xls"]:
                    res = read_excel_enhanced(path)
                    content = res[0] if isinstance(res, tuple) else res
                elif ext in [".py", ".txt", ".json", ".md", ".csv"]:
                    with open(path, "r", encoding="utf-8") as f:
                        content = f.read()
                
                # Now 'content' is guaranteed to exist (either as a string or None)
                if content and content.strip():
                    all_content += f"\n\n--- ATTACHMENT: {os.path.basename(path)} ---\n"
                    all_content += content
                    file_count += 1
            except Exception as e:
                logger.error(f"Error reading {path} for standard context: {e}")

        if all_content:
            self.pending_file_content = all_content
            self._insert_formatted(f"Standard context updated with {file_count} file(s).", role="system")
        else:
            self.pending_file_content = None


    def open_rag_selection_ui(self):
        """Window to manage attached files and activate RAG."""
        rag_win = tk.Toplevel(self.root)
        rag_win.title("Attached Files & RAG Manager")
        rag_win.geometry("1200x1400")
        rag_win.configure(bg="#1e1e1e")

        tk.Label(rag_win, text="ATTACHED FILES", bg="#1e1e1e", fg="#4ea1ff", font=("Arial", 12, "bold")).pack(pady=10)
        tk.Label(rag_win, text="Click a file to toggle 'Activate RAG' (Green = Selected)", 
                 bg="#1e1e1e", fg="#aaaaaa", font=("Arial", 9, "italic")).pack()

        # Scrollable area for file buttons
        container = tk.Frame(rag_win, bg="#1e1e1e")
        container.pack(fill=tk.BOTH, expand=True, padx=20, pady=10)
        
        canvas = tk.Canvas(container, bg="#1e1e1e", highlightthickness=0)
        scrollbar = ttk.Scrollbar(container, orient="vertical", command=canvas.yview)
        scroll_frame = tk.Frame(canvas, bg="#1e1e1e")

        scroll_frame.bind("<Configure>", lambda e: canvas.configure(scrollregion=canvas.bbox("all")))
        canvas.create_window((0, 0), window=scroll_frame, anchor="nw", width=540)
        canvas.configure(yscrollcommand=scrollbar.set)
        
        canvas.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        scrollbar.pack(side=tk.RIGHT, fill=tk.Y)

        # Draw buttons for all persistent files
        for file_path in self.attached_files_list:
            fname = os.path.basename(file_path)
            
            # Use green background if file is already active
            initial_bg = "#27ae60" if file_path in self.active_rag_files else "#34495e"
            
            btn = tk.Button(scroll_frame, text=fname, bg=initial_bg, fg="white", 
                            font=("Segoe UI", 10), anchor="w", padx=10, pady=8, relief="flat")
            
            def toggle(p=file_path, b=btn):
                if p in self.active_rag_files:
                    self.active_rag_files.remove(p)
                    b.config(bg="#34495e")
                else:
                    self.active_rag_files.add(p)
                    b.config(bg="#27ae60") # GREEN BACKGROUND when selected

            btn.config(command=toggle)
            btn.pack(fill=tk.X, pady=2)

        def confirm_and_embed():
            rag_win.destroy()
            if not self.active_rag_files:
                self._insert_formatted("No files selected for RAG. Files remain attached as raw context.", role="system")
                return

            self._start_loading_animation()
            self._insert_formatted(f"Indexing {len(self.active_rag_files)} files for RAG...", role="system")
            
            def background_work():
                # Clear manager and rebuild only selected files
                self.doc_rag_manager.active_indices = {}
                for f in self.active_rag_files:
                    try:
                        self.doc_rag_manager.embed_file(f)
                    except Exception as e:
                        logger.error(f"Error embedding {f}: {e}")
                
                self.root.after(0, self._stop_loading_animation)
                self.root.after(0, lambda: self._insert_formatted("RAG Activated. These files will now be searched semantically.", role="system"))

            threading.Thread(target=background_work, daemon=True).start()

        # Bottom Buttons
        btn_frame = tk.Frame(rag_win, bg="#1e1e1e")
        btn_frame.pack(fill=tk.X, pady=20, padx=20)

        tk.Button(btn_frame, text="Add More Files", command=lambda: [rag_win.destroy(), self.open_drag_drop_dialog()],
                  bg="#3498db", fg="white", font=("Arial", 10, "bold"), width=15).pack(side=tk.LEFT)
        
        tk.Button(btn_frame, text="ACTIVATE RAG", command=confirm_and_embed,
                  bg="#2ecc71", fg="white", font=("Arial", 10, "bold"), width=25).pack(side=tk.RIGHT)

    def print_file_list(self):
        """Prints the list of currently attached files to console and chat."""
        if not self.attached_files_list:
            print("No files attached.")
            self._insert_formatted("No files currently in memory.", role="system")
            return

        print("-" * 30)
        print(f"Attached Files ({len(self.attached_files_list)}):")
        msg = "Current Files:\n"
        for f in self.attached_files_list:
            print(f)
            msg += f"- {os.path.basename(f)}\n"
        print("-" * 30)
        self._insert_formatted(msg, role="system")



    def open_creative_mode(self):
        creative_window=tk.Toplevel(self.root, bg="#1e1e1e")
        creative_window.title("Creative Mode")
        creative_window.geometry("3200x2400")
        tk.Label(creative_window, text="Choose an option:", bg="#1e1e1e", fg="#ffffff").pack(pady=10)
        button_frame = tk.Frame(creative_window, bg="#1e1e1e")
        button_frame.pack(pady=10)

        
        def img_audio_to_vid():
            print("OK")



        tk.Button(
            button_frame,
            text="Img+Audio to Vid",
            command=img_audio_to_vid,
            bg="#3498db",
            fg="#ffffff",
            padx=10
        ).pack(side=tk.LEFT, padx=5)
       
    def toggle_pin(self):
        """Toggle pinning the chat window on top"""
        global pinned
        pinned = not pinned
        self.root.attributes("-topmost", pinned)
        if pinned:
            self.toggle_pin_button.config(bg="#27ae60", text="Pinned: ON")
        else:
            self.toggle_pin_button.config(bg="#e74c3c", text="Pinned: OFF")

    def coding_mode(self):
        """Open coding mode window"""
        coding_window = tk.Toplevel(self.root, bg="#1e1e1e")
        coding_window.title("Coding Mode")
        coding_window.geometry("3200x2400")
        tk.Label(coding_window, text="Coding Mode Activated", bg="#1e1e1e", fg="#12e220").pack(pady=10)
        # Additional coding mode UI components can be added here
        codeduo_button = tk.Button(
            coding_window,
            text="CodeDuo",
            command=self.ask_codeduo_query,
            bg="#000000",
            fg="#B300F4",
            padx=10
        )
        codeduo_button.pack(pady=5)
    def ask_codeduo_query(self):
        """Ask a coding question to CodeDuo"""
        question = simpledialog.askstring("CodeDuo", "Enter your coding request:")
        if question:
            self._insert_formatted(f"CodeDuo: {question}", role="user")
            self._insert_formatted("Please wait...", role="assistant")
            coding_window = self.root.winfo_children()[-1]
            coding_window.destroy()
            asyncio.run(self.send_codeduo_question(question))
            
    async def send_codeduo_question(self, question):
        await asyncio.sleep(3)
        code_filename=codeduo(question)
        self._insert_formatted(f"Code ultimated. You can find it at {code_filename}", role="assistant")



    def toggle_dataframe_reader(self):
        """Toggle Dataframe Reader mode"""
        global dataframe_reader_mode
        dataframe_reader_mode = not dataframe_reader_mode
        if dataframe_reader_mode:
            self.dataframe_reader_button.config(bg="#27ae60", text="Dataframe Reader: ON")
        else:
            self.dataframe_reader_button.config(bg="#e74c3c", text="Dataframe Reader: OFF")

    def toggle_bypass_router(self):
        """Toggle bypass router mode"""
        global routing
        routing = not routing
        if routing:
            self.bypass_router_button.config(bg="#27ae60", text="Advanced Tasks: ON")
        else:
            self.bypass_router_button.config(bg="#f39c12", text="Advanced Tasks: OFF")

    def setup_text_tags(self):
        """Setup text tags for different content types"""
        self.chat_display.tag_configure("user", foreground="#35FFFC", font=("Segoe UI", 11, "bold"), justify="right")
        self.chat_display.tag_configure("user_message", foreground="#b9e7f0", font=("Segoe UI", 10), justify="right")
        self.chat_display.tag_configure("assistant", foreground="#fcd88c", font=("Segoe UI", 11, "bold"))
        self.chat_display.tag_configure("header", font=("Segoe UI", 12, "bold"), foreground="#4ea1ff")
        self.chat_display.tag_configure("list", font=("Segoe UI", 10), foreground="#bdc3c7")
        self.chat_display.tag_configure("code", font=("Consolas", 10), background="#3a3a3a", foreground="#15f5aa", selectbackground="#8E3C7C")
        self.chat_display.tag_configure("italic", font=("Segoe UI", 10, "italic"))
        self.chat_display.tag_configure("link", foreground="#00aeff", underline=True, font=("Segoe UI", 10, "bold"))
        self.chat_display.tag_configure("bold", font=("Segoe UI", 10, "bold"))
        self.chat_display.tag_configure("system", foreground="#7a9983", font=("Segoe UI", 10, "bold"))
        self.chat_display.tag_configure("system_user", foreground="#7a9983", font=("Segoe UI", 10, "bold"), justify="right")

    def bind_events(self):
        """Bind events to UI components"""
        self.root.bind("<Configure>", self.on_window_resize)
        self.input_entry.bind("<Return>", self._on_return_pressed)
    def _on_return_pressed(self, event):
        """Handle Ctrl+Enter to send, Enter/Shift+Enter for newline"""
        # Check if Control key is pressed (Mask 0x0004)
        if (event.state & 0x0004):
            self.send_prompt(event)
            return "break"  # Prevent newline insertion from the Enter key
        
        # For standard Enter or Shift+Enter, return None to allow 
        # the default Tkinter behavior (inserting a newline)
        return None



    def on_window_resize(self, event):
        """Handle window resize events"""
        pass

    def send_prompt(self, event=None):
        """Enhanced send prompt with semantic memory integration."""
        # 1. Capture the raw user query specifically
        raw_prompt = self.input_entry.get("1.0", tk.END).strip()
        
        if not raw_prompt:
            return
            
        self.input_entry.delete("1.0", tk.END)
        self.prefix_added = False
        
        last_response = self.memory_manager.get_last_assistant_message()
        last_question = self.memory_manager.get_last_user_message() 

        # 2. Build the enhanced_prompt (Raw query + Attached File Context)
        enhanced_prompt = raw_prompt

        if self.excel_structure and self.pending_file_content:
            excel_context = create_excel_query_context(self.excel_structure, raw_prompt)
            enhanced_prompt = f"{excel_context}\n{self.pending_file_content}\nUser Query: {raw_prompt}"
            # Clear pending content after use
            self.pending_file_content = None
            
        elif self.pending_file_content:
            enhanced_prompt = f"{self.pending_file_content}\n{raw_prompt}"
            # Clear pending content after use
            self.pending_file_content = None
            
        elif self.excel_sheets_data is not None:
            summary = "Excel Data Summary:\n"
            for sheet_name, sheet_data in self.excel_sheets_data.items():
                summary += f"Sheet: {sheet_name}\n"
                summary += f"Shape: {sheet_data.shape}\n"
                summary += f"Columns: {list(sheet_data.columns)}\n"
                summary += "Sample data (first 5 rows):\n"
                summary += sheet_data.head(5).to_string(index=False, max_cols=5, max_colwidth=20)
                summary += "\n"
            enhanced_prompt = f"{summary}\nUser Query: {raw_prompt}"
            
        # 3. Handle Slash Commands (using raw_prompt)
        if raw_prompt.lower().startswith("/"):
            self.handle_command_prompt(raw_prompt, last_response, last_question)
            return
            
        # 4. Update UI (Show only what the user typed, not the hidden context)
        self._insert_formatted(raw_prompt, role="user")
        
        # 5. Handle specialized Dataframe Reader Mode
        if dataframe_reader_mode and self.current_excel_path:
            python_result, python_script = self.handle_dataframe_reader_mode(raw_prompt)
            
            # Manually save interaction to memory since we bypass the main pipeline
            user_message = HumanMessage(content=raw_prompt)
            ai_content = f"Path to the file: {self.current_excel_path}\nGenerated Python script: {python_script}\nResult: {python_result.stdout if hasattr(python_result, 'stdout') else str(python_result)}"
            ai_message = AIMessage(content=ai_content)
            
            self.memory_manager.add_message(user_message)
            self.memory_manager.add_message(ai_message)
            self.memory_manager.save_memory()
            
            self.memory_manager.log_event("dataframe_execution", {
                "prompt": raw_prompt,
                "script_length": len(python_script),
                "success": "error" not in ai_content.lower()
            })
        else:
            # 6. Start Processing: Pass BOTH enhanced_prompt (for LLM) AND raw_prompt (for Metadata/Topics)
            self._start_processing_prompt(enhanced_prompt, last_response, last_question, raw_prompt=raw_prompt)

    def handle_dataframe_reader_mode(self, prompt):
        """Handle Dataframe Reader mode by generating and executing Python code."""
        self._start_loading_animation()
        self.prefix_added = False
        for i in range(5):
            print(f"Coding Attempt {i}...")
            script_content = self.generate_dataframe_script(prompt, error_msg=self.code_error)
            script_path = os.path.join(OUTPUT_DIR, "excel_script.py")
            try:
                script_content = extract_python_code(script_content)
                with open(script_path, "w") as f:
                    f.write(script_content)
                result = subprocess.run(
                    ["python", script_path, self.current_excel_path],
                    capture_output=True,
                    text=True,
                    check=True
                )
                self._insert_formatted(result.stdout, role="assistant")
                self.code_error = ""
                self._stop_loading_animation()
                return result, script_content
            except subprocess.CalledProcessError as e:
                error_msg = f"Error executing script: {str(e)}\nScript output:\n{e.stdout}\nError output:\n{e.stderr}"
                self._insert_formatted(error_msg, role="assistant")
                self.code_error = error_msg
            except Exception as e:
                self._insert_formatted(f"Error in Dataframe Reader mode: {str(e)}", role="assistant")
                self.code_error = str(e)
        self._stop_loading_animation()
        return None, script_content

    def generate_dataframe_script(self, prompt, error_msg=""):
        """Generate a Python script to process Excel data based on the prompt."""
        script_template = """import pandas as pd
import sys

def process_excel(file_path):
    # Load the Excel file
    df = pd.read_excel(file_path)
    {code_block}
    return result

if __name__ == "__main__":
    file_path = sys.argv[1]
    try:
        result = process_excel(file_path)
        print(result)
    except Exception as e:
        print(f"Error processing Excel file: {{str(e)}}")
"""
        code_block = self.generate_code_from_prompt(prompt, error_msg)
        return script_template.format(code_block=code_block)

    def generate_code_from_prompt(self, prompt, error_msg=""):
        """Generate Python code based on the user's prompt."""
        output_code = asyncio.run(self.coder.generate_code(prompt, file_path=self.current_excel_path, error_msg=error_msg))
        return output_code

    def handle_command_prompt(self, prompt, last_response, last_question):
        """Handle command prompts like /google, /write_word, etc."""
        command = prompt.split(' ')[0][1:].lower()
        remaining_prompt = ' '.join(prompt.split(' ')[1:]) if len(prompt.split(' ')) > 1 else ""
        if command == "google":
            self.google_search(remaining_prompt)
        elif command == "writeword":
            self.write_word()
        elif command == "clear":
            self.clear_chat()
        elif command == "excel" and self.excel_structure:
            self.show_excel_analytics()
        else:
            self._insert_formatted(prompt, role="user")
            self._start_processing_prompt(prompt, last_response, last_question)

    def _start_processing_prompt(self, prompt, last_response, last_question, raw_prompt=None):
        """Start processing a prompt with appropriate routing"""
        self._start_loading_animation()
        self.prefix_added = False
        
        # If raw_prompt wasn't passed (legacy call), assume prompt is raw
        if raw_prompt is None: 
            raw_prompt = prompt

        thread = threading.Thread(
            target=self._process_prompt_async,
            # Add raw_prompt to arguments
            args=(prompt, last_response, last_question, raw_prompt) 
        )
        thread.daemon = True
        thread.start()

    def _process_prompt_async(self, prompt, last_response, last_question, raw_prompt):
        """Process prompt asynchronously"""
        global routing
        if routing:
            try:
                # Router uses raw_prompt to decide, not the massive file dump
                route_result = asyncio.run(self.router.route(raw_prompt, last_response))
                if route_result == "google":
                    self._google_search_async(raw_prompt, last_response, last_question)
                elif route_result == "writeword":
                    self._write_word_async()
                else:
                    self._normal_process_async(prompt, raw_prompt)
            except Exception as e:
                logger.error(f"Processing error: {e}")
                self._insert_formatted(f"Error: {str(e)}", role="assistant")
                self._stop_loading_animation()
        else:
            self._normal_process_async(prompt, raw_prompt)

    def _normal_process_async(self, prompt, raw_prompt):
        try:
            asyncio.run(
                self.prompthandler.crew_prompt(
                    prompt, 
                    callback=self._normal_stream_response, 
                    use_semantic_memory=True,
                    raw_query=raw_prompt,
                    doc_rag_manager=self.doc_rag_manager 
                )
            )
        except Exception as e:
            logger.error(f"Processing error: {e}")
            self._insert_formatted(f"Error: {str(e)}", role="assistant")
            self._stop_loading_animation()

    def _normal_stream_response(self, response_text, role):
        """Handle streaming for normal prompts with proper AI prefix"""
        if role is None:
            self._stop_loading_animation()
            self.chat_display.config(state="normal")
            self.chat_display.insert("end", "\n")
            self.chat_display.config(state="disabled")
            self.prompthandler.save_memory()
            self.prefix_added = False
            self._reload_from_memory()
        elif role == "assistant" and response_text:
            if not self.prefix_added:
                timestamp = datetime.now().strftime("%H:%M")
                self.chat_display.config(state="normal")
                self._insert_formatted(f"\n[{timestamp}]", role="system")
                #self._insert_formatted(f"M.A.I.A.: ", role="assistant")
                self.chat_display.config(state="disabled")
                self.prefix_added = True
            self._stream_text_only(response_text)

    def process_file_list(self, files):
        """Centralized logic to process a list of files and append to memory."""
        if not files:
            return

        processed_text = ""
        file_count = 0
        self._start_loading_animation()

        try:
            for file_path in files:
                file_path = os.path.normpath(file_path)
                if not os.path.exists(file_path): continue
                
                # Update the tracked file list
                if file_path not in self.attached_files_list:
                    self.attached_files_list.append(file_path)
                
                ext = os.path.splitext(file_path)[1].lower()
                content = None

                if ext == ".pdf":
                    content = read_pdf(file_path)
                elif ext == ".docx":
                    content = read_word(file_path)
                elif ext in [".xlsx", ".xls"]:
                    res = read_excel_enhanced(file_path)
                    if isinstance(res, tuple):
                        content = res[0]
                        if not self.excel_structure: self.excel_structure = res[1]
                        if not self.excel_sheets_data: self.excel_sheets_data = res[2]
                        self.current_excel_path = file_path
                    else:
                        content = res
                elif ext in [".py", ".txt", ".json", ".md", ".csv"]:
                    try:
                        with open(file_path, "r", encoding="utf-8") as f:
                            content = f.read()
                    except Exception as e:
                        logger.error(f"Error reading text file: {e}")

                if content:
                    processed_text += f"\n\n--- ATTACHMENT: {os.path.basename(file_path)} ---\n"
                    processed_text += content
                    file_count += 1
            
            if processed_text:
                if self.pending_file_content:
                    self.pending_file_content += processed_text
                else:
                    self.pending_file_content = processed_text
                
                self._insert_formatted(f"Attached {file_count} file(s) to next prompt.", role="system")
            
        except Exception as e:
            self._insert_formatted(f"Error processing files: {e}", role="assistant")
        finally:
            self._stop_loading_animation()



    # Now we simplify the individual uploaders to use the new logic:
    def upload_pdf(self):
        file_path = filedialog.askopenfilename(filetypes=[("PDF files", "*.pdf")])
        if file_path: self.process_file_list([file_path])

    def upload_excel(self):
        file_path = filedialog.askopenfilename(filetypes=[("Excel files", "*.xlsx *.xls")])
        if file_path: self.process_file_list([file_path])

    def upload_word(self):
        file_path = filedialog.askopenfilename(filetypes=[("Word files", "*.docx")])
        if file_path: self.process_file_list([file_path])

    def _google_search_async(self, query, last_response="", last_question=""):
        """Perform Google search asynchronously"""
        try:
            # Get Config
            retrieval_conf = LLM_CONFIG.get("retrieval", DEFAULT_CONFIG["retrieval"])
            num_links_val = int(retrieval_conf.get("google_k", 5))

            refined_query = asyncio.run(self.refiner.refine_query(query, last_response, last_question))
            refined_query = refined_query.replace('"', "'")
            logger.info(f"Googling query: {refined_query} (Limit: {num_links_val})")
            
            try:
                # --- PASS THE PARAMETER HERE ---
                embedded_prompt = inquire(refined_query, num_links=num_links_val)
                
                if not embedded_prompt or not isinstance(embedded_prompt, str):
                    logger.warning("inquire() did not produce valid prompt, using refined_query.")
                    embedded_prompt = f"Answer the following refined query:\n{refined_query}"
            except Exception as e:
                logger.error(f"Inquire error: {e}")
                embedded_prompt = f"Answer the following refined query:\n{refined_query}"
            
            asyncio.run(
                self.prompthandler.crew_prompt(embedded_prompt, callback=self._google_stream_response, use_semantic_memory=True)
            )
        except Exception as e:
            logger.error(f"Search error: {str(e)}")
            self._insert_formatted(f"Search error: {str(e)}", role="assistant")
            self._stop_loading_animation()



    def _google_stream_response(self, response_text, role):
        """Handle streaming for Google search results"""
        if role is None:
            self._stop_loading_animation()
            self.chat_display.config(state="normal")
            self.chat_display.insert("end", "\n")
            self.chat_display.config(state="disabled")
            self.prompthandler.save_memory()
            self.prefix_added = False
            self._reload_from_memory()
        elif role == "assistant" and response_text:
            if not self.prefix_added:
                timestamp = datetime.now().strftime("%H:%M")
                self.chat_display.config(state="normal")
                self.chat_display.insert("end", f"[{timestamp}] M.A.I.A.: ")
                self.chat_display.config(state="disabled")
                self.prefix_added = True
            self._stream_text_only(response_text)

    def _write_word_async(self):
        """Write word asynchronously"""
        try:
            last_response = self.memory_manager.get_last_assistant_message()
            if last_response:
                writeword(last_response)
                self._insert_formatted("Content saved to Word document.", role="assistant")
            else:
                self._insert_formatted("No content to save.", role="assistant")
        except Exception as e:
            logger.error(f"Write Word error: {str(e)}")
            self._insert_formatted(f"Error: {str(e)}", role="assistant")
        finally:
            self._stop_loading_animation()

    def google_search(self, query=None):
        """Initiate Google search"""
        if query is None:
            query = simpledialog.askstring("Google Search", "Enter search query:")
            if not query:
                return
        self._insert_formatted(f"Searching Google for: {query}", role="user")
        self._start_loading_animation()
        self.prefix_added = False
        thread = threading.Thread(target=self._direct_google_search, args=(query,))
        thread.daemon = True
        thread.start()

    def _direct_google_search(self, query):
        """Perform direct Google search without routing or refining"""
        try:
            logger.info(f"Direct googling query: {query}")
            try:
                embedded_prompt = inquire(query)
                if not embedded_prompt or not isinstance(embedded_prompt, str):
                    logger.warning("inquire() did not produce valid prompt, using original query.")
                    embedded_prompt = f"Answer the following query:\n{query}"
            except:
                embedded_prompt = f"Answer the following query:\n{query}"
            asyncio.run(
                self.prompthandler.crew_prompt(embedded_prompt, callback=self._google_stream_response, use_semantic_memory=True)
            )
        except Exception as e:
            logger.error(f"Direct search error: {str(e)}")
            self._insert_formatted(f"Search error: {str(e)}", role="assistant")
            self._stop_loading_animation()

    def write_word(self):
        """Initiate word writing"""
        try:
            last_response = self.memory_manager.get_last_assistant_message()
            if last_response:
                writeword(last_response)
                self._insert_formatted("Content saved to Word document.", role="assistant")
            else:
                self._insert_formatted("No content to save.", role="assistant")
        except Exception as e:
            self._insert_formatted(f"Error saving to Word: {str(e)}", role="assistant")

    def _insert_formatted(self, text: str, role="assistant"):
        """Insert formatted text into chat display with proper markdown processing"""
        self.chat_display.config(state="normal")
        if role == "system_user":
            self.chat_display.insert("end", f"{text}\n", "system_user")
        if role == "user":
            self.chat_display.insert("end", "\U0001F64B"+f"You: \n", "user")
            self.chat_display.insert("end", f"{text}\n", "user_message")
        elif role == "system":
            self.chat_display.insert("end", f"{text}\n", "system")
        elif role == "assistant":
            self.chat_display.insert("end", "\U0001F916"+f" M.A.I.A.: \n", "assistant")
            code_blocks = re.split(r"(```.*?```)", text, flags=re.S)
            for block in code_blocks:
                if block.startswith("```") and block.endswith("```"):
                    code_lines = block.strip("`").splitlines()
                    if len(code_lines) > 0 and re.match(r"^\w+$", code_lines[0]):
                        code_lines = code_lines[1:]
                    for line in code_lines:
                        line = line.rstrip()
                        url_pattern = r"(https?://[^\s\]\)\*`]+|www\.[^\s\]\)\*`]+)"
                        last_pos = 0
                        for m in re.finditer(url_pattern, line):
                            raw_url = m.group()
                            clean_link = clean_url(raw_url.strip("*_`"))
                            tag_name = f"link-{clean_link}"
                            if m.start() > last_pos:
                                self.chat_display.insert("end", line[last_pos:m.start()], ("code",))
                            self.chat_display.insert("end", clean_link, ("code", "link", tag_name))
                            self.chat_display.tag_bind(tag_name, "<Button-1>",
                                                        lambda e, url=clean_link: self._open_link(url))
                            last_pos = m.end()
                        if last_pos < len(line):
                            self.chat_display.insert("end", line[last_pos:] + "\n", ("code",))
                        else:
                            self.chat_display.insert("end", "\n", ("code",))
                else:
                    lines = block.split("\n")
                    for line in lines:
                        line = line.rstrip()
                        clean_line = line.strip()
                        url_pattern = r"(https?://[^\s\]\)\*`]+|www\.[^\s\]\)\*`]+)"
                        if re.search(url_pattern, clean_line):
                            last_pos = 0
                            for m in re.finditer(url_pattern, clean_line):
                                raw_url = m.group()
                                clean_link = clean_url(raw_url.strip("*_`"))
                                tag_name = f"link-{clean_link}"
                                if m.start() > last_pos:
                                    self.chat_display.insert("end", clean_line[last_pos:m.start()])
                                self.chat_display.insert("end", clean_link, ("link", tag_name))
                                self.chat_display.tag_bind(tag_name, "<Button-1>",
                                                            lambda e, url=clean_link: self._open_link(url))
                                last_pos = m.end()
                            if last_pos < len(clean_line):
                                self.chat_display.insert("end", clean_line[last_pos:] + "\n")
                            else:
                                self.chat_display.insert("end", "\n")
                        else:
                            cleaned_text = clean_line.replace("`", "")
                            if cleaned_text.startswith("#"):
                                header_text = cleaned_text.lstrip("# ").strip()
                                if "**" in header_text:
                                    process_markdown_bold(header_text, self.chat_display)
                                    self.chat_display.insert("end", "\n")
                                else:
                                    self.chat_display.insert("end", header_text + "\n", ("header",))
                            elif cleaned_text.startswith("- "):
                                list_text = cleaned_text[2:].strip()
                                if "**" in list_text:
                                    self.chat_display.insert("end", "- ")
                                    process_markdown_bold_simple(list_text, self.chat_display)
                                    self.chat_display.insert("end", "\n")
                                else:
                                    self.chat_display.insert("end", f"- {list_text}\n", ("list",))
                            elif "**" in cleaned_text:
                                process_markdown_bold_simple(cleaned_text, self.chat_display)
                                self.chat_display.insert("end", "\n")
                            elif "_" in cleaned_text:
                                parts = re.split(r"(_.*?_)", cleaned_text)
                                for p in parts:
                                    if p.startswith("_") and p.endswith("_"):
                                        self.chat_display.insert("end", p.strip("_"), ("italic",))
                                    else:
                                        self.chat_display.insert("end", p)
                                self.chat_display.insert("end", "\n")
                            else:
                                self.chat_display.insert("end", cleaned_text + "\n")
            self.chat_display.insert("end", "\n")
        self.chat_display.see("end")
        self.chat_display.config(state="disabled")

    def _open_link(self, url):
        """Open URL in browser"""
        try:
            webbrowser.open(url)
        except Exception as e:
            logger.error(f"Error opening link: {str(e)}")

    def clear_chat(self):
        """Clear the chat display and start a new chat"""
        self.prompthandler.save_memory()
        # Finalize embeddings for current session
        if self.embedding_manager:
            self.embedding_manager.finalize_session(self.memory_file)
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        new_memory_file = os.path.join(OUTPUT_DIR, f"autochat_{timestamp}.json")
        self.memory_file = new_memory_file
        self.memory_manager.memory_file = new_memory_file
        self.memory_manager.messages = []
        self.memory_manager.event_log = []
        self.memory_manager.last_embedded_count = 0
        self.chat_display.config(state="normal")
        self.chat_display.delete(1.0, "end")
        self.chat_display.config(state="disabled")
        self.excel_structure = None
        self.excel_sheets_data = None
        self.pending_file_content = None
        self.pending_dataframe = None

    def load_chat(self):
        """Load existing chat from memory"""
        self._reload_from_memory()


    def _reload_from_memory(self):
        """Reload chat display from memory"""
        self.chat_display.config(state="normal")
        self.chat_display.delete("1.0", "end")
        for message in self.memory_manager.messages:
            if not isinstance(message, SystemMessage):
                msg_timestamp=datetime.fromisoformat(message.additional_kwargs["timestamp"]).strftime("%d/%m/%Y - %H:%M")
            if isinstance(message, HumanMessage):
                self._insert_formatted(f"\n[{msg_timestamp}]", role="system_user")
                self._insert_formatted(message.content, role="user")
            elif isinstance(message, AIMessage):
                self._insert_formatted(f"\n[{msg_timestamp}]", role="system")
                self._insert_formatted(message.content, role="assistant")
        self.chat_display.config(state="disabled")

    def save_chat(self):
        """Save current chat to file"""
        self.prompthandler.save_memory()
        if self.embedding_manager:
            self.embedding_manager.finalize_session(self.memory_file)

    def _start_loading_animation(self):
        """Start loading animation"""
        self.loading_animation = True
        self.loading_dots = 0
        self._animate_loading()

    def _stop_loading_animation(self):
        """Stop loading animation"""
        self.loading_animation = False
        self.loading_label.config(text="")

    def _animate_loading(self):
        """Animate loading indicator"""
        if self.loading_animation:
            dots = "." * (self.loading_dots % 4)
            self.loading_label.config(text="\U0001F916"+f" M.A.I.A. is thinking{dots}")
            self.loading_dots += 1
            self.root.after(500, self._animate_loading)


class RouterAgent:
    def __init__(self, model=DEFAULT_MODEL):
        # We don't store the client in self.client anymore to avoid staleness
        pass

    async def route(self, prompt: str, last_response: str = "") -> str:
        client = get_llm_client("router")
        if client is None: return "normalprompt"
        
        logger.info(f"Routing prompt: {prompt}")
        system_msg = SystemMessage(content=""" You're an AI router. You must decide what command needs to be attached to the user's prompt among the following options: 'writeword' if the prompt asks to print or save in Word the previous answer; 'google' if the prompt asks to search for extra information or to google something. Take note our internal knowledge is not updated, so this command is also recommended for updates, latest news, innovation and breakthroughs; 'normalprompt' in all the other cases. Answer ONLY with the chosen command. Verbose = 0. """)
        user_msg = HumanMessage(content=f"Prompt: {prompt}\nLast Answer: {last_response}")
        
        try:
            response = ""
            # Unified stream
            async for text in unified_astream(client, [system_msg, user_msg]):
                response += text
                
            command = response.strip().lower()
            if "writeword" in command: return "writeword"
            if "google" in command: return "google"
            return "normalprompt"
        except Exception as e:
            logger.error(f"Error routing: {str(e)}")
            return "normalprompt"

class QueryRefinerAgent:
    def __init__(self, model=DEFAULT_MODEL):
        pass

    async def refine_query(self, prompt: str, last_response="", last_question="") -> str:
        client = get_llm_client("refiner")
        if client is None: return prompt
        
        logger.info(f"Refining query: {prompt}")
        system_msg = SystemMessage(content=f""" You are a query refiner. Your task is to take the user's prompt and your last response and refine it into a better Google search query. The refined query should be concise, specific, and optimized for web search. Do NOT insert dates or years unless explicitly asked by the user. For context, here is the last response you gave to the user: {last_response}. Here is the last question the user asked: {last_question}. Answer ONLY with ONE and ONLY refined query. Do not add any explanations or additional text. """)
        user_msg = HumanMessage(content=f"Prompt: {prompt}")
        try:
            response = ""
            async for text in unified_astream(client, [system_msg, user_msg]):
                response += text
            return response.strip()
        except Exception as e:
            logger.error(f"Error refining query: {str(e)}")
            return prompt
        
   
class CodingAgent:
    def __init__(self, model=DEFAULT_MODEL):
        pass

    async def generate_code(self, prompt: str, last_response="", last_question="", file_path="", sample_data="", error_msg="") -> str:
        client = get_llm_client("coder")
        if client is None: return prompt
        
        logger.info(f"Generating code for prompt: {prompt}")
        system_msg = SystemMessage(content=f""" You are a senior software developer and highly skilled python coder. Your task is to create a python code whose output answers the user's prompt. The code must be concise, specific, and optimized for execution. Do NOT insert dates or years unless explicitly asked by the user. The data you have to analyze is in the following file path: {file_path}. It's a xlsx file. Here is a sample of the data: {sample_data}. You can read the names of the columns and the first 5 rows. The prompt is: {prompt}. Answer ONLY with the python code. Do not add any explanations or additional text. Make sure to include all necessary imports. Use pandas for data manipulation and analysis. This command might have been called before, and may be being called right now because of an error. Here's the current error to handle: {error_msg}. """)
        user_msg = HumanMessage(content=f"Prompt: {prompt}")
        try:
            response = ""
            async for text in unified_astream(client, [system_msg, user_msg]):
                response += text
            return response.strip()
        except Exception as e:
            logger.error(f"Error generating code: {str(e)}")
            return prompt   





class SummaryRAG:
    def __init__(self, input_dir):
        self.input_folder = input_dir
        
        # 1. Setup Directories
        for folder in [SUMMARY_FOLDER, DB_PATH]:
            if not os.path.exists(folder):
                os.makedirs(folder)

        self.registry_path = os.path.join(SUMMARY_FOLDER, "summary_registry.json")

        # 2. Initialize ChromaDB
        self.chroma_client = chromadb.PersistentClient(path=str(DB_PATH))
        self.emb_fn = embedding_functions.DefaultEmbeddingFunction()
        
        self.collection = self.chroma_client.get_or_create_collection(
            name="chat_summaries_archive",
            embedding_function=self.emb_fn
        )

        # 3. Initialize LLM
        self.llm = get_llm_client("summarizer")
        
        # --- FIX: Disable System Prompt Override ---
        # We manually clear the wrapper's system prompt so it respects the 
        # specific prompt templates defined below.
        if hasattr(self.llm, 'system_prompt'):
            self.llm.system_prompt = None
        # -------------------------------------------
        
        self.system_instructions = """
You are an AI assistant that analyzes chat logs. 
You must output VALID JSON only. Do not output markdown blocks, backticks, or explanatory text.
Your JSON object must have exactly these two keys:
1. "topics": A list of strings (3-5 main topics).
2. "summary": A detailed string paragraph summarizing the conversation.
"""
        self.base_prompt = ChatPromptTemplate.from_messages([
            ("system", self.system_instructions),
            ("user", "ANALYZE THIS CHAT:\n{text_content}")
        ])

        self.correction_prompt = ChatPromptTemplate.from_messages([
            ("system", self.system_instructions),
            ("user", "Your previous JSON was invalid or missing keys. Fix it.\n\nERROR: {error}\n\nBAD OUTPUT: {bad_output}\n\nSOURCE CHAT (Context): {text_content}\n\nProvide corrected JSON only.")
        ])

    def _load_registry(self):
        if os.path.exists(self.registry_path):
            with open(self.registry_path, 'r', encoding='utf-8') as f:
                try: return json.load(f)
                except json.JSONDecodeError: return {}
        return {}

    def _save_registry(self, data):
        with open(self.registry_path, 'w', encoding='utf-8') as f:
            json.dump(data, f, indent=4, ensure_ascii=False)

    def _extract_chat_data(self, filepath):
        try:
            with open(filepath, 'r', encoding='utf-8') as f:
                data = json.load(f)
            
            conversation = []
            if isinstance(data, dict) and "messages" in data:
                data = data["messages"]
            if not isinstance(data, list):
                return "", []

            unique_dates = set()
            for entry in data:
                role = entry.get("role", "unknown")
                content = entry.get("content", "")
                if role != "system":
                    conversation.append(f"{role.upper()}: {content}")
                
                meta = entry.get("metadata", {})
                ts = meta.get("timestamp")
                if ts:
                    try:
                        dt = datetime.fromisoformat(ts)
                        unique_dates.add(dt.strftime("%d-%m-%y"))
                    except: pass
            
            return "\n".join(conversation), sorted(list(unique_dates))
        except Exception:
            return "", []

    def _clean_and_parse_json(self, raw_text):
        """Robustly find and parse JSON from the LLM response."""
        try:
            # 1. Try cleaning markdown code blocks
            text = re.sub(r'```json\s*', '', raw_text, flags=re.IGNORECASE)
            text = re.sub(r'```', '', text).strip()
            return self._validate_json(json.loads(text))
        except json.JSONDecodeError:
            # 2. Try regex extraction of the JSON object
            match = re.search(r"\{.*\}", raw_text, re.DOTALL)
            if match:
                try:
                    return self._validate_json(json.loads(match.group(0)))
                except json.JSONDecodeError:
                    pass
            raise ValueError(f"Could not parse JSON. Raw: {raw_text[:50]}...")

    def _validate_json(self, data):
        if "topics" not in data or "summary" not in data:
            raise ValueError(f"Missing keys 'topics' or 'summary'. Found: {list(data.keys())}")
        return data

    def _generate_summary_with_retry(self, chat_text, log_callback):
        is_hf = isinstance(self.llm, HuggingFacePipeline)
        bad_output = ""
        error_msg = ""
        truncated_text = chat_text[:3000]

        for attempt in range(1, 6):
            raw_content = ""
            try:
                if attempt == 1:
                    if is_hf:
                        prompt = self.system_instructions + "\nANALYZE THIS CHAT:\n" + truncated_text
                        raw_content = self.llm.invoke(prompt)
                    else:
                        messages = self.base_prompt.format_messages(text_content=truncated_text)
                        response = self.llm.invoke(messages)
                        
                        if isinstance(response, str): raw_content = response
                        elif hasattr(response, 'content'): raw_content = response.content
                        else: raw_content = str(response)
                else:
                    log_callback(f"      [Logic] AI Fixing error: {error_msg}")
                    messages = self.correction_prompt.format_messages(
                        text_content=truncated_text, 
                        bad_output=bad_output,
                        error=error_msg
                    )
                    response = self.llm.invoke(messages)
                    
                    if isinstance(response, str): raw_content = response
                    elif hasattr(response, 'content'): raw_content = response.content
                    else: raw_content = str(response)

                return self._clean_and_parse_json(raw_content)

            except Exception as e:
                error_msg = str(e)
                bad_output = raw_content if raw_content else "Empty"
                if attempt == 5: return None

    def _process_file(self, file_path, registry, log_callback):
        filename = os.path.basename(file_path)
        log_callback(f"Processing: {filename}...")

        chat_text, date_list = self._extract_chat_data(file_path)
        if not chat_text: 
            return False

        result = self._generate_summary_with_retry(chat_text, log_callback)
        if not result:
            log_callback("   -> AI failed to generate summary.")
            return False
            
        try:
            summary_text = result['summary']
            topics = result['topics']
            summary_filename = f"summary_{filename}"
            output_path = os.path.join(SUMMARY_FOLDER, summary_filename)
            
            with open(output_path, 'w', encoding='utf-8') as f:
                json.dump({
                    "original_file": filename,
                    "topics": topics,
                    "summary": summary_text,
                    "dates": date_list
                }, f, indent=4, ensure_ascii=False)

            registry[filename] = {
                "original_file": filename,
                "summary_file": summary_filename,
                "dates": date_list
            }
            self._save_registry(registry)

            self.collection.upsert(
                documents=[summary_text],
                metadatas=[{"filename": filename, "topics": ", ".join(topics), "dates": ", ".join(date_list)}],
                ids=[summary_filename]
            )
            log_callback(f"   -> Success. Saved.")
            return True
        except Exception as e:
            log_callback(f"   -> CRITICAL SAVE ERROR: {e}")
            return False

    def process_and_index(self, log_callback):
        log_callback("--- Starting Full Processing Pipeline ---")
        registry = self._load_registry()
        files = glob.glob(os.path.join(self.input_folder, "*.json"))
        
        if not files:
            log_callback("No files found.")
            return

        for i, file_path in enumerate(files):
            log_callback(f"[{i+1}/{len(files)}]")
            self._process_file(file_path, registry, log_callback)
        log_callback("--- Done ---")

    def process_missing_files(self, log_callback):
        log_callback("--- Checking for Missing Files ---")
        registry = self._load_registry()
        all_files = glob.glob(os.path.join(self.input_folder, "*.json"))
        missing_files = [f for f in all_files if os.path.basename(f) not in registry]
        
        if not missing_files:
            log_callback("All files processed.")
            return

        for i, file_path in enumerate(missing_files):
            log_callback(f"[{i+1}/{len(missing_files)}]")
            self._process_file(file_path, registry, log_callback)
        log_callback("--- Done ---")

    def search_summaries(self, query, log_callback):
        log_callback(f"Searching: '{query}'")
        results = self.collection.query(query_texts=[query], n_results=5)
        
        if not results['documents']:
            log_callback("No matches.")
            return
            
        docs = results['documents'][0]
        distances = results['distances'][0]
        metas = results['metadatas'][0]
        
        combined = sorted(zip(docs, distances, metas), key=lambda x: x[1])
        found = False
        
        for doc, dist, meta in combined:
            if dist < RELEVANCE_THRESHOLD:
                log_callback(f"\n--- File: {meta['filename']} (Dist: {dist:.4f}) ---")
                log_callback(f"Dates: {meta['dates']}")
                log_callback(f"Summary: {doc}\n")
                found = True
        
        if not found:
            log_callback("No results within threshold.")

















def restore_memory():
        """Opens the Memory Handler Panel with multiple options."""
        root=ROOT
        memory_panel_win = tk.Toplevel(root)
        memory_panel_win.geometry("1200x800") # Increased height for 3 rows
        memory_panel_win.title("Memory Handler")
        memory_panel_win.configure(bg="#1e1e1e")

        # --- Helper: Create Log Window ---
        def create_log_window(title):
            win = tk.Toplevel(memory_panel_win)
            win.title(title)
            win.geometry("1300x1000")
            win.configure(bg="#1e1e1e")
            
            lbl = tk.Label(win, text="Initializing...", bg="#1e1e1e", fg="white", font=("Segoe UI", 10))
            lbl.pack(pady=(10, 5))
            
            log_frame = tk.Frame(win, bg="#1e1e1e")
            log_frame.pack(padx=10, pady=5, fill=tk.BOTH, expand=True)

            scrollbar = tk.Scrollbar(log_frame)
            scrollbar.pack(side=tk.RIGHT, fill=tk.Y)

            log_text = tk.Text(
                log_frame, 
                height=15, 
                width=70, 
                bg="#2d2d2d", 
                fg="#aaaaaa", 
                font=("Consolas", 8),
                yscrollcommand=scrollbar.set 
            )
            log_text.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
            scrollbar.config(command=log_text.yview)
            
            return win, lbl, log_text

        # --- Button Functions ---

        def reset_and_reindex_ui():
            """Deletes old embeddings and re-indexes all chats with the new model."""
            
            # 1. Safety Confirmation
            if not messagebox.askyesno("Confirm Reset", 
                                       "WARNING: This will delete all existing semantic embeddings (FAISS index).\n\n"
                                       "It will then regenerate them using the new Multilingual Model.\n"
                                       "This process may take time depending on your chat history size.\n\n"
                                       "Continue?"):
                return

            progress_win, lbl, log_text = create_log_window("Resetting & Re-indexing Embeddings")
            
            def log_callback(msg):
                try:
                    progress_win.after(0, lambda: lbl.config(text=msg[:60]))
                    def append():
                        log_text.insert(tk.END, str(msg) + "\n")
                        log_text.see(tk.END)
                    progress_win.after(0, append)
                except Exception: pass

            def run_task():
                try:
                    # A. DELETE OLD FILES
                    log_callback("--- STEP 1: DELETING OLD INDICES ---")
                    files_to_delete = [
                        "global_index.faiss", 
                        "global_metadata.pkl", 
                        "embedded_turns.json", 
                        "embeddings_map.json"
                    ]
                    
                    for fname in files_to_delete:
                        fpath = os.path.join(EMBEDDINGS_DIR, fname)
                        if os.path.exists(fpath):
                            try:
                                os.remove(fpath)
                                log_callback(f"Deleted: {fname}")
                            except Exception as e:
                                log_callback(f"Error deleting {fname}: {e}")
                        else:
                            log_callback(f"File not found (already clean): {fname}")

                    # B. RE-INITIALIZE MANAGER
                    log_callback("\n--- STEP 2: INITIALIZING NEW MODEL ---")
                    # This automatically creates fresh, empty index files because we deleted the old ones
                    # Ensure your EmbeddingManager class is updated to the new model name before running this!
                    new_manager = EmbeddingManager(model_name='sentence-transformers/paraphrase-multilingual-MiniLM-L12-v2')
                    log_callback("New Embedding Manager initialized.")

                    # C. RE-INDEX ALL CHATS
                    log_callback("\n--- STEP 3: RE-INDEXING CHATS ---")
                    chat_files = glob.glob(os.path.join(OUTPUT_DIR, "*.json"))
                    log_callback(f"Found {len(chat_files)} chat files.")

                    total_turns = 0
                    
                    for i, file_path in enumerate(chat_files):
                        filename = os.path.basename(file_path)
                        try:
                            with open(file_path, 'r', encoding='utf-8') as f:
                                data = json.load(f)
                            
                            # Normalize data structure (handle if it's a list or a dict with "messages")
                            messages = data if isinstance(data, list) else data.get("messages", [])
                            
                            turns_in_file = 0
                            idx = 0
                            while idx < len(messages) - 1:
                                curr = messages[idx]
                                nxt = messages[idx+1]
                                
                                # Find User -> Assistant pairs
                                if curr.get('role') == 'user' and nxt.get('role') == 'assistant':
                                    user_text = curr.get('content', '')
                                    ai_text = nxt.get('content', '')
                                    
                                    # Extract or generate timestamps
                                    u_meta = curr.get('metadata', {})
                                    a_meta = nxt.get('metadata', {})
                                    
                                    # Handle case where metadata might be None
                                    if u_meta is None: u_meta = {}
                                    if a_meta is None: a_meta = {}

                                    u_ts = u_meta.get('timestamp')
                                    a_ts = a_meta.get('timestamp')

                                    if user_text and ai_text:
                                        new_manager.embed_conversation_turn(
                                            user_msg=user_text,
                                            ai_msg=ai_text,
                                            chat_file=file_path,
                                            user_timestamp=u_ts,
                                            ai_timestamp=a_ts
                                        )
                                        turns_in_file += 1
                                    idx += 2
                                else:
                                    idx += 1
                            
                            total_turns += turns_in_file
                            if i % 5 == 0:
                                log_callback(f"[{i}/{len(chat_files)}] Processed {filename} (+{turns_in_file} turns)")

                        except Exception as e:
                            log_callback(f"Error processing {filename}: {e}")

                    # D. SAVE
                    new_manager._save_global_index()
                    new_manager._save_embedded_turns()
                    
                    log_callback("\n" + "="*30)
                    log_callback(f"COMPLETED. Total turns embedded: {total_turns}")
                    log_callback("Please Restart M.A.I.A to fully load the new index.")
                    log_callback("="*30)

                except Exception as e:
                    log_callback(f"CRITICAL ERROR: {e}")
                    import traceback
                    log_callback(traceback.format_exc())

            threading.Thread(target=run_task, daemon=True).start()



        def elaborate_daily_episodes():
            progress_win, lbl, log_text = create_log_window("Elaborating Daily Episodes")
            def log_callback(msg):
                try:
                    progress_win.after(0, lambda: lbl.config(text=msg if len(msg) < 60 else "Processing..."))
                    def append():
                        log_text.insert(tk.END, msg + "\n")
                        log_text.see(tk.END)
                    progress_win.after(0, append)
                except Exception: pass

            def run_task():
                try:
                    manager = EpisodicMemoryManager()
                    manager.analyze_and_store(log_callback, chats='today')
                    log_callback("--- DONE ---")
                except Exception as e: log_callback(f"ERROR: {e}")
            threading.Thread(target=run_task, daemon=True).start()

        def elaborate_all_episodes():
            progress_win, lbl, log_text = create_log_window("Restoring All Memory")
            def log_callback(msg):
                try:
                    progress_win.after(0, lambda: lbl.config(text=msg if len(msg) < 60 else "Processing..."))
                    def append():
                        log_text.insert(tk.END, msg + "\n")
                        log_text.see(tk.END)
                    progress_win.after(0, append)
                except Exception: pass

            def run_task():
                try:
                    manager = EpisodicMemoryManager()
                    manager.analyze_and_store(log_callback, chats='all')
                    log_callback("--- DONE ---")
                except Exception as e: log_callback(f"ERROR: {e}")
            threading.Thread(target=run_task, daemon=True).start()

        def consolidate_episodes():
            progress_win, lbl, log_text = create_log_window("Consolidating Episodes")
            def log_callback(msg):
                try:
                    progress_win.after(0, lambda: lbl.config(text=msg if len(msg) < 60 else "Processing..."))
                    def append():
                        log_text.insert(tk.END, msg + "\n")
                        log_text.see(tk.END)
                    progress_win.after(0, append)
                except Exception: pass

            def run_task():
                try:
                    manager = EpisodicMemoryManager()
                    manager.consolidate_memories(log_callback)
                    log_callback("--- DONE ---")
                except Exception as e: log_callback(f"ERROR: {e}")
            threading.Thread(target=run_task, daemon=True).start()

        # --- NEW SUMMARY FUNCTIONS ---

        def summarize_all_chats():
            """Runs SummaryRAG process_and_index on all files."""
            progress_win, lbl, log_text = create_log_window("Summarizing ALL Chats")
            
            def log_callback(msg):
                try:
                    progress_win.after(0, lambda: lbl.config(text=msg[:60]))
                    def append():
                        log_text.insert(tk.END, str(msg) + "\n")
                        log_text.see(tk.END)
                    progress_win.after(0, append)
                except Exception: pass

            def run_task():
                try:
                    rag = SummaryRAG(input_dir=OUTPUT_DIR)
                    rag.process_and_index(log_callback)
                    log_callback("--- SUMMARY COMPLETE ---")
                except Exception as e:
                    log_callback(f"CRITICAL ERROR: {e}")
            
            threading.Thread(target=run_task, daemon=True).start()

        def summarize_missing_chats():
            """Runs SummaryRAG process_missing_files."""
            progress_win, lbl, log_text = create_log_window("Summarizing Missing Chats")
            
            def log_callback(msg):
                try:
                    progress_win.after(0, lambda: lbl.config(text=msg[:60]))
                    def append():
                        log_text.insert(tk.END, str(msg) + "\n")
                        log_text.see(tk.END)
                    progress_win.after(0, append)
                except Exception: pass

            def run_task():
                try:
                    rag = SummaryRAG(input_dir=OUTPUT_DIR)
                    rag.process_missing_files(log_callback)
                    log_callback("--- SUMMARY COMPLETE ---")
                except Exception as e:
                    log_callback(f"CRITICAL ERROR: {e}")
            
            threading.Thread(target=run_task, daemon=True).start()

        def search_summaries_ui():
            """Asks for query then searches RAG."""
            query = simpledialog.askstring("Search Summaries", "Enter topic or keyword:")
            if not query:
                return

            progress_win, lbl, log_text = create_log_window(f"Search Results: {query}")
            
            def log_callback(msg):
                try:
                    # Don't update label for every line of text result, just status
                    if "---" in msg:
                        progress_win.after(0, lambda: lbl.config(text="Found match..."))
                    def append():
                        log_text.insert(tk.END, str(msg) + "\n")
                        log_text.see(tk.END)
                    progress_win.after(0, append)
                except Exception: pass

            def run_task():
                try:
                    log_callback("Initializing Database...")
                    rag = SummaryRAG(input_dir=OUTPUT_DIR)
                    rag.search_summaries(query, log_callback)
                    log_callback("--- SEARCH FINISHED ---")
                except Exception as e:
                    log_callback(f"CRITICAL ERROR: {e}")
            
            threading.Thread(target=run_task, daemon=True).start()

        # --- Grid Layout for Memory Panel ---
        
        btn_style = {"font": ("Segoe UI", 10), "padx": 10, "pady": 15, "fg": "white", "width": 25, "relief": "flat"}
        
        # Column 1 Buttons
        tk.Button(memory_panel_win, text="Elaborate Daily Episodes", command=elaborate_daily_episodes, bg="#2980b9", **btn_style).grid(row=0, column=0, padx=20, pady=10)
        tk.Button(memory_panel_win, text="Elaborate All Episodes", command=elaborate_all_episodes, bg="#d35400", **btn_style).grid(row=1, column=0, padx=20, pady=10)
        tk.Button(memory_panel_win, text="Consolidate Episodes", command=consolidate_episodes, bg="#8e44ad", **btn_style).grid(row=2, column=0, padx=20, pady=10)

        # Column 2 Buttons
        tk.Button(memory_panel_win, text="Summarize All Chats", command=summarize_all_chats, bg="#27ae60", **btn_style).grid(row=0, column=1, padx=20, pady=10)
        tk.Button(memory_panel_win, text="Summarize Missing Chats", command=summarize_missing_chats, bg="#16a085", **btn_style).grid(row=1, column=1, padx=20, pady=10)
        # New 6th Button
        tk.Button(memory_panel_win, text="Search Summaries", command=search_summaries_ui, bg="#f39c12", **btn_style).grid(row=2, column=1, padx=20, pady=10)
        tk.Button(memory_panel_win, text="Search Summaries", command=search_summaries_ui, bg="#f39c12", **btn_style).grid(row=2, column=1, padx=20, pady=10)

        # --- 7TH BUTTON: RESET & RE-INDEX ---
        tk.Button(
            memory_panel_win, 
            text="\u26A0\uFE0F RESET & RE-INDEX ALL EMBEDDINGS \u26A0\uFE0F", 
            command=reset_and_reindex_ui, 
            bg="#c0392b", # Red color for caution
            font=("Segoe UI", 10, "bold"),
            padx=10, 
            pady=15, 
            fg="white", 
            relief="flat",
            width=55 # Wider to span
        ).grid(row=3, column=0, columnspan=2, padx=20, pady=(20, 10)) # Row 3, Span 2 cols
        
        memory_panel_win.grid_columnconfigure(0, weight=1)
        memory_panel_win.grid_columnconfigure(1, weight=1)


# ======================
# MAIN FUNCTION
# ======================

def open_other_settings_ui(parent_win):
    """UI for configuring Retrieval, Search, and Memory parameters."""
    win = tk.Toplevel(parent_win)
    win.title("Other Settings (Retrieval & Search)")
    win.geometry("1200x1800") # Increased size slightly
    win.configure(bg="#2d2d2d")

    # Load defaults
    current_conf = LLM_CONFIG.get("retrieval", DEFAULT_CONFIG["retrieval"])
    
    # --- Variables ---
    # RAG Counts
    rag_var = tk.StringVar(value=str(current_conf.get("rag_k", 3)))
    episodic_var = tk.StringVar(value=str(current_conf.get("episodic_k", 12)))
    summary_var = tk.StringVar(value=str(current_conf.get("summary_k", 3)))
    
    # --- NEW VARIABLES FOR DOCUMENT RAG ---
    doc_k_var = tk.StringVar(value=str(current_conf.get("doc_rag_k", 5)))
    doc_thresh_var = tk.StringVar(value=str(current_conf.get("doc_rag_threshold", 2.0)))
    # --------------------------------------

    # Search
    google_var = tk.StringVar(value=str(current_conf.get("google_k", 5)))
    
    # Thresholds & Limits
    rel_thresh_var = tk.StringVar(value=str(current_conf.get("relevance_threshold", 2.0)))
    cons_thresh_var = tk.StringVar(value=str(current_conf.get("consolidation_threshold", 0.8)))
    history_var = tk.StringVar(value=str(current_conf.get("max_history", 1000)))
    topic_trunc_var = tk.StringVar(value=str(current_conf.get("topic_truncation", 100)))
    sem_trunc_var = tk.StringVar(value=str(current_conf.get("semantic_truncation", 600)))

    # --- UI Generators ---
    def add_section(text):
        tk.Label(win, text=text, bg="#2d2d2d", fg="#4ea1ff", font=("Segoe UI", 11, "bold")).pack(pady=(20, 5))

    def add_setting_row(label_text, var, tooltip=""):
        frame = tk.Frame(win, bg="#2d2d2d")
        frame.pack(fill="x", padx=30, pady=5)
        lbl = tk.Label(frame, text=label_text, bg="#2d2d2d", fg="white", width=40, anchor="w")
        lbl.pack(side="left")
        tk.Entry(frame, textvariable=var, bg="#3a3a3a", fg="white", insertbackground="white", width=10).pack(side="right")

    # --- Layout ---
    add_section("Context Window (Items to Retrieve)")
    add_setting_row("Semantic Context (From Past Chats):", rag_var)
    add_setting_row("Episodic Memories (Facts):", episodic_var)
    add_setting_row("Chat Summaries:", summary_var)

    # --- NEW SECTION IN UI ---
    add_section("Document RAG Settings")
    add_setting_row("Doc Chunks per File (K):", doc_k_var)
    add_setting_row("Doc Distance Threshold (Lower=Stricter):", doc_thresh_var)
    # -------------------------

    add_section("Context Truncation (Max Characters)")
    add_setting_row("Topic/Metadata Context Length:", topic_trunc_var)
    add_setting_row("Main Semantic Context Length:", sem_trunc_var)

    add_section("Google Search")
    add_setting_row("Max Links to Scrape (num_links):", google_var)

    add_section("Memory & Thresholds")
    add_setting_row("Relevance Distance (Lower = Stricter):", rel_thresh_var)
    add_setting_row("Consolidation Distance (Lower = Stricter):", cons_thresh_var)
    add_setting_row("Max Active Chat History (Messages):", history_var)

    def save_settings():
        try:
            # Update global config
            new_conf = {
                "rag_k": int(rag_var.get()),
                "episodic_k": int(episodic_var.get()),
                "summary_k": int(summary_var.get()),
                "google_k": int(google_var.get()),
                "relevance_threshold": float(rel_thresh_var.get()),
                "consolidation_threshold": float(cons_thresh_var.get()),
                "max_history": int(history_var.get()),
                "topic_truncation": int(topic_trunc_var.get()),
                "semantic_truncation": int(sem_trunc_var.get()),
                
                # --- SAVE NEW VALUES ---
                "doc_rag_k": int(doc_k_var.get()),
                "doc_rag_threshold": float(doc_thresh_var.get())
            }
            LLM_CONFIG["retrieval"] = new_conf
            
            # Update Global Constants immediately
            global MAX_MEMORY_SIZE, RELEVANCE_THRESHOLD
            MAX_MEMORY_SIZE = new_conf["max_history"]
            RELEVANCE_THRESHOLD = new_conf["relevance_threshold"]
            
            save_config(LLM_CONFIG)
            messagebox.showinfo("Success", "Settings saved successfully.")
            win.destroy()
        except ValueError:
            messagebox.showerror("Error", "Please enter valid numbers.")

    tk.Button(win, text="Save Settings", command=save_settings, bg="#2ecc71", fg="white", font=("Segoe UI", 10, "bold"), padx=20, pady=10).pack(pady=30)

def open_directory_settings_ui(parent_win):
    """UI for configuring Local LLM Directories (HF & Ollama)."""
    win = tk.Toplevel(parent_win)
    win.title("Define Local Model Directories")
    win.geometry("1400x700")
    win.configure(bg="#2d2d2d")

    # Load current values
    current_dirs = LLM_CONFIG.get("directories", DEFAULT_CONFIG["directories"])
    
    hf_var = tk.StringVar(value=current_dirs.get("hf_home", ""))
    ollama_var = tk.StringVar(value=current_dirs.get("ollama_models", ""))

    def browse_dir(var):
        path = filedialog.askdirectory()
        if path:
            var.set(path)

    def add_row(label_text, var, tooltip_text):
        frame = tk.Frame(win, bg="#2d2d2d")
        frame.pack(fill="x", padx=20, pady=10)
        
        lbl = tk.Label(frame, text=label_text, bg="#2d2d2d", fg="#4ea1ff", font=("Segoe UI", 10, "bold"), width=20, anchor="w")
        lbl.pack(side="left")
        
        entry = tk.Entry(frame, textvariable=var, bg="#3a3a3a", fg="white", insertbackground="white")
        entry.pack(side="left", fill="x", expand=True, padx=5)
        
        btn = tk.Button(frame, text="Browse", command=lambda: browse_dir(var), bg="#555", fg="white")
        btn.pack(side="right")
        
        # Tooltip/Info line
        info_lbl = tk.Label(win, text=tooltip_text, bg="#2d2d2d", fg="#aaaaaa", font=("Segoe UI", 8, "italic"), anchor="w")
        info_lbl.pack(fill="x", padx=25, pady=(0, 10))

    tk.Label(win, text="Local Model Storage Paths", bg="#2d2d2d", fg="white", font=("Segoe UI", 12, "bold")).pack(pady=15)

    # HuggingFace Row
    add_row("HuggingFace Home:", hf_var, 
            "Destination for downloaded Transformers/HF models. (Sets HF_HOME)")

    # Ollama Row
    add_row("Ollama Models:", ollama_var, 
            "Destination for Ollama blobs/manifests. (Sets OLLAMA_MODELS). Restart Ollama service after changing.")

    def save_dirs():
        new_dirs = {
            "hf_home": hf_var.get().strip(),
            "ollama_models": ollama_var.get().strip()
        }
        
        # Save to config
        LLM_CONFIG["directories"] = new_dirs
        save_config(LLM_CONFIG)
        
        # Apply immediately to current session env vars
        if new_dirs["hf_home"]: 
            os.environ["HF_HOME"] = new_dirs["hf_home"]
            os.environ["TRANSFORMERS_CACHE"] = new_dirs["hf_home"]
            
        if new_dirs["ollama_models"]:
            os.environ["OLLAMA_MODELS"] = new_dirs["ollama_models"]
            
        messagebox.showinfo("Saved", "Directories saved.\n\nFor HuggingFace: Applied immediately.\nFor Ollama: You may need to restart the Ollama service for this to take effect.")
        win.destroy()

    btn_frame = tk.Frame(win, bg="#2d2d2d")
    btn_frame.pack(fill="x", pady=20)
    
    tk.Button(btn_frame, text="Save & Apply", command=save_dirs, bg="#2ecc71", fg="white", font=("Segoe UI", 10, "bold"), width=20).pack()

def open_custom_llm_ui(root_window):
    """Configuration UI."""
    config_win = tk.Toplevel(root_window)
    config_win.title("LLM Configuration")
    config_win.geometry("1600x1800") 
    config_win.configure(bg="#1e1e1e")

    tk.Label(config_win, text="LLM Settings", bg="#1e1e1e", fg="white", font=("Segoe UI", 12, "bold")).pack(pady=10)

    # ============================================================
    # 1. SETUP FOOTER (Buttons)
    # ============================================================
    btn_frame = tk.Frame(config_win, bg="#1e1e1e")
    btn_frame.pack(side="bottom", fill="x", pady=10) 

    # --- MODIFIED: SELECTIVE UNLOAD BUTTON ---
    tk.Button(btn_frame, text="\u26A0\uFE0F Manage Loaded Models (Selective Unload) \u26A0\uFE0F", 
              command=lambda: open_unload_selector(config_win),
              bg="#c0392b", fg="white", font=("Segoe UI", 10, "bold"), 
              relief="flat", pady=5).pack(side="top", fill="x", padx=20, pady=5)
    # -----------------------------------------

    # Bottom Button: Other Settings
    tk.Button(btn_frame, text="Other Settings (Retrieval & Search)", 
              command=lambda: open_other_settings_ui(config_win),
              bg="#8e44ad", fg="white", font=("Segoe UI", 10, "bold")).pack(side="bottom", fill="x", padx=20, pady=(2, 0))

    # Top Button (in footer): Directories
    tk.Button(btn_frame, text="Define Directories (Download Paths)", 
            command=lambda: open_directory_settings_ui(config_win),
            bg="#2980b9", fg="white", font=("Segoe UI", 10, "bold")).pack(side="bottom", fill="x", padx=20, pady=(5, 2))

    # ============================================================
    # 2. SETUP SCROLLABLE AREA
    # ============================================================
    canvas = tk.Canvas(config_win, bg="#1e1e1e", highlightthickness=0)
    scrollbar = ttk.Scrollbar(config_win, orient="vertical", command=canvas.yview)
    scrollable_frame = tk.Frame(canvas, bg="#1e1e1e")

    scrollable_frame.bind("<Configure>", lambda e: canvas.configure(scrollregion=canvas.bbox("all")))
    
    frame_id = canvas.create_window((0, 0), window=scrollable_frame, anchor="nw")
    
    def resize_frame(event):
        canvas.itemconfig(frame_id, width=event.width)
    canvas.bind("<Configure>", resize_frame)

    canvas.configure(yscrollcommand=scrollbar.set)

    scrollbar.pack(side="right", fill="y") 
    canvas.pack(side="left", fill="both", expand=True)

    # ============================================================
    # 3. POPULATE LIST
    # ============================================================
    components = ["main", "router", "refiner", "coder", "summarizer", "memory_analyzer", "memory_consolidator"]
    
    for comp in components:
        frame = tk.Frame(scrollable_frame, bg="#1e1e1e", pady=5, padx=10)
        frame.pack(fill=tk.X)
        
        lbl_text = comp.replace("_", " ").title()
        tk.Label(frame, text=lbl_text, bg="#1e1e1e", fg="#aaaaaa", width=30, anchor="w").pack(side=tk.LEFT)
        
        curr = LLM_CONFIG.get(comp, DEFAULT_CONFIG.get("main"))
        status = f"{curr.get('provider', 'ollama')} : {curr.get('model', 'default')}"
        
        tk.Button(frame, text=status, command=lambda c=comp: configure_component(c, config_win),
                  bg="#34495e", fg="white", width=30).pack(side=tk.RIGHT)
    




import tkinter as tk
from tkinter import ttk

def configure_component(component, parent_win):
    """Detail config window."""
    win = tk.Toplevel(parent_win)
    win.title(f"Config: {component}")
    win.geometry("1600x1800")
    win.configure(bg="#2d2d2d")
    
    # Load defaults
    default_vals = DEFAULT_CONFIG.get(component, DEFAULT_CONFIG["main"])
    current_conf = LLM_CONFIG.get(component, default_vals)
    
    # --- Variables ---
    provider_var = tk.StringVar(value=current_conf.get("provider", "ollama"))
    model_var = tk.StringVar(value=current_conf.get("model", DEFAULT_MODEL))
    apikey_var = tk.StringVar(value=current_conf.get("api_key", ""))
    endpoint_var = tk.StringVar(value=current_conf.get("endpoint", ""))
    apiversion_var = tk.StringVar(value=current_conf.get("api_version", ""))
    
    temp_var = tk.StringVar(value=str(current_conf.get("temperature", 0.7)))
    topp_var = tk.StringVar(value=str(current_conf.get("top_p", 0.9)))
    tokens_var = tk.StringVar(value=str(current_conf.get("max_tokens", 1024)))

    # --- UI Layout (Canvas + Scrollbar) ---
    main_canvas = tk.Canvas(win, bg="#2d2d2d", highlightthickness=0)
    scrollbar = ttk.Scrollbar(win, orient="vertical", command=main_canvas.yview)
    
    # The Scrollable Frame
    scrollable_frame = tk.Frame(main_canvas, bg="#2d2d2d")

    # 1. Update scrollregion when content changes size
    scrollable_frame.bind(
        "<Configure>",
        lambda e: main_canvas.configure(scrollregion=main_canvas.bbox("all"))
    )

    # 2. CREATE THE WINDOW AND STORE THE ID
    frame_id = main_canvas.create_window((0, 0), window=scrollable_frame, anchor="nw")
    
    # 3. CRITICAL FIX: Force inner frame to match canvas width
    def resize_frame(event):
        main_canvas.itemconfig(frame_id, width=event.width)
    
    main_canvas.bind("<Configure>", resize_frame)

    main_canvas.configure(yscrollcommand=scrollbar.set)

    # Pack Main Layout
    scrollbar.pack(side="right", fill="y")
    main_canvas.pack(side="left", fill="both", expand=True)

    # --- Content Inside Scrollable Frame ---
    
    # Use a container frame with padding
    content_padding = tk.Frame(scrollable_frame, bg="#2d2d2d")
    content_padding.pack(fill="both", expand=True, padx=20, pady=10)

    # Provider
    # CHANGE: Parent set to 'content_padding' instead of 'scrollable_frame'
    tk.Label(content_padding, text="Provider:", bg="#2d2d2d", fg="white", font=("Segoe UI", 10, "bold")).pack(pady=(10,5), anchor="w")
    
    # CHANGE: Parent set to 'content_padding' instead of 'scrollable_frame'
    p_frame = tk.Frame(content_padding, bg="#2d2d2d")
    p_frame.pack(anchor="w")
    
    providers = ["ollama", "huggingface", "openai", "anthropic", "mistral", "gemini", "azure"] 
    
    for p in providers:
        tk.Radiobutton(p_frame, text=p.title(), variable=provider_var, value=p, bg="#2d2d2d", fg="white", selectcolor="#2d2d2d").pack(side="left", padx=5)

    # --- UI Layout for Model, API Key, AND ENDPOINT ---
    tk.Label(content_padding, text="Model Name:", bg="#2d2d2d", fg="white", font=("Segoe UI", 10, "bold")).pack(pady=(15,5), anchor="w")
    tk.Entry(content_padding, textvariable=model_var, bg="#3a3a3a", fg="white", insertbackground="white").pack(anchor="w", fill="x")
    
    # New Endpoint Field
    tk.Label(content_padding, text="Endpoint URL (e.g. https://api.ollama.com or http://localhost:11434):", bg="#2d2d2d", fg="white", font=("Segoe UI", 10, "bold")).pack(pady=(15,5), anchor="w")
    tk.Entry(content_padding, textvariable=endpoint_var, bg="#3a3a3a", fg="white", insertbackground="white").pack(anchor="w", fill="x")

    tk.Label(content_padding, text="API Key (Ollama Cloud / OpenAI / etc.):", bg="#2d2d2d", fg="white", font=("Segoe UI", 10, "bold")).pack(pady=(15,5), anchor="w")
    tk.Entry(content_padding, textvariable=apikey_var, show="*", bg="#3a3a3a", fg="white", insertbackground="white").pack(anchor="w", fill="x")
    
    tk.Label(content_padding, text="API Version (Azure only, e.g., 2024-02-15-preview):", bg="#2d2d2d", fg="white", font=("Segoe UI", 10, "bold")).pack(pady=(15,5), anchor="w")
    tk.Entry(content_padding, textvariable=apiversion_var, bg="#3a3a3a", fg="white", insertbackground="white").pack(anchor="w", fill="x")

    # Numeric Settings
    settings_frame = tk.Frame(content_padding, bg="#2d2d2d")
    settings_frame.pack(pady=20, fill="x")

    settings_frame.columnconfigure(1, weight=1)
    settings_frame.columnconfigure(3, weight=1)
    settings_frame.columnconfigure(5, weight=1)

    tk.Label(settings_frame, text="Temperature:", bg="#2d2d2d", fg="white").grid(row=0, column=0, sticky="e", padx=5)
    tk.Entry(settings_frame, textvariable=temp_var, width=10, bg="#3a3a3a", fg="white", insertbackground="white").grid(row=0, column=1, sticky="w")

    tk.Label(settings_frame, text="Top P:", bg="#2d2d2d", fg="white").grid(row=0, column=2, sticky="e", padx=5)
    tk.Entry(settings_frame, textvariable=topp_var, width=10, bg="#3a3a3a", fg="white", insertbackground="white").grid(row=0, column=3, sticky="w")

    tk.Label(settings_frame, text="Max Tokens:", bg="#2d2d2d", fg="white").grid(row=0, column=4, sticky="e", padx=5)
    tk.Entry(settings_frame, textvariable=tokens_var, width=10, bg="#3a3a3a", fg="white", insertbackground="white").grid(row=0, column=5, sticky="w")

    # --- SYSTEM PROMPT SECTION ---
    tk.Label(content_padding, text="System Prompt (Instructions):", bg="#2d2d2d", fg="#4ea1ff", font=("Segoe UI", 10, "bold")).pack(pady=(20,5), anchor="w")
    
    text_container = tk.Frame(content_padding, bg="#3a3a3a")
    text_container.pack(fill="x", expand=True)

    text_scroll = ttk.Scrollbar(text_container, orient="vertical")
    
    prompt_text_widget = tk.Text(text_container, height=25, bg="#3a3a3a", fg="white", 
                                 font=("Consolas", 10), wrap=tk.WORD, 
                                 yscrollcommand=text_scroll.set,
                                 insertbackground="white")
    
    text_scroll.config(command=prompt_text_widget.yview)
    
    text_scroll.pack(side="right", fill="y")
    prompt_text_widget.pack(side="left", fill="both", expand=True)
    
    # Load existing system prompt
    existing_sys_prompt = current_conf.get("system_prompt", default_vals.get("system_prompt", ""))
    prompt_text_widget.insert("1.0", existing_sys_prompt)

    # --- SAVE LOGIC ---
    def save_and_reload():
        print("Saving configuration for", component)
        prov = provider_var.get()
        mod = model_var.get().strip()
        
        try:
            t_val = float(temp_var.get())
            p_val = float(topp_var.get())
            tok_val = int(tokens_var.get())
        except ValueError:
            messagebox.showerror("Input Error", "Numeric fields must be valid numbers.")
            return

        # Get Text from widget
        sys_prompt_val = prompt_text_widget.get("1.0", "end-1c").strip()

        new_conf = {
            "provider": prov,
            "model": mod,
            "api_key": apikey_var.get().strip(),
            "endpoint": endpoint_var.get().strip(),
            "api_version": apiversion_var.get().strip(),
            "temperature": t_val,
            "top_p": p_val,
            "max_tokens": tok_val,
            "system_prompt": sys_prompt_val # Save the prompt
        }
        
        if prov == "ollama":
            validate_ollama_model(mod)
        
        LLM_CONFIG[component] = new_conf
        save_config(LLM_CONFIG)
        
        win.destroy()
        parent_win.destroy()
        open_custom_llm_ui(parent_win.master)

    tk.Button(content_padding, text="Save Configuration", command=save_and_reload, bg="#2ecc71", fg="white", font=("Segoe UI", 11, "bold"), padx=20, pady=10).pack(pady=30)

    




def select_or_new_chat():
    """Dialog to select existing chat, create new one, restore memory, or configure LLMs."""
    root = TkinterDnD.Tk()
    root.title("M.A.I.A. Start")
    root.configure(bg="#1e1e1e")
    root.geometry("1300x600") # Increased height
    choice = {"file": None}
    global ROOT
    ROOT=root

    def select_chat():
        file_path = filedialog.askopenfilename(initialdir=OUTPUT_DIR, filetypes=[("JSON files", "*.json")])
        if file_path:
            choice["file"] = file_path
            root.destroy()

    def new_chat():
        title = simpledialog.askstring("New Chat", "Session Title:")
        if title:
            safe_title = re.sub(r"[^a-zA-Z0-9_-]", "", title)
            timestamp = datetime.now().strftime("%Y%m%d%H%M%S")
            file_path = os.path.join(OUTPUT_DIR, f"{safe_title}_{timestamp}.json")
            choice["file"] = file_path
            root.destroy()

    # --- UI LAYOUT ---
    tk.Label(root, text="Select an option to start M.A.I.A.", bg="#1e1e1e", fg="#ffffff", font=("Segoe UI", 14, "bold")).pack(pady=20)
    
    button_frame = tk.Frame(root, bg="#1e1e1e")
    button_frame.pack(pady=10)
    
    btn_style = {"font": ("Segoe UI", 10), "padx": 20, "pady": 10, "fg": "white", "width": 20}

    # Button 1: Existing Chat
    tk.Button(button_frame, text="Load Existing Chat", command=select_chat, bg="#3498db", **btn_style).grid(row=0, column=0, padx=10)
    
    # Button 2: New Chat
    tk.Button(button_frame, text="Start New Chat", command=new_chat, bg="#2ecc71", **btn_style).grid(row=0, column=1, padx=10)
    
    # NEW Button 3: Custom LLM
    tk.Button(root, text="Custom LLM Configuration", command=lambda: open_custom_llm_ui(root), bg="#8e44ad", font=("Segoe UI", 10), padx=20, pady=10, fg="white", width=30).pack(pady=10)

    # Button 4: Restore Memory
    tk.Button(root, text="Memory Handler", command=restore_memory, bg="#d35400", font=("Segoe UI", 10), padx=20, pady=10, fg="white", width=30).pack(pady=10)

    # Center window
    root.update_idletasks()
    width = root.winfo_width()
    height = root.winfo_height()
    x = (root.winfo_screenwidth() // 2) - (width // 2)
    y = (root.winfo_screenheight() // 2) - (height // 2)
    root.geometry(f"{width}x{height}+{x}+{y}")
    
    root.mainloop()
    return choice["file"]

def perform_shutdown_cleanup():
    """
    Clears RAM, VRAM, and unloads Ollama models before exit.
    """
    print("\n--- INITIATING SHUTDOWN CLEANUP ---")
    
    # 1. HUGGINGFACE & PYTORCH CLEANUP
    global LOADED_HF_PIPELINES
    if LOADED_HF_PIPELINES:
        print(f"Unloading {len(LOADED_HF_PIPELINES)} HF Pipelines...")
        LOADED_HF_PIPELINES.clear()
    
    import gc
    gc.collect()
    
    try:
        import torch
        if torch.cuda.is_available():
            torch.cuda.empty_cache()
            torch.cuda.ipc_collect()
            print("CUDA VRAM cache cleared.")
        elif torch.backends.mps.is_available():
            torch.mps.empty_cache()
            print("MPS (Mac) RAM cache cleared.")
    except ImportError:
        pass

    # 2. OLLAMA CLEANUP
    # We use the helper function we added earlier to find what's running
    try:
        active_ollama = get_ollama_loaded_models()
        if active_ollama:
            print(f"Unloading active Ollama models: {active_ollama}")
            for model in active_ollama:
                # We reuse your existing unload function
                unload_specific_model("ollama", model)
    except Exception as e:
        print(f"Error cleaning up Ollama: {e}")

    print("--- CLEANUP COMPLETE ---")



def main():
    try:
        load_config() # Ensure config is loaded
        memory_file = select_or_new_chat()
        if not memory_file:
            return
        
        episodic_manager = EpisodicMemoryManager()
        embedding_manager = EmbeddingManager()
        memory_manager = MemoryManager(memory_file, embedding_manager)
        prompthandler = PromptHandler(memory_manager, episodic_manager)

        # Agents don't need model args passed, they check config
        router = RouterAgent() 
        refiner = QueryRefinerAgent()
        coder = CodingAgent()
        
        root = TkinterDnD.Tk()
        app = CrewChatUI(root, router, refiner, coder, memory_manager, embedding_manager, prompthandler, episodic_manager)

        def on_closing():
            # Optional: Ask for confirmation
            if messagebox.askokcancel("Quit", "Do you want to save & quit?"):
                # 1. Update UI to show we are closing
                root.title("M.A.I.A. - Shutting down...")
                
                # 2. Save Data
                app.save_chat()
                
                # 3. Perform Memory Cleanup
                perform_shutdown_cleanup()
                
                # 4. Destroy Window
                root.destroy()
                
                # 5. Force Kill (Optional, ensures all threads stop)
                import sys
                sys.exit(0)
    
        root.protocol("WM_DELETE_WINDOW", on_closing)
        root.mainloop()
    except Exception as e:
        logger.error(f"Error in main: {str(e)}")
        # We can't use messagebox here if root is not created yet or destroyed
        print(f"Application error: {str(e)}")

if __name__ == "__main__":
    main()